#!/home/kavinjey/.virtualenvs/myvenv/bin/python
from fastapi import FastAPI, File, UploadFile, HTTPException, Request, Header, Depends, Form, Query, BackgroundTasks
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pypdf import PdfReader
import docx
from bs4 import BeautifulSoup
import io  # Import io for BytesIO
from mangum import Mangum
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import List, Optional, Union, Dict, Any
import openai
from openai import OpenAI
import numpy as np
from neo4j import GraphDatabase
from dotenv import load_dotenv
from fastapi.middleware.cors import CORSMiddleware
import os
import json
import asyncio
import httpx
import time
import logging
import hashlib
import sys
from collections import defaultdict
from sanitization import (
    sanitize_user_message, sanitize_chat_id, sanitize_api_key,
    sanitize_text, sanitize_with_length, SanitizedUserMessage,
    SanitizedChatId, SanitizedApiKey, sanitize_request_data,
    sanitize_with_xss_detection, detect_xss, sanitize_filename
)

# ==============================================================================
# Redis Queue Setup for Background File Processing
# ==============================================================================
try:
    from redis import Redis
    from rq import Queue, Worker
    REDIS_AVAILABLE = True
except ImportError:
    REDIS_AVAILABLE = False
    print("‚ö†Ô∏è Redis/RQ not installed. Background processing disabled. Install with: pip install redis rq")

# Load environment variables early for Redis setup
load_dotenv()

# Redis connection setup
REDIS_URL = os.getenv("REDIS_URL", "redis://localhost:6379")
UPSTASH_REDIS_REST_URL = os.getenv("UPSTASH_REDIS_REST_URL")
UPSTASH_REDIS_REST_TOKEN = os.getenv("UPSTASH_REDIS_REST_TOKEN")

# Initialize Redis connection
redis_conn = None
file_queue = None

def get_redis_connection():
    """Get Redis connection, creating it if needed"""
    global redis_conn
    if redis_conn is not None:
        return redis_conn

    if not REDIS_AVAILABLE:
        return None

    try:
        redis_url = REDIS_URL

        # For Upstash and other TLS Redis, convert redis:// to rediss:// for SSL
        is_upstash = redis_url and "upstash.io" in redis_url
        if is_upstash and redis_url.startswith("redis://"):
            redis_url = redis_url.replace("redis://", "rediss://", 1)
            print(f"üîí Using TLS connection for Upstash Redis")

        if redis_url:
            # Use ssl_cert_reqs=None for Upstash/TLS connections to avoid cert verification issues
            if is_upstash or redis_url.startswith("rediss://"):
                import ssl
                redis_conn = Redis.from_url(redis_url, ssl_cert_reqs=None)
            else:
                redis_conn = Redis.from_url(redis_url)
        else:
            redis_conn = Redis(host='localhost', port=6379)

        # Test connection
        redis_conn.ping()
        print(f"‚úÖ Redis connected successfully")
        return redis_conn
    except Exception as e:
        print(f"‚ö†Ô∏è Redis connection failed: {e}")
        return None

def get_file_queue():
    """Get the file processing queue"""
    global file_queue
    if file_queue is not None:
        return file_queue

    conn = get_redis_connection()
    if conn:
        file_queue = Queue("file_ingestion", connection=conn, default_timeout=1800)  # 30 min timeout
        return file_queue
    return None

# Secure Privacy-First API Models
class AppAuthorizationRequest(BaseModel):
    end_user_id: str
    capabilities: List[str]
    redirect_uri: str

class SecureUserQueryRequest(BaseModel):
    question: str
    include_citations: Optional[bool] = True

class SecureUploadRequest(BaseModel):
    filename: str
    file_type: str

class AppUserRequest(BaseModel):
    end_user_id: str
    capabilities: List[str]

class PrivacyFirstQueryRequest(BaseModel):
    end_user_id: str
    question: str
    include_citations: Optional[bool] = False

class DirectUploadRequest(BaseModel):
    end_user_id: str
    filename: str
    file_type: str

class UserAuthVerification(BaseModel):
    user_auth_token: str

class TokenClaims(BaseModel):
    app_id: str
    end_user_id: str
    chat_id: str
    capabilities: List[str]
    exp: float
    iat: float

# Data Models
class ChunkScore(BaseModel):
    chunk_id: str
    chunk_text: str
    score: float

class AnswerWithContext(BaseModel):
    answer: str = ""
    context: List[ChunkScore] = []

class QuestionRequest(BaseModel):
    question: str
    chat_id: str
    selected_model: str = "gpt-4o-mini"  # Default model
    custom_prompt: Union[str, None] = None  # Optional custom system prompt
    temperature: Optional[float] = 0.7  # Default temperature
    max_tokens: Optional[int] = 1000  # Default max tokens
    scope_filters: Optional[Dict[str, Union[str, int, bool]]] = {}  # Optional scope filters
    unhinged_mode: Optional[bool] = False  # Use Grok's unhinged AI model

class CreateNodesAndEmbeddingsRequest(BaseModel):
    pdf_text: str
    pdf_id: str
    chat_id: str
    filename: str
    scope_values: Optional[Dict[str, Union[str, int, bool]]] = {}

# API Models for external access
class ApiQuestionRequest(BaseModel):
    question: str
    selected_model: Optional[str] = None  # Will use chat's saved model if not provided
    custom_prompt: Optional[str] = None   # Will use chat's saved prompt if not provided
    temperature: Optional[float] = None   # Will use chat's saved temperature if not provided
    max_tokens: Optional[int] = None      # Will use chat's saved max_tokens if not provided
    scope_filters: Optional[Dict[str, Union[str, int, bool]]] = {}  # Optional scope filters

class ApiAnswerResponse(BaseModel):
    answer: str
    context: List[ChunkScore]
    chat_id: str
    model: str
    usage: dict

# ==============================================================================
# Custom Scoping System
# ==============================================================================

class ScopeDefinition(BaseModel):
    """Definition of a custom scope field"""
    name: str  # e.g., "playlist_id", "workspace_id", "project_id"
    type: str = "string"  # string, number, boolean
    required: bool = False
    description: Optional[str] = None

class AppScopeConfig(BaseModel):
    """Configuration for custom scopes in an app/chat"""
    scopes: List[ScopeDefinition] = []

class ScopeValues(BaseModel):
    """Values for custom scopes when uploading or querying"""
    values: Dict[str, Union[str, int, bool]] = {}

class CreateNodesWithScopesRequest(BaseModel):
    """Extended request with custom scope support"""
    pdf_text: str
    pdf_id: str
    chat_id: str
    filename: str
    scope_values: Optional[Dict[str, Union[str, int, bool]]] = {}

class ApiQuestionWithScopesRequest(BaseModel):
    """Extended API question request with scope filtering"""
    question: str
    selected_model: Optional[str] = None
    custom_prompt: Optional[str] = None
    temperature: Optional[float] = None
    max_tokens: Optional[int] = None
    scope_filters: Optional[Dict[str, Union[str, int, bool]]] = {}

# In-memory scope configuration storage (in production, this would be in a database)
SCOPE_CONFIGS: Dict[str, AppScopeConfig] = {}

def validate_scope_name(name: str) -> bool:
    """Validate that scope name follows conventions"""
    import re
    # Allow alphanumeric, underscores, and hyphens
    return bool(re.match(r'^[a-zA-Z][a-zA-Z0-9_-]{0,63}$', name))

def validate_scope_value(value: Union[str, int, bool], scope_type: str) -> bool:
    """Validate that a scope value matches its type"""
    if scope_type == "string":
        return isinstance(value, str) and len(value) <= 255
    elif scope_type == "number":
        return isinstance(value, (int, float))
    elif scope_type == "boolean":
        return isinstance(value, bool)
    return False

def sanitize_scope_value(value: Union[str, int, bool]) -> str:
    """Sanitize scope value for Neo4j query"""
    if isinstance(value, str):
        # Sanitize string values for Neo4j
        safe_value = value.replace('\\', '\\\\').replace("'", "\\'").replace('"', '\\"')
        return f"'{safe_value}'"
    elif isinstance(value, bool):
        return "true" if value else "false"
    else:
        return str(value)

def build_scope_properties(scope_values: Dict[str, Union[str, int, bool]],
                          scope_config: Optional[AppScopeConfig] = None,
                          node_var: str = None,
                          for_set_clause: bool = False) -> tuple[str, Dict[str, Union[str, int, bool]]]:
    """Build Neo4j property string for scopes using parameterized queries

    Args:
        scope_values: Dictionary of scope key-value pairs
        scope_config: Optional scope configuration for validation
        node_var: Node variable name (e.g., "d" for Document, "c" for Chunk)
        for_set_clause: If True, format for SET clause (d.key = $param), else for CREATE (key: $param)

    Returns:
        tuple: (property_string, params_dict)
        Example SET: (",\n    d.test = $scope_test", {"scope_test": "dev"})
        Example CREATE: (", test: $scope_test", {"scope_test": "dev"})
    """
    if not scope_values:
        return "", {}

    properties = []
    params = {}

    for key, value in scope_values.items():
        if not validate_scope_name(key):
            continue

        # Validate against config if provided
        if scope_config:
            scope_def = next((s for s in scope_config.scopes if s.name == key), None)
            if scope_def and not validate_scope_value(value, scope_def.type):
                continue

        # Use parameterized query instead of string interpolation
        param_name = f"scope_{key}"

        if for_set_clause and node_var:
            # For SET clause: d.test = $scope_test
            properties.append(f"{node_var}.{key} = ${param_name}")
        else:
            # For CREATE clause: test: $scope_test
            properties.append(f"{key}: ${param_name}")

        params[param_name] = value

    if for_set_clause:
        prop_string = ",\n    " + ",\n    ".join(properties) if properties else ""
    else:
        prop_string = ", " + ", ".join(properties) if properties else ""

    return prop_string, params

def build_scope_where_clause(scope_filters: Dict[str, Union[str, int, bool]],
                             node_var: str = "c",
                             scope_config: Optional[AppScopeConfig] = None) -> str:
    """Build WHERE clause for scope filtering"""
    if not scope_filters:
        return ""

    conditions = []
    for key, value in scope_filters.items():
        if not validate_scope_name(key):
            continue

        # Validate against config if provided
        if scope_config:
            scope_def = next((s for s in scope_config.scopes if s.name == key), None)
            if scope_def and not validate_scope_value(value, scope_def.type):
                continue

        sanitized_value = sanitize_scope_value(value)
        conditions.append(f"{node_var}.{key} = {sanitized_value}")

    return " AND " + " AND ".join(conditions) if conditions else ""

async def get_scope_config(chat_id: str) -> Optional[AppScopeConfig]:
    """Get scope configuration for a chat from Convex"""
    try:
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
        async with httpx.AsyncClient() as client:
            # For V1 subchats, extract app_id and get scope config from parent app
            if chat_id.startswith("subchat_app_"):
                # Parse subchat ID: subchat_app_{app_id}_user_{user_id}_{timestamp}
                import re
                match = re.match(r'subchat_(app_[a-zA-Z0-9_]+)_user_', chat_id)
                if match:
                    app_id = match.group(1)
                    logger.info(f"üîç V1 subchat detected, extracted app_id: {app_id}")

                    # Get app details to find parent chat
                    app_response = await client.post(
                        f"{convex_url}/api/run/app_management/getAppWithSettings",
                        json={
                            "args": {"appId": app_id},
                            "format": "json"
                        },
                        headers={"Content-Type": "application/json"}
                    )

                    if app_response.status_code == 200:
                        app_result = app_response.json()
                        app_data = app_result.get("value")

                        if app_data:
                            parent_chat_id = app_data.get("parentChatId")
                            logger.info(f"üîç Found parent chat ID for app {app_id}: {parent_chat_id}")

                            if parent_chat_id:
                                # Get scope config from parent chat (using exposed endpoint)
                                parent_response = await client.post(
                                    f"{convex_url}/api/run/chats/getChatByIdExposed",
                                    json={
                                        "args": {"id": parent_chat_id},
                                        "format": "json"
                                    },
                                    headers={"Content-Type": "application/json"}
                                )

                                if parent_response.status_code == 200:
                                    parent_result = parent_response.json()
                                    parent_data = parent_result.get("value")

                                    if parent_data:
                                        logger.info(f"üîç Parent chat data keys: {list(parent_data.keys())}")
                                        if "scopeConfig" in parent_data:
                                            scope_data = parent_data["scopeConfig"]
                                            logger.info(f"‚úÖ Found scope config from parent chat {parent_chat_id}: {scope_data}")
                                            return AppScopeConfig(**scope_data)
                                        else:
                                            logger.info(f"‚ö†Ô∏è  Parent chat {parent_chat_id} has no scopeConfig field")
                                    else:
                                        logger.info(f"‚ö†Ô∏è  Parent chat {parent_chat_id} returned null/empty data")

            # Standard flow for regular chats (using exposed endpoint)
            response = await client.post(
                f"{convex_url}/api/run/chats/getChatByIdExposed",
                json={
                    "args": {"id": chat_id},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            if response.status_code == 200:
                result = response.json()
                chat_data = result.get("value")

                logger.info(f"üîç Chat data for {chat_id}: has_data={chat_data is not None}, has_scopeConfig={'scopeConfig' in chat_data if chat_data else False}")

                if chat_data and "scopeConfig" in chat_data:
                    scope_data = chat_data["scopeConfig"]
                    logger.info(f"‚úÖ Found scope config directly on chat {chat_id}: {scope_data}")
                    return AppScopeConfig(**scope_data)

                # If this is a subchat and no scope config found, try parent chat
                if chat_id.startswith("subchat_") and chat_data:
                    parent_chat_id = chat_data.get("parentChatId")
                    logger.info(f"üîç Subchat detected, parent_chat_id={parent_chat_id}")
                    if parent_chat_id:
                        logger.info(f"üîç Subchat {chat_id} has no scope config, checking parent {parent_chat_id}")
                        parent_response = await client.post(
                            f"{convex_url}/api/run/chats/getChatByIdExposed",
                            json={
                                "args": {"id": parent_chat_id},
                                "format": "json"
                            },
                            headers={"Content-Type": "application/json"}
                        )

                        if parent_response.status_code == 200:
                            parent_result = parent_response.json()
                            parent_data = parent_result.get("value")
                            if parent_data and "scopeConfig" in parent_data:
                                scope_data = parent_data["scopeConfig"]
                                logger.info(f"‚úÖ Inherited scope config from parent chat: {list(scope_data.get('scopes', []))}")
                                return AppScopeConfig(**scope_data)
    except Exception as e:
        logger.warning(f"Could not load scope config for chat {chat_id}: {e}")

    # Check in-memory cache as fallback
    return SCOPE_CONFIGS.get(chat_id)

async def save_scope_config(chat_id: str, scope_config: AppScopeConfig) -> bool:
    """Save scope configuration for a chat"""
    try:
        # Store in memory (in production, this would update Convex)
        SCOPE_CONFIGS[chat_id] = scope_config

        # Also update in Convex
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"{convex_url}/api/run/chats/updateChatScopeConfig",
                json={
                    "args": {
                        "chatId": chat_id,
                        "scopeConfig": scope_config.dict()
                    },
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            return response.status_code == 200
    except Exception as e:
        logger.error(f"Failed to save scope config: {e}")
        return False

app = FastAPI(
    title="Trainly API with V1 Trusted Issuer Authentication",
    description="Privacy-first GraphRAG backend with user-controlled authentication",
    version="1.0.0"
)
handler = Mangum(app)

# Configure CORS middleware (adjust origins as needed for production)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Change this in production!
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# NOTE: Duplicate detection removed - users can upload the same file multiple times

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# V1 Trusted Issuer Authentication - Integrated directly into main app
# This replaces the separate router files for simplicity

# Additional imports for V1 auth
import requests
from cachetools import TTLCache
from typing import Dict, List, Optional, Any
from datetime import datetime, timedelta

# V1 Auth: JWKS cache and app configuration storage
JWKS_CACHE = TTLCache(maxsize=100, ttl=3600)  # 1 hour TTL
V1_APP_CONFIGS = {}  # In production, this would be a database

# V1 Auth: App configuration models
class V1AppConfig:
    def __init__(self, app_id: str, issuer: str, allowed_audiences: List[str],
                 alg_allowlist: List[str] = None, jwks_uri: str = None):
        self.app_id = app_id
        self.issuer = issuer
        self.jwks_uri = jwks_uri
        self.allowed_audiences = allowed_audiences
        self.alg_allowlist = alg_allowlist or ["RS256", "ES256"]
        self.created_at = datetime.utcnow()

# V1 Auth: Helper functions
async def discover_jwks_uri(issuer: str) -> str:
    """Discover JWKS URI from OIDC well-known configuration"""
    well_known_url = f"{issuer.rstrip('/')}/.well-known/openid-configuration"

    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get(well_known_url)
            response.raise_for_status()
            config = response.json()
            jwks_uri = config.get("jwks_uri")
            if not jwks_uri:
                raise Exception(f"No jwks_uri found in {well_known_url}")
            return jwks_uri
    except Exception as e:
        logger.error(f"Failed to discover JWKS URI for {issuer}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to discover JWKS URI: {str(e)}")

async def get_jwks(jwks_uri: str) -> Dict[str, Any]:
    """Fetch and cache JWKS (JSON Web Key Set)"""
    cache_key = f"jwks_{hashlib.md5(jwks_uri.encode()).hexdigest()}"

    if cache_key in JWKS_CACHE:
        return JWKS_CACHE[cache_key]

    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get(jwks_uri)
            response.raise_for_status()
            jwks = response.json()
            JWKS_CACHE[cache_key] = jwks
            logger.info(f"Fetched and cached JWKS from {jwks_uri}")
            return jwks
    except Exception as e:
        logger.error(f"Failed to fetch JWKS from {jwks_uri}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch JWKS: {str(e)}")

def find_key_by_kid(jwks: Dict[str, Any], kid: str) -> Optional[Dict[str, Any]]:
    """Find a specific key by kid (key ID) in JWKS"""
    keys = jwks.get("keys", [])
    for key in keys:
        if key.get("kid") == kid:
            return key
    return None

async def verify_v1_jwt_token(token: str, app_config: V1AppConfig) -> Dict[str, Any]:
    """Verify JWT token against app's OIDC configuration"""
    try:
        # Decode header to get kid and alg
        unverified_header = jwt.get_unverified_header(token)
        kid = unverified_header.get("kid")
        alg = unverified_header.get("alg")

        # Security checks
        if alg not in app_config.alg_allowlist:
            raise HTTPException(status_code=401, detail=f"Algorithm {alg} not allowed")
        if alg == "none":
            raise HTTPException(status_code=401, detail="Algorithm 'none' is not allowed")

        # Handle dynamic issuer and audience detection
        actual_issuer = app_config.issuer
        actual_audiences = app_config.allowed_audiences

        if app_config.issuer == "DYNAMIC" or app_config.allowed_audiences == ["DYNAMIC"]:
            # Extract issuer and audience from token payload without verification
            unverified_payload = jwt.decode(token, key="", algorithms=["RS256"], options={"verify_signature": False})

            if app_config.issuer == "DYNAMIC":
                actual_issuer = unverified_payload.get("iss")
                if not actual_issuer:
                    raise HTTPException(status_code=401, detail="Token missing issuer claim")
                logger.info(f"üîç Dynamic issuer detected: {actual_issuer}")

            if app_config.allowed_audiences == ["DYNAMIC"]:
                token_audience = unverified_payload.get("aud")
                if token_audience:
                    # Handle both single audience (string) and multiple audiences (array)
                    if isinstance(token_audience, str):
                        actual_audiences = [token_audience]
                    else:
                        actual_audiences = token_audience
                    logger.info(f"üîç Dynamic audience detected: {actual_audiences}")
                else:
                    # Some tokens don't have explicit audience, so we'll be flexible
                    actual_audiences = None
                    logger.info(f"üîç No audience in token, skipping audience validation")

        # Get JWKS URI
        jwks_uri = app_config.jwks_uri
        if not jwks_uri:
            jwks_uri = await discover_jwks_uri(actual_issuer)

        # Fetch JWKS and find key
        jwks = await get_jwks(jwks_uri)
        key = find_key_by_kid(jwks, kid) if kid else None

        if not key:
            # Refresh cache and try again (handle key rotation)
            cache_key = f"jwks_{hashlib.md5(jwks_uri.encode()).hexdigest()}"
            if cache_key in JWKS_CACHE:
                del JWKS_CACHE[cache_key]
            jwks = await get_jwks(jwks_uri)
            key = find_key_by_kid(jwks, kid) if kid else None

            if not key:
                raise HTTPException(status_code=401, detail=f"Key with kid '{kid}' not found")

        # Convert JWK to PEM and verify token
        public_key = jwt.algorithms.RSAAlgorithm.from_jwk(key)

        # Prepare JWT decode options
        decode_options = {
            "verify_exp": True,
            "verify_iat": True,
            "verify_signature": True,
            "verify_aud": actual_audiences is not None,  # Skip audience verification if None
            "verify_iss": True,
            "leeway": 300  # 5 minute clock skew tolerance
        }

        # Prepare JWT decode call with correct PyJWT syntax
        if actual_audiences is not None:
            decoded_token = jwt.decode(
                token,
                public_key,
                algorithms=[alg],
                issuer=actual_issuer,
                audience=actual_audiences,
                options=decode_options
            )
        else:
            decoded_token = jwt.decode(
                token,
                public_key,
                algorithms=[alg],
                issuer=actual_issuer,
                options=decode_options
            )

        return decoded_token

    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token has expired")
    except jwt.InvalidTokenError as e:
        raise HTTPException(status_code=401, detail=f"Invalid token: {str(e)}")
    except Exception as e:
        logger.error(f"JWT verification failed: {str(e)}")
        raise HTTPException(status_code=401, detail=f"Token verification failed: {str(e)}")

def derive_v1_user_identity(token_claims: Dict[str, Any], app_id: str) -> Dict[str, str]:
    """Derive internal user identity from verified token claims"""
    external_user_id = token_claims.get("sub")
    if not external_user_id:
        raise HTTPException(status_code=401, detail="Token missing 'sub' claim")

    # Generate deterministic internal user_id and chat_id
    user_id_input = f"{app_id}::{external_user_id}"
    user_id = f"user_v1_{hashlib.sha256(user_id_input.encode()).hexdigest()[:16]}"

    chat_id_input = f"{app_id}::{external_user_id}::chat"
    chat_id = f"chat_v1_{hashlib.sha256(chat_id_input.encode()).hexdigest()[:16]}"

    return {
        "app_id": app_id,
        "external_user_id": external_user_id,
        "user_id": user_id,
        "chat_id": chat_id,
        "iss": token_claims["iss"],
        "aud": token_claims.get("aud", "none")  # Handle missing audience gracefully
    }

async def get_app_config_from_convex(app_id: str) -> Optional[V1AppConfig]:
    """Get app configuration from Convex system for V1 auth"""
    try:
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
        logger.info(f"üîç Checking Convex for app {app_id} at {convex_url}")

        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"{convex_url}/api/run/app_management/getAppWithSettings",
                json={
                    "args": {"appId": app_id},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            logger.info(f"üîç Convex response: {response.status_code} - {response.text[:200]}")

            if response.status_code == 200:
                result = response.json()
                app_data = result.get("value")

                logger.info(f"üîç App data: {app_data}")

                if app_data and app_data.get("isActive"):
                    # Check if API access is disabled for this app
                    if app_data.get("isApiDisabled", False):
                        logger.warning(f"üö´ App {app_id} has API access disabled")
                        raise HTTPException(
                            status_code=403,
                            detail=f"API access is disabled for this app. Please contact the app developer to enable API access."
                        )

                    logger.info(f"‚úÖ Found active app {app_id} in Convex, creating dynamic V1 config")
                    # For Convex apps, we'll determine the issuer and audience dynamically from the JWT token
                    # This allows users to use any OAuth provider without configuration
                    return V1AppConfig(
                        app_id=app_id,
                        issuer="DYNAMIC",  # Special marker for dynamic issuer detection
                        allowed_audiences=["DYNAMIC"],  # Special marker for dynamic audience detection
                        alg_allowlist=["RS256"]
                    )
                else:
                    logger.warning(f"‚ö†Ô∏è App {app_id} found but not active or missing data")
            else:
                logger.warning(f"‚ö†Ô∏è Convex returned non-200 status: {response.status_code}")
    except HTTPException:
        # Re-raise HTTP exceptions (like API disabled)
        raise
    except Exception as e:
        logger.error(f"Failed to get app config from Convex: {e}")

    return None

async def get_parent_chat_id_from_app(app_id: str) -> Optional[str]:
    """Get parent chat ID from app ID"""
    try:
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"{convex_url}/api/run/app_management/getAppWithSettings",
                json={
                    "args": {"appId": app_id},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            if response.status_code == 200:
                result = response.json()
                app_data = result.get("value")

                if app_data:
                    # The parentChatId is already the chat string ID we need for Neo4j
                    parent_chat_id = app_data.get("parentChatId")
                    logger.info(f"üîç Parent chat ID from app data: {parent_chat_id}")
                    if parent_chat_id:
                        logger.info(f"üîç Found parent chat: {parent_chat_id} for app {app_id}")
                        return parent_chat_id
                    else:
                        logger.info(f"‚ÑπÔ∏è App {app_id} has no parentChatId configured")
                else:
                    logger.warning(f"‚ö†Ô∏è No app data returned for app {app_id}")
    except Exception as e:
        logger.error(f"Failed to get parent chat ID from app {app_id}: {e}")

    return None

async def authenticate_v1_user(authorization: str, app_id: str) -> Dict[str, str]:
    """Main V1 authentication: validate OAuth ID token and derive user identity"""
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid Authorization header")

    token = authorization[7:]  # Remove "Bearer " prefix

    # Get app configuration - first try V1 registry, then Convex system
    app_config = V1_APP_CONFIGS.get(app_id)
    if not app_config:
        # Try to get configuration from Convex app system
        app_config = await get_app_config_from_convex(app_id)
        if not app_config:
            raise HTTPException(status_code=404, detail=f"App {app_id} not found in V1 registry or Convex system")

    # Verify JWT token
    token_claims = await verify_v1_jwt_token(token, app_config)

    # Derive user identity
    user_identity = derive_v1_user_identity(token_claims, app_id)

    logger.info(f"V1 Auth: Authenticated user {user_identity['user_id']} for app {app_id}")
    return user_identity

# Privacy-First API Helper Functions
import jwt
import time
import hashlib

# Temporary storage for OAuth authorization codes (use Redis in production)
auth_requests = {}

async def verify_app_secret(api_key: str) -> Dict[str, Any]:
    """
    Verify app secret for server-to-server calls.
    App secrets can create users and mint tokens, but CANNOT read raw data.
    """
    # Sanitize the secret
    sanitized_secret = sanitize_api_key(api_key)
    if not sanitized_secret:
        raise HTTPException(status_code=401, detail="Invalid app secret format")

    # Verify with Convex apps table
    try:
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

        # Query Convex to verify app secret
        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"{convex_url}/api/run/app_management/verifyAppSecret",
                json={
                    "args": {"appSecret": sanitized_secret},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            if response.status_code != 200:
                raise HTTPException(status_code=401, detail="Invalid app secret")

            result = response.json()

            # Check if Convex returned an error
            if result.get("status") == "error":
                # Convex is having issues, use fallback
                if sanitized_secret == "as_demo_secret_123":
                    return {
                        "appId": "app_demo_123",
                        "developerId": "dev_demo",
                        "isActive": True,
                        "allowedCapabilities": ["ask", "upload"]
                    }
                # Removed hardcoded fallback - using Convex system instead
                raise HTTPException(status_code=401, detail="Unable to verify app secret - Convex error")

            app_data = result.get("value")

            if not app_data or not app_data.get("isActive"):
                raise HTTPException(status_code=401, detail="App not found or inactive")

            return {
                "appId": app_data["appId"],
                "developerId": app_data["developerId"],
                "isActive": app_data["isActive"],
                "allowedCapabilities": app_data.get("settings", {}).get("allowedCapabilities", ["ask", "upload"]),
                "parentChatSettings": app_data.get("parentChatSettings")
            }

    except httpx.RequestError:
        # Fallback for development - allow hardcoded demo secret
        if sanitized_secret == "as_demo_secret_123":
            return {
                "appId": "app_demo_123",
                "developerId": "dev_demo",
                "isActive": True,
                "allowedCapabilities": ["ask", "upload"]
            }
        raise HTTPException(status_code=401, detail="Unable to verify app secret")

async def verify_scoped_token(token: str) -> TokenClaims:
    """
    Verify a scoped JWT token for user-specific operations.
    These tokens are short-lived and scoped to specific (app_id, user_id, chat_id).
    """
    try:
        secret_key = os.getenv("JWT_SECRET_KEY", "your-secret-key")
        payload = jwt.decode(token, secret_key, algorithms=["HS256"])

        claims = TokenClaims(**payload)

        # Check expiration
        if claims.exp < time.time():
            raise HTTPException(status_code=401, detail="Token expired")

        return claims

    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

def generate_scoped_token(app_id: str, end_user_id: str, chat_id: str, capabilities: List[str]) -> str:
    """Generate a short-lived scoped token for specific user operations"""
    secret_key = os.getenv("JWT_SECRET_KEY", "your-secret-key")

    payload = {
        "app_id": app_id,
        "end_user_id": end_user_id,
        "chat_id": chat_id,
        "capabilities": capabilities,
        "iat": time.time(),
        "exp": time.time() + (15 * 60)  # 15 minutes expiry
    }

    return jwt.encode(payload, secret_key, algorithm="HS256")

async def get_or_create_user_subchat(app_id: str, end_user_id: str) -> Dict[str, Any]:
    """Get or create a private sub-chat for a user under an app"""
    convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

    try:
        # First, check if user already has a subchat for this app
        async with httpx.AsyncClient() as client:
            # Query existing user subchat
            check_response = await client.post(
                f"{convex_url}/api/run/app_management/getUserSubChat",
                json={
                    "args": {
                        "appId": app_id,
                        "endUserId": end_user_id
                    },
                    "format": "json"
                },
                headers={"Content-Type": "application/json"},
                timeout=10.0
            )

            if check_response.status_code == 200:
                check_result = check_response.json()
                existing_subchat = check_result.get("value")

                if existing_subchat:
                    # User already has a subchat - return existing
                    logger.info(f"üë§ Returning existing subchat for user {end_user_id[:8]}... in app {app_id}")
                    logger.debug(f"üîç Existing subchat structure: {existing_subchat}")

                    # Handle different possible response structures
                    chat_string_id = existing_subchat.get("chatStringId") or existing_subchat.get("chatId")
                    if not chat_string_id:
                        logger.error(f"‚ùå No chatStringId or chatId found in subchat response: {existing_subchat}")
                        raise ValueError("Invalid subchat response structure")

                    return {
                        "userChatId": existing_subchat["userChatId"],
                        "chatId": chat_string_id,
                        "chatStringId": chat_string_id,
                        "isNew": False,  # Existing user
                        "capabilities": existing_subchat.get("capabilities", ["ask", "upload"])
                    }

            # User doesn't have a subchat yet - create new one
            logger.info(f"üÜï Creating new subchat for user {end_user_id[:8]}... in app {app_id}")

            create_response = await client.post(
                f"{convex_url}/api/run/app_management/createUserSubChat",
                json={
                    "args": {
                        "appId": app_id,
                        "endUserId": end_user_id,
                        "capabilities": ["ask", "upload"]
                    },
                    "format": "json"
                },
                headers={"Content-Type": "application/json"},
                timeout=10.0
            )

            if create_response.status_code == 200:
                create_result = create_response.json()
                new_subchat = create_result.get("value")

                if new_subchat:
                    logger.debug(f"üîç New subchat structure: {new_subchat}")

                    # Handle different possible response structures
                    chat_string_id = new_subchat.get("chatStringId") or new_subchat.get("chatId")
                    if not chat_string_id:
                        logger.error(f"‚ùå No chatStringId or chatId found in new subchat response: {new_subchat}")
                        raise ValueError("Invalid new subchat response structure")

                    return {
                        "userChatId": new_subchat["userChatId"],
                        "chatId": chat_string_id,
                        "chatStringId": chat_string_id,
                        "isNew": new_subchat.get("isNew", True),
                        "capabilities": new_subchat.get("capabilities", ["ask", "upload"])
                    }

            # Fallback if Convex calls fail
            logger.warning(f"‚ö†Ô∏è Convex calls failed, using fallback for user {end_user_id[:8]}...")
            user_chat_id = f"subchat_{app_id}_{end_user_id}_{int(time.time())}"
            return {
                "userChatId": f"uc_{app_id}_{end_user_id}",
                "chatId": user_chat_id,
                "chatStringId": user_chat_id,
                "isNew": True,  # Assume new since we can't verify
                "capabilities": ["ask", "upload"]
            }

    except Exception as e:
        logger.error(f"Error in get_or_create_user_subchat: {e}")
        # Fallback to basic implementation if there's an error
        user_chat_id = f"subchat_{app_id}_{end_user_id}_{int(time.time())}"
        return {
            "userChatId": f"uc_{app_id}_{end_user_id}",
            "chatId": user_chat_id,
            "chatStringId": user_chat_id,
            "isNew": True,  # Assume new since we can't verify
            "capabilities": ["ask", "upload"]
        }

async def log_app_access(
    app_id: str,
    end_user_id: str,
    chat_id: str,
    action: str,
    capability: str,
    allowed: bool,
    request: Request = None,
    **metadata
):
    """Log all app access attempts for audit trail"""
    try:
        # In production, this would call Convex to log the access
        logger.info(f"App Access: {app_id} | User: {end_user_id} | Action: {action} | Allowed: {allowed}")
    except Exception as e:
        logger.error(f"Failed to log app access: {e}")

async def track_subchat_creation(app_id: str, end_user_id: str, chat_id: str):
    """Track subchat creation for analytics"""
    try:
        # In production, this would call Convex to update chat metadata
        # For demonstration, we'll call a Convex function to track this
        convex_base_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
        convex_url = f"{convex_base_url}/api/run/chat_analytics/trackSubchatCreation"

        async with httpx.AsyncClient() as client:
            await client.post(
                convex_url,
                json={
                    "args": {
                        "appId": app_id,
                        "endUserId": end_user_id,
                        "chatId": chat_id
                    },
                    "format": "json"
                },
                timeout=5.0
            )

        logger.info(f"üìä Analytics: New subchat created | App: {app_id} | User: {end_user_id[:8]}...")
    except Exception as e:
        logger.error(f"Failed to track subchat creation: {e}")

async def track_file_upload(app_id: str, end_user_id: str, filename: str, file_size: int, chat_id: str, file_hash: str = None):
    """Track file upload for analytics"""
    try:
        # Call Convex to update metadata
        convex_base_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
        convex_url = f"{convex_base_url}/api/run/chat_analytics/trackFileUpload"

        args = {
            "appId": app_id,
            "endUserId": end_user_id,
            "filename": filename,
            "fileSize": file_size,
            "chatId": chat_id  # This should be the Convex ID, not the string ID
        }

        # Add file hash if provided (for deduplication)
        if file_hash:
            args["fileHash"] = file_hash

        async with httpx.AsyncClient() as client:
            await client.post(
                convex_url,
                json={
                    "args": args,
                    "format": "json"
                },
                timeout=5.0
            )

        logger.info(f"üìä Analytics: File uploaded | App: {app_id} | Size: {format_bytes(file_size)} | Type: {get_file_type(filename)}")
    except Exception as e:
        logger.error(f"Failed to track file upload: {e}")

async def track_file_deletion(app_id: str, end_user_id: str, filename: str, file_size: int, chat_id: str):
    """Track file deletion for analytics"""
    try:
        # Call Convex to update metadata and decrease storage counts
        convex_base_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
        convex_url = f"{convex_base_url}/api/run/chat_analytics/trackFileDeletion"

        args = {
            "appId": app_id,
            "endUserId": end_user_id,
            "filename": filename,
            "fileSize": file_size,
            "chatId": chat_id
        }

        async with httpx.AsyncClient() as client:
            await client.post(
                convex_url,
                json={
                    "args": args,
                    "format": "json"
                },
                timeout=5.0
            )

        logger.info(f"üìä Analytics: File deleted | App: {app_id} | Size: {format_bytes(file_size)} freed | Type: {get_file_type(filename)}")
    except Exception as e:
        logger.error(f"Failed to track file deletion: {e}")

async def track_upload_analytics(chat_id: str, filename: str, pdf_text: str, actual_file_size: int = None):
    """Track upload analytics by extracting info from chat_id and determining if it's a sub-chat"""
    try:
        logger.info(f"üîç Tracking analytics for chat_id: {chat_id}, filename: {filename}")

        # Use actual file size if provided, otherwise estimate from text length
        if actual_file_size is not None:
            file_size = actual_file_size
            logger.info(f"üìè Using actual file size: {format_bytes(file_size)}")
        else:
            file_size = len(pdf_text.encode('utf-8'))
            logger.info(f"üìè Estimated file size from text: {format_bytes(file_size)}")

        # Calculate tokens from extracted text (same approximation as frontend: Math.ceil(text.length / 4))
        import math
        tokens_ingested = math.ceil(len(pdf_text) / 4)
        logger.info(f"üî¢ Calculated tokens from extracted text: {tokens_ingested:,} tokens")

        # Check if this is a sub-chat by querying Convex
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
        logger.info(f"üåê Querying Convex at: {convex_url}")

        async with httpx.AsyncClient() as client:
            # Get chat details to see if it's a sub-chat (using chatId string field)
            chat_response = await client.post(
                f"{convex_url}/api/run/backend_credits/getChatWithApp",
                json={
                    "args": {"chatId": chat_id},
                    "format": "json"
                },
                timeout=10.0
            )

            logger.info(f"üì° Convex response status: {chat_response.status_code}")

            if chat_response.status_code == 200:
                chat_data = chat_response.json()
                chat = chat_data.get("value")
                logger.info(f"üí¨ Chat data retrieved: chatType={chat.get('chatType') if chat else None}, parentAppId={chat.get('parentAppId') if chat else None}")

                # Determine userId for token tracking
                # For subchats (API uploads), track against developer's account (parent chat owner)
                # For regular chats, track against chat owner
                user_id_for_tokens = None

                if chat and chat.get("chatType") == "app_subchat" and chat.get("parentAppId"):
                    # This is a sub-chat! Extract the necessary info
                    app_id = chat.get("parentAppId")
                    end_user_id = chat.get("userId")  # The end user who owns the sub-chat
                    convex_chat_id = chat.get("_id")  # Get the Convex document ID
                    logger.info(f"üéØ Sub-chat detected! app_id={app_id}, end_user_id={end_user_id[:8] if end_user_id else None}..., convex_id={convex_chat_id}")

                    if app_id and end_user_id:
                        # Call the existing tracking function with string chat ID
                        await track_file_upload(
                            app_id=app_id,
                            end_user_id=end_user_id,
                            filename=filename,
                            file_size=file_size,
                            chat_id=chat_id  # Use string chat ID (Convex function now handles both)
                        )
                        logger.info(f"‚úÖ Sub-chat upload tracked: {filename} in {chat_id}")

                        # For token tracking, get the developer's userId (parent chat owner)
                        # Get parent chat ID from app
                        parent_chat_id = await get_parent_chat_id_from_app(app_id)
                        if parent_chat_id:
                            # Get parent chat to get developer's userId
                            parent_chat_response = await client.post(
                                f"{convex_url}/api/run/backend_credits/getChatWithApp",
                                json={
                                    "args": {"chatId": parent_chat_id},
                                    "format": "json"
                                },
                                timeout=10.0
                            )
                            if parent_chat_response.status_code == 200:
                                parent_chat_data = parent_chat_response.json()
                                parent_chat = parent_chat_data.get("value")
                                if parent_chat:
                                    user_id_for_tokens = parent_chat.get("userId")
                                    logger.info(f"üîç Found developer userId for token tracking: {user_id_for_tokens[:8] if user_id_for_tokens else None}...")
                                else:
                                    logger.warning(f"‚ö†Ô∏è  Parent chat {parent_chat_id} not found")
                            else:
                                logger.warning(f"‚ö†Ô∏è  Failed to get parent chat: {parent_chat_response.status_code}")
                        else:
                            logger.warning(f"‚ö†Ô∏è  No parent chat ID found for app {app_id}")

                    else:
                        logger.warning(f"‚ùå Sub-chat missing required fields - app_id: {app_id}, user_id: {end_user_id}")
                else:
                    # Not a sub-chat, use chat owner's userId for token tracking
                    user_id_for_tokens = chat.get("userId") if chat else None
                    logger.info(f"‚ÑπÔ∏è  Regular chat upload - tracking tokens for chat owner: {user_id_for_tokens[:8] if user_id_for_tokens else None}...")

                # Track token ingestion for the appropriate user
                # For subchats: developer's account (parent chat owner)
                # For regular chats: chat owner
                if user_id_for_tokens:
                    try:
                        token_tracking_response = await client.post(
                            f"{convex_url}/api/run/fileStorage/addTokenIngestion",
                            json={
                                "args": {
                                    "userId": user_id_for_tokens,
                                    "tokens": tokens_ingested
                                },
                                "format": "json"
                            },
                            timeout=10.0
                        )
                        if token_tracking_response.status_code == 200:
                            logger.info(f"‚úÖ Token ingestion tracked: {tokens_ingested:,} tokens ({tokens_ingested // 500} KU) for user {user_id_for_tokens[:8]}...")
                        else:
                            logger.warning(f"‚ö†Ô∏è  Failed to track token ingestion: {token_tracking_response.status_code}")
                            if token_tracking_response.status_code != 200:
                                response_text = await token_tracking_response.text()
                                logger.warning(f"Response body: {response_text[:500]}")
                    except Exception as token_error:
                        logger.error(f"‚ùå Error tracking token ingestion: {token_error}")
                        import traceback
                        logger.error(f"Traceback: {traceback.format_exc()}")
                        # Don't fail the upload if token tracking fails
                else:
                    logger.warning(f"‚ö†Ô∏è  No user_id found for token tracking, skipping token ingestion tracking")
            else:
                logger.warning(f"‚ùå Failed to get chat details for analytics tracking: {chat_response.status_code}")
                if chat_response.status_code != 200:
                    response_text = await chat_response.text()
                    logger.warning(f"Response body: {response_text[:500]}")

    except Exception as e:
        logger.error(f"Failed to track upload analytics for {chat_id}: {e}")
        # Don't re-raise - this should not break the upload process

async def track_api_query(app_id: str, end_user_id: str, response_time: float, success: bool):
    """Track API query for performance analytics"""
    try:
        # Call Convex to update metadata
        convex_base_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
        convex_url = f"{convex_base_url}/api/run/chat_analytics/trackApiQuery"

        async with httpx.AsyncClient() as client:
            await client.post(
                convex_url,
                json={
                    "args": {
                        "appId": app_id,
                        "endUserId": end_user_id,
                        "responseTime": response_time,
                        "success": success
                    },
                    "format": "json"
                },
                timeout=5.0
            )

        logger.info(f"üìä Analytics: Query tracked | App: {app_id} | Time: {response_time:.0f}ms | Success: {success}")
    except Exception as e:
        logger.error(f"Failed to track API query: {e}")

def format_bytes(bytes_val: int) -> str:
    """Format bytes into human readable string"""
    if bytes_val == 0:
        return "0 B"

    sizes = ["B", "KB", "MB", "GB", "TB"]
    i = 0
    while bytes_val >= 1024 and i < len(sizes) - 1:
        bytes_val /= 1024
        i += 1

    return f"{bytes_val:.1f} {sizes[i]}"

def get_file_type(filename: str) -> str:
    """Get file type category from filename"""
    ext = filename.lower().split('.')[-1] if '.' in filename else ''

    if ext in ['pdf']:
        return 'pdf'
    elif ext in ['doc', 'docx']:
        return 'docx'
    elif ext in ['txt', 'md', 'csv', 'json']:
        return 'txt'
    elif ext in ['jpg', 'jpeg', 'png', 'gif', 'webp']:
        return 'images'
    else:
        return 'other'

print("‚úÖ Privacy-First API functions loaded successfully")

# Security
security = HTTPBearer()

# Simple rate limiting (in-memory)
request_counts = {}
RATE_LIMIT_WINDOW = 60  # seconds
RATE_LIMIT_MAX = 60  # requests per window

# Increased to 500MB to match highest tier limit (scale/enterprise)
# Frontend enforces tier-specific limits (50MB free, 200MB pro, 500MB scale/enterprise)
MAX_FILE_SIZE = 500 * 1024 * 1024  # 500MB
PEEK_BYTES = 1024  # Larger peek window for better type detection
EXTRACT_SEM = asyncio.Semaphore(8)  # Module-level semaphore for backpressure

class ReadFiles:
    def __init__(self):
        self.supported_file_types = [
            ".pdf", ".docx", ".txt", ".md", ".ts", ".csv", ".json", ".html",
            ".xml", ".yaml", ".yml", ".js", ".py", ".java", ".cpp", ".c", ".h",
            ".cs", ".php", ".rb", ".sh", ".bat", ".ps1", ".psm1", ".psd1",
            ".ps1xml", ".pssc"
        ]
        # Map file extensions to their respective extraction methods
        self.extractors = {
            ".pdf": self.extract_text_from_pdf,
            ".docx": self.extract_text_from_docx,
            ".txt": self.extract_text_from_text,
            ".md": self.extract_text_from_text,
            ".ts": self.extract_text_from_text,
            ".csv": self.extract_text_from_text,
            ".json": self.extract_text_from_text,
            ".html": self.extract_text_from_html,
            ".xml": self.extract_text_from_text,
            ".yaml": self.extract_text_from_text,
            ".yml": self.extract_text_from_text,
            ".js": self.extract_text_from_text,
            ".py": self.extract_text_from_text,
            ".java": self.extract_text_from_text,
            ".cpp": self.extract_text_from_text,
            ".c": self.extract_text_from_text,
            ".h": self.extract_text_from_text,
            ".cs": self.extract_text_from_text,
            ".php": self.extract_text_from_text,
            ".rb": self.extract_text_from_text,
            ".sh": self.extract_text_from_text,
            ".bat": self.extract_text_from_text,
            ".ps1": self.extract_text_from_text,
            ".psm1": self.extract_text_from_text,
            ".psd1": self.extract_text_from_text,
            ".ps1xml": self.extract_text_from_text,
            ".pssc": self.extract_text_from_text,
        }

    def get_file_size(self, upload_file: UploadFile) -> int:
        current_pos = upload_file.file.tell()
        upload_file.file.seek(0, 2)   # Go to the end of the file
        size = upload_file.file.tell()
        upload_file.file.seek(current_pos)  # Go back to original position
        return size

    def extract_text_from_pdf(self, file: UploadFile):
        try:
            # file_size = self.get_file_size(file)
            # if file_size > MAX_FILE_SIZE:
            #     raise HTTPException(status_code=413, detail="File too large (max 5 MB).")
            file.file.seek(0)
            reader = PdfReader(file.file)
            text = ""
            for page in reader.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + "\n"
            return text
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing PDF file: {str(e)}")
        finally:
            file.file.close()

    def extract_text_from_docx(self, file: UploadFile):
        try:
            # if file.size() > MAX_FILE_SIZE:
            #     raise HTTPException(status_code=413, detail="File too large (max 5 MB).")
            # Read the entire file into bytes
            file.file.seek(0)
            data = file.file.read()

            # Use BytesIO to create a file-like object
            doc = docx.Document(io.BytesIO(data))

            # Extract text from paragraphs
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing DOCX file: {str(e)}")
        finally:
            file.file.close()

    def extract_text_from_html(self, file: UploadFile):
        try:
            # if file.size() > MAX_FILE_SIZE:
            #     raise HTTPException(status_code=413, detail="File too large (max 5 MB).")
            file.file.seek(0)
            content = file.file.read().decode('utf-8', errors='ignore')
            soup = BeautifulSoup(content, "html.parser")
            text = soup.get_text(separator='\n')
            return text
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing HTML file: {str(e)}")
        finally:
            file.file.close()

    def extract_text_from_text(self, file: UploadFile):
        try:
            # if file.size() > MAX_FILE_SIZE:
            #     raise HTTPException(status_code=413, detail="File too large (max 5 MB).")
            file.file.seek(0)
            # Attempt to detect encoding; default to utf-8
            content = file.file.read().decode('utf-8', errors='ignore')
            print("HERE", content)
            return content
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing text file: {str(e)}")
        finally:
            file.file.close()

    def extract_text(self, file: UploadFile):
        # Identify the file extension
        # if file.size() > MAX_FILE_SIZE:
        #         raise HTTPException(status_code=413, detail="File too large (max 5 MB).")
        filename = file.filename.lower()
        matched_extension = None
        for ext in self.supported_file_types:
            if filename.endswith(ext):
                matched_extension = ext
                break

        if not matched_extension:
            raise HTTPException(status_code=400, detail="Unsupported file type.")

        extractor = self.extractors.get(matched_extension)
        if extractor:
            return extractor(file)
        else:
            raise HTTPException(status_code=400, detail=f"No extractor available for file type: {matched_extension}")

read_files = ReadFiles()

def _detect_file_type(head: bytes, filename: str) -> str:
    """
    Detect file type using magic bytes and filename.
    Uses larger peek window for accurate detection.
    """
    fn = filename.lower()

    # PDF magic bytes
    if head.startswith(b"%PDF-"):
        return "pdf"

    # DOCX magic bytes (ZIP signature + .docx extension)
    if head.startswith(b"PK\x03\x04") and fn.endswith(".docx"):
        return "docx"

    # HTML detection (with whitespace handling)
    h = head.lstrip()
    if h.startswith(b"<!DOCTYPE") or h.startswith(b"<html"):
        return "html"

    # Text detection - use full sample, not just first 8 bytes
    sample = head[:PEEK_BYTES] if len(head) >= PEEK_BYTES else head
    if sample:
        # Check if content is printable text (more robust check)
        try:
            # Try to decode as UTF-8 and check for printable characters
            decoded = sample.decode('utf-8', errors='ignore')
            if decoded and all(c.isprintable() or c in '\t\n\r' for c in decoded[:500]):
                return "text"
        except:
            pass

        # Fallback: byte-level printable check
        if all((32 <= b <= 126) or b in (9, 10, 13) for b in sample):
            return "text"

    # Extension-based fallback for known text formats
    text_extensions = (".txt", ".md", ".csv", ".json", ".xml", ".yaml", ".yml",
                      ".js", ".py", ".java", ".cpp", ".c", ".h", ".cs", ".php", ".rb")
    for ext in text_extensions:
        if fn.endswith(ext):
            return "text"

    return "unknown"

@app.post("/extract-pdf-text")
async def extract_text_endpoint(
    file: UploadFile = File(...),
    content_length: Optional[int] = Header(None, alias="content-length")
):
    """
    Optimized endpoint to upload a file and extract its text based on the file type.
    Non-blocking: File parsing & heavy sanitization run in threadpool.
    Fixed: content-type logic, magic bytes detection, double-read elimination.
    """

    # 0) Early rejections
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded.")

    # Early reject via Content-Length if provided
    if content_length:
        try:
            if int(content_length) > MAX_FILE_SIZE:
                raise HTTPException(status_code=413, detail=f"File too large (max {MAX_FILE_SIZE // (1024 * 1024)} MB).")
        except (ValueError, TypeError):
            pass

    # Sanitize filename for security
    sanitized_filename = sanitize_filename(file.filename)
    if not sanitized_filename:
        raise HTTPException(status_code=400, detail="Invalid filename.")

    # 1) Magic bytes detection with larger peek window
    await file.seek(0)
    head = await file.read(PEEK_BYTES)
    await file.seek(0)

    detected_type = _detect_file_type(head, file.filename)
    if detected_type == "unknown":
        raise HTTPException(status_code=415, detail="Unsupported or invalid file format.")

    # 2) Stream to compute size + hash in one pass (memory efficient)
    size = 0
    hasher = hashlib.sha256()
    chunk_size = 65536  # 64KB chunks

    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        size += len(chunk)
        if size > MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail=f"File too large (max {MAX_FILE_SIZE // (1024 * 1024)} MB).")
        hasher.update(chunk)

    file_hash = hasher.hexdigest()
    await file.seek(0)  # Reset for extraction

    # 3) Non-blocking extraction with module-level semaphore
    def _extract_sync():
        """Extract text without unnecessary double-read"""
        # Use the underlying SpooledTemporaryFile directly (no extra read)
        file.file.seek(0)
        return read_files.extract_text(file)

    try:
        async with EXTRACT_SEM:  # Use module-level semaphore
            text = await asyncio.wait_for(
                asyncio.to_thread(_extract_sync),
                timeout=30.0
            )
    except asyncio.TimeoutError:
        raise HTTPException(status_code=408, detail="File processing timeout.")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Text extraction failed: {str(e)}")
        raise HTTPException(status_code=400, detail="Failed to extract text from file.")

    if not text or not text.strip():
        upload_timestamp = int(time.time() * 1000)  # Unix timestamp in milliseconds
        return JSONResponse(content={
            "text": "",
            "file_hash": file_hash,
            "size_bytes": size,
            "filename": sanitized_filename,
            "uploaded_at": upload_timestamp,
            "extracted_text_length": 0  # No text extracted
        })

    # 4) Sanitize with truncation first to cap processing work
    MAX_TEXT = 1_000_000
    to_sanitize = text[:MAX_TEXT] if len(text) > MAX_TEXT else text

    def _sync_sanitize():
        return sanitize_with_xss_detection(
            to_sanitize,
            allow_html=False,
            max_length=MAX_TEXT,
            context="file_extraction"
        )

    try:
        sanitized_text = await asyncio.wait_for(
            asyncio.to_thread(_sync_sanitize),
            timeout=10.0
        )
    except asyncio.TimeoutError:
        raise HTTPException(status_code=408, detail="Text sanitization timeout.")
    except Exception as e:
        logger.error(f"Sanitization error: {str(e)}")
        raise HTTPException(status_code=400, detail="Text sanitization failed.")

    if not sanitized_text and text:
        raise HTTPException(status_code=400, detail="File contains potentially malicious content.")

    # 5) Fire-and-forget analytics (skip for now to avoid errors)
    # The main file processing doesn't depend on analytics
    async def _track():
        try:
            # Only track if we have valid app/user context (not demo data)
            # For now, just log the successful extraction
            logger.info(f"üìä File processed successfully: {sanitized_filename} ({format_bytes(size)})")
        except Exception as e:
            logger.warning(f"Analytics failed: {str(e)}")

    asyncio.create_task(_track())

    # 6) Return optimized response with timestamp
    upload_timestamp = int(time.time() * 1000)  # Unix timestamp in milliseconds
    extracted_text_length = len(sanitized_text) if sanitized_text else 0

    return JSONResponse(content={
        "text": sanitized_text,
        "file_hash": file_hash,
        "size_bytes": size,
        "filename": sanitized_filename,
        "uploaded_at": upload_timestamp,
        "extracted_text_length": extracted_text_length,  # Length of extracted text for accurate KU calculation
        "processing": "async_nonblocking_fixed"
    })


# Load environment variables
load_dotenv()

openai.api_key = os.getenv("OPENAI_API_KEY")
neo4j_uri = os.getenv("NEO4J_URI")
neo4j_user = os.getenv("NEO4J_USER")
neo4j_password = os.getenv("NEO4J_PASSWORD")

# Automatically detect environment and set Convex URL
def get_convex_url():
    """
    Automatically detect if running locally or in production and return appropriate Convex URL
    """
    # Check if we're running locally
    is_local = (
        os.getenv("ENVIRONMENT") == "development" or
        os.getenv("NODE_ENV") == "development" or
        not os.getenv("RAILWAY_ENVIRONMENT") and  # Not on Railway
        not os.getenv("RENDER") and              # Not on Render
        not os.getenv("VERCEL") and              # Not on Vercel
        not os.getenv("HEROKU_APP_NAME")         # Not on Heroku
    )

    # Use environment variable for Convex URL
    convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

    if is_local:
        logger.info(f"üîß Using Convex deployment (local environment detected): {convex_url}")
    else:
        logger.info(f"üöÄ Using Convex deployment (production environment detected): {convex_url}")

    return f"{convex_url}/api/run/chats/getChatByIdExposed"

# Get the appropriate Convex URL
CONVEX_URL = get_convex_url()

# ==============================================================================
# Credit Management System
# ==============================================================================

# Model multipliers for credit calculation (GPT-4o-mini as 1x baseline)
MODEL_MULTIPLIERS = {
    # OpenAI Models
    'gpt-4o-mini': 1,         # Baseline: 1 credit = 1000 tokens
    'gpt-3.5-turbo': 0.7,     # 0.7x of 4o-mini
    'gpt-4': 18,              # 18x more expensive than 4o-mini
    'gpt-4-turbo': 12,        # 12x more expensive than 4o-mini
    'gpt-4o': 15,             # 15x more expensive than 4o-mini

    # Anthropic Claude Models
    'claude-3-haiku': 1,      # Similar to 4o-mini
    'claude-3-sonnet': 8,     # 8x more expensive than 4o-mini
    'claude-3-opus': 20,      # 20x more expensive than 4o-mini
    'claude-3.5-sonnet': 10,  # 10x more expensive than 4o-mini

    # Google Gemini Models
    'gemini-pro': 3,          # 3x more expensive than 4o-mini
    'gemini-ultra': 12,       # 12x more expensive than 4o-mini
    'gemini-1.5-pro': 4,      # 4x more expensive than 4o-mini

    # Open Source Models (cheaper)
    'llama-3': 0.5,           # 0.5x cheaper than 4o-mini
    'llama-3.1': 0.6,         # 0.6x of 4o-mini
    'mistral-7b': 0.5,        # 0.5x cheaper than 4o-mini
    'mixtral-8x7b': 0.8,      # 0.8x of 4o-mini
}

def calculate_credits_used(tokens: int, model: str) -> float:
    """
    Calculate credits used based on tokens and model.
    Uses the EXACT same calculation as frontend: (tokens / 1000) * multiplier
    Returns fractional credits, NOT rounded.
    """
    multiplier = MODEL_MULTIPLIERS.get(model, 1)
    return (tokens / 1000) * multiplier

class InsufficientCreditsError(Exception):
    """Raised when user doesn't have enough credits"""
    def __init__(self, required: float, available: float):
        self.required = required
        self.available = available
        super().__init__(f"Insufficient credits: need {required}, have {available}")

async def check_user_credits(user_id: str, required_credits: float) -> dict:
    """Check if user has enough credits and return credit info"""
    try:
        logger.info(f"üí≥ Checking credits for user {user_id}: need {required_credits}")

        # Use backend-safe credit checking function
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"{convex_url}/api/run/backend_credits/checkDeveloperCredits",
                json={
                    "args": {"developerId": user_id},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            if response.status_code == 200:
                result = response.json()
                credit_data = result.get("value")

                # Unified credit system: returns 500 default if not found (same as frontend)
                if credit_data:
                    remaining = credit_data.get("remainingCredits", 500)
                    total = credit_data.get("totalCredits", 500)
                    used = credit_data.get("usedCredits", 0)
                    found = credit_data.get("found", False)

                    has_sufficient = remaining >= required_credits

                    if not found:
                        logger.info(f"üí≥ No credit record found for user {user_id}, using default 500 credits")
                    else:
                        logger.info(f"üí≥ Credit check: {remaining}/{total} remaining, need {required_credits}, sufficient: {has_sufficient}")

                    return {
                        "has_sufficient": has_sufficient,
                        "remaining": remaining,
                        "total": total,
                        "used": used,
                        "required": required_credits
                    }
                else:
                    logger.warning(f"üí≥ No credit data returned for user {user_id}, using default 500")
                    return {"has_sufficient": 500 >= required_credits, "remaining": 500, "total": 500, "used": 0, "required": required_credits}
            else:
                logger.error(f"üí≥ Credit check failed: {response.status_code}")
                return {"has_sufficient": False, "remaining": 0, "total": 0, "used": 0, "required": required_credits}

    except Exception as e:
        logger.error(f"Error checking user credits: {e}")
        return {
            "has_sufficient": False,
            "remaining": 0,
            "total": 0,
            "used": 0
        }

async def consume_user_credits(
    user_id: str,
    credits: float,
    model: str,
    tokens_used: int,
    chat_id: str = None,
    description: str = "AI model usage"
) -> bool:
    """Consume credits from user's balance"""
    try:
        logger.info(f"üí≥ Attempting to consume {credits} credits for user {user_id}: {description}")
        logger.info(f"   Model: {model}, Tokens: {tokens_used}, Chat: {chat_id}")

        # Use the backend-safe updateUserCredits function to consume credits
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

        async with httpx.AsyncClient() as client:
            # First get current credits by querying user_credits table directly
            # We'll use a more direct approach since we can't use the auth-based functions

            # Use the backend-safe credit consumption function
            response = await client.post(
                f"{convex_url}/api/run/backend_credits/consumeDeveloperCredits",
                json={
                    "args": {
                        "developerId": user_id,
                        "credits": credits,
                        "model": model,
                        "tokensUsed": tokens_used,
                        "description": description,
                        "chatId": chat_id
                    },
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            if response.status_code == 200:
                result = response.json()
                value = result.get("value", {})

                if value.get("success"):
                    credits_consumed = value.get("creditsUsed", credits)
                    remaining = value.get("remainingCredits", 0)
                    was_initialized = value.get("wasInitialized", False)

                    if was_initialized:
                        logger.info(f"‚úÖ Initialized user {user_id} with credits and consumed {credits_consumed}")
                    else:
                        logger.info(f"‚úÖ Successfully deducted {credits_consumed} credits from user {user_id}")

                    logger.info(f"üí≥ User {user_id} now has {remaining} credits remaining")
                    return True
                else:
                    error_msg = value.get("error", "Unknown error")
                    required = value.get("required", credits)
                    available = value.get("remainingCredits", 0)
                    logger.error(f"‚ùå Credit consumption failed: {error_msg}")
                    logger.error(f"üí≥ Required: {required}, Available: {available}")
                    return False
            else:
                logger.error(f"‚ùå Failed to consume credits: {response.status_code} - {response.text}")
                return False

    except Exception as e:
        logger.error(f"Error consuming user credits: {e}")
        return False

async def get_developer_id_from_chat(chat_id: str) -> str:
    """
    Get the developer ID responsible for a chat's API costs.
    Returns the developer ID or the original chat owner ID as fallback.
    """
    try:
        logger.info(f"üîç Looking up developer ID for chat {chat_id}")
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

        async with httpx.AsyncClient() as client:
            # First get the chat data
            chat_response = await client.post(
                f"{convex_url}/api/run/backend_credits/getChatWithApp",
                json={
                    "args": {"chatId": chat_id},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            if chat_response.status_code == 200:
                chat_data = chat_response.json()
                value = chat_data.get("value")

                if value:
                    # If chat has a parentAppId, get the developer from that app
                    if value.get("parentAppId"):
                        app_id = value["parentAppId"]
                        logger.info(f"üîç Chat {chat_id} is linked to app {app_id}")

                        # Get app info to find developer
                        app_response = await client.post(
                            f"{convex_url}/api/run/app_management/getAppWithSettings",
                            json={
                                "args": {"appId": app_id},
                                "format": "json"
                            },
                            headers={"Content-Type": "application/json"}
                        )

                        if app_response.status_code == 200:
                            app_data = app_response.json()
                            app_value = app_data.get("value")
                            if app_value and app_value.get("developerId"):
                                developer_id = app_value["developerId"]
                                logger.info(f"‚úÖ Found developer ID {developer_id} for chat {chat_id} via app {app_id}")
                                return developer_id

                    # If no app association, check if this chat has apps created from it
                    chat_owner_id = value.get("userId")
                    if chat_owner_id:
                        # Check if there are any apps with this chat as parentChatId
                        apps_response = await client.post(
                            f"{convex_url}/api/run/backend_credits/getAppsForParentChat",
                            json={
                                "args": {"parentChatId": chat_id},
                                "format": "json"
                            },
                            headers={"Content-Type": "application/json"}
                        )

                        if apps_response.status_code == 200:
                            apps_data = apps_response.json()
                            apps_value = apps_data.get("value", [])
                            if apps_value and len(apps_value) > 0:
                                # Use the developer ID from the first app (they should all be the same developer)
                                developer_id = apps_value[0].get("developerId")
                                if developer_id:
                                    logger.info(f"‚úÖ Found developer ID {developer_id} for chat {chat_id} as chat owner with apps")
                                    return developer_id

                        # Fallback: use the chat owner as the one responsible for credits
                        logger.info(f"üí° Using chat owner {chat_owner_id} as developer for chat {chat_id}")
                        return chat_owner_id

    except Exception as e:
        logger.error(f"‚ùå Error looking up developer for chat {chat_id}: {e}")

    # Ultimate fallback: use chat_id as user_id
    logger.warning(f"‚ö†Ô∏è Using chat_id {chat_id} as fallback developer ID")
    return chat_id

async def validate_and_consume_credits(
    user_id: str,
    model: str,
    estimated_tokens: int,
    chat_id: str = None
) -> int:
    """
    Validate user has enough credits and consume them.
    Returns the number of credits consumed.
    """
    required_credits = calculate_credits_used(estimated_tokens, model)

    # Determine who should be charged for the credits
    actual_user_id = user_id
    if chat_id:
        # Try to get the developer ID who should be charged
        developer_id = await get_developer_id_from_chat(chat_id)
        if developer_id != chat_id:  # Only use if we found a real developer ID
            actual_user_id = developer_id
            logger.info(f"üí≥ Charging developer {developer_id} instead of original user {user_id}")

    # Try to consume the credits (this will auto-initialize if user doesn't exist)
    success = await consume_user_credits(
        actual_user_id,
        required_credits,
        model,
        estimated_tokens,
        chat_id,
        f"AI response using {model} (chat: {chat_id})"
    )

    if not success:
        # If consumption failed, check why and provide helpful error
        credit_info = await check_user_credits(actual_user_id, required_credits)
        if not credit_info["has_sufficient"]:
            raise InsufficientCreditsError(required_credits, credit_info["remaining"])
        else:
            raise Exception("Failed to consume credits for unknown reason")

    return required_credits

async def consume_credits_for_actual_usage(
    user_id: str,
    model: str,
    question: str,
    response: str,
    chat_id: str = None,
    skip_developer_lookup: bool = False
) -> float:
    """
    Consume credits based on actual token usage from AI response.
    Uses the EXACT same calculation as the frontend Trainly chat.
    Returns the number of credits consumed.

    Args:
        user_id: User ID to charge credits from
        model: Model used for the response
        question: User's question
        response: AI's response
        chat_id: Optional chat ID for logging
        skip_developer_lookup: If True, use user_id directly without developer lookup (for API calls)
    """
    # Calculate actual tokens used (EXACT same method as frontend)
    # Frontend: Math.ceil(question.length / 4) + Math.ceil(response.length / 4)
    import math
    question_tokens = math.ceil(len(question) / 4)
    response_tokens = math.ceil(len(response) / 4)
    total_actual_tokens = question_tokens + response_tokens

    # Calculate credits based on actual usage (EXACT same method as frontend)
    # Frontend: (tokens / 1000) * multiplier (NO Math.ceil!)
    multiplier = MODEL_MULTIPLIERS.get(model, 1)
    actual_credits = (total_actual_tokens / 1000) * multiplier

    # Determine who should be charged for the credits
    actual_user_id = user_id
    if chat_id and not skip_developer_lookup:
        # Try to get the developer ID who should be charged (for regular chat usage)
        developer_id = await get_developer_id_from_chat(chat_id)
        if developer_id != chat_id:  # Only use if we found a real developer ID
            actual_user_id = developer_id
            logger.info(f"üí≥ Charging developer {developer_id} instead of original user {user_id}")

    # Consume the credits based on actual usage
    success = await consume_user_credits(
        actual_user_id,
        actual_credits,
        model,
        total_actual_tokens,
        chat_id,
        f"AI response using {model} - {total_actual_tokens} tokens (chat: {chat_id})"
    )

    if not success:
        # If consumption failed, check why and provide helpful error
        credit_info = await check_user_credits(actual_user_id, actual_credits)
        if not credit_info["has_sufficient"]:
            raise InsufficientCreditsError(actual_credits, credit_info["remaining"])
        else:
            raise Exception("Failed to consume credits for unknown reason")

    logger.info(f"üí≥ Consumed {actual_credits} credits for {total_actual_tokens} actual tokens ({model})")
    return actual_credits

async def ensure_developer_has_credits(developer_user_id: str, min_credits: int = 1000):
    """Ensure developer has credits, initialize with default amount if needed"""
    try:
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

        async with httpx.AsyncClient() as client:
            # Check current credits
            response = await client.post(
                f"{convex_url}/api/run/subscriptions/getUserCredits",
                json={
                    "args": {"userId": developer_user_id},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            current_credits = 0
            if response.status_code == 200:
                result = response.json()
                if result.get("value"):
                    current_credits = result["value"].get("remainingCredits", 0)

            # If no credits, initialize with default amount
            if current_credits == 0:
                logger.info(f"üí≥ Initializing credits for developer {developer_user_id} with {min_credits} credits")
                init_response = await client.post(
                    f"{convex_url}/api/run/subscriptions/initializeUserCredits",
                    json={
                        "args": {"userId": developer_user_id, "initialCredits": min_credits},
                        "format": "json"
                    },
                    headers={"Content-Type": "application/json"}
                )

                if init_response.status_code == 200:
                    logger.info(f"‚úÖ Developer {developer_user_id} initialized with {min_credits} credits")
                else:
                    logger.warning(f"‚ö†Ô∏è Could not initialize credits for developer: {init_response.status_code}")
            else:
                logger.info(f"üí≥ Developer {developer_user_id} has {current_credits} credits")

    except Exception as e:
        logger.error(f"Error ensuring developer credits: {e}")

def sanitize_for_neo4j(text: str) -> str:
    safe_text = text
    safe_text = safe_text.replace('\\', '\\\\')
    safe_text = safe_text.replace('"', '\\"')
    safe_text = safe_text.replace("'", "\\'")
    safe_text = safe_text.replace('\n', '\\n')
    safe_text = safe_text.replace('\r', '\\r')
    return safe_text

def chunk_text(full_text: str, max_chars: int = 2000) -> List[str]:
    """
    Chunk text into pieces for embedding.

    PERFORMANCE OPTIMIZATION: Increased from 500 to 2000 chars.
    This reduces thousands of embedding calls to hundreds for large files.

    Args:
        full_text: The full text to chunk
        max_chars: Maximum characters per chunk (default 2000, was 500)

    Returns:
        List of text chunks
    """
    chunks = []
    start = 0
    while start < len(full_text):
        end = start + max_chars
        if end > len(full_text):
            chunks.append(full_text[start:])
            break

        # Try to break at sentence boundary for better semantic coherence
        # Look back up to 200 chars for a period, question mark, or exclamation
        if end < len(full_text):
            best_break = end
            for i in range(end, max(start + max_chars - 200, start), -1):
                if full_text[i-1] in '.!?\n':
                    best_break = i
                    break
            end = best_break

        chunks.append(full_text[start:end])
        start = end
    return chunks

def get_embedding(text: str) -> List[float]:
    """Get embedding for a single text. Use get_embeddings_batch for multiple texts."""
    response = openai.embeddings.create(
        model="text-embedding-3-small",
        input=text
    )
    return response.data[0].embedding

def get_embeddings_batch(texts: List[str], batch_size: int = 32) -> List[List[float]]:
    """
    Get embeddings for multiple texts in batches.

    PERFORMANCE OPTIMIZATION: Batches embedding calls instead of one-by-one.
    Reduces API overhead dramatically - e.g., 1000 chunks becomes ~32 API calls
    instead of 1000.

    Args:
        texts: List of texts to embed
        batch_size: Number of texts per API call (default 32, max ~100 for OpenAI)

    Returns:
        List of embeddings in the same order as input texts
    """
    all_embeddings = []

    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        try:
            response = openai.embeddings.create(
                model="text-embedding-3-small",
                input=batch
            )
            # Embeddings come back in order
            batch_embeddings = [item.embedding for item in response.data]
            all_embeddings.extend(batch_embeddings)

            # Log progress for large batches
            if len(texts) > batch_size:
                logger.info(f"üìä Embedded batch {i//batch_size + 1}/{(len(texts) + batch_size - 1)//batch_size}")
        except Exception as e:
            logger.error(f"Batch embedding failed at index {i}: {e}")
            # Fallback to individual calls for this batch
            for text in batch:
                try:
                    embedding = get_embedding(text)
                    all_embeddings.append(embedding)
                except Exception as inner_e:
                    logger.error(f"Individual embedding failed: {inner_e}")
                    # Return zeros as fallback (better than crashing)
                    all_embeddings.append([0.0] * 1536)

    return all_embeddings

def cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))

# ==============================================================================
# Queue-Based File Processing System
# ==============================================================================

async def update_convex_file_status(
    file_queue_id: str,
    status: str,
    progress: int = 0,
    error: Optional[str] = None,
    extracted_text_length: Optional[int] = None,
    knowledge_units: Optional[float] = None,
    nodes_created: Optional[int] = None
):
    """
    Update file processing status in Convex database.
    This provides persistent status tracking across workers and restarts.
    """
    try:
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

        # Build the update payload
        update_args = {
            "fileQueueId": file_queue_id,
            "status": status,
            "progress": progress,
        }

        if error:
            update_args["error"] = error
        if extracted_text_length is not None:
            update_args["extractedTextLength"] = extracted_text_length
        if knowledge_units is not None:
            update_args["knowledgeUnits"] = knowledge_units
        if nodes_created is not None:
            update_args["nodesCreated"] = nodes_created

        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"{convex_url}/api/mutation",
                json={
                    "path": "fileQueue:updateFileProgressByQueueId",
                    "args": update_args,
                    "format": "json"
                },
                headers={"Content-Type": "application/json"},
                timeout=10.0
            )

            if response.status_code == 200:
                logger.info(f"üìä Updated Convex file status: {file_queue_id} -> {status} ({progress}%)")
            else:
                logger.warning(f"‚ö†Ô∏è Failed to update Convex status: {response.status_code}")

    except Exception as e:
        logger.error(f"‚ùå Error updating Convex file status: {e}")

def update_convex_file_status_sync(
    file_queue_id: str,
    status: str,
    progress: int = 0,
    error: Optional[str] = None,
    extracted_text_length: Optional[int] = None,
    knowledge_units: Optional[float] = None,
    nodes_created: Optional[int] = None
):
    """
    Synchronous version of update_convex_file_status for use in RQ workers.
    """
    import requests

    try:
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

        update_args = {
            "fileQueueId": file_queue_id,
            "status": status,
            "progress": progress,
        }

        if error:
            update_args["error"] = error
        if extracted_text_length is not None:
            update_args["extractedTextLength"] = extracted_text_length
        if knowledge_units is not None:
            update_args["knowledgeUnits"] = knowledge_units
        if nodes_created is not None:
            update_args["nodesCreated"] = nodes_created

        response = requests.post(
            f"{convex_url}/api/mutation",
            json={
                "path": "fileQueue:updateFileProgressByQueueId",
                "args": update_args,
                "format": "json"
            },
            headers={"Content-Type": "application/json"},
            timeout=10.0
        )

        if response.status_code == 200:
            print(f"üìä Updated Convex file status: {file_queue_id} -> {status} ({progress}%)")
        else:
            print(f"‚ö†Ô∏è Failed to update Convex status: {response.status_code}")

    except Exception as e:
        print(f"‚ùå Error updating Convex file status: {e}")

def process_file_job(
    file_queue_id: str,
    chat_id: str,
    filename: str,
    text_content: str,
    file_size: int,
    scope_values: Optional[Dict[str, Union[str, int, bool]]] = None,
    file_id: Optional[str] = None
):
    """
    Pure processor function for file ingestion - called by RQ worker.

    This function:
    1. Updates status to PROCESSING
    2. Chunks the text
    3. Gets embeddings in batches
    4. Creates nodes in Neo4j
    5. Updates status to READY or FAILED

    Args:
        file_queue_id: The Convex file_upload_queue document ID for status updates
        chat_id: The chat this file belongs to
        filename: The filename
        text_content: The extracted text content
        file_size: The file size in bytes
        scope_values: Optional custom scope values for the file
        file_id: Optional pre-generated file ID (if not provided, will be generated)
    """
    import time as time_module
    start_time = time_module.time()

    # Load environment variables for worker process
    load_dotenv()

    # Initialize OpenAI and Neo4j from environment
    openai.api_key = os.getenv("OPENAI_API_KEY")
    neo4j_uri_local = os.getenv("NEO4J_URI")
    neo4j_user_local = os.getenv("NEO4J_USER")
    neo4j_password_local = os.getenv("NEO4J_PASSWORD")

    print(f"üîÑ Processing file job: {filename} for chat {chat_id}")
    print(f"   Text length: {len(text_content)} chars, File size: {file_size} bytes")

    # Generate file ID if not provided
    if not file_id:
        file_id = f"{chat_id}_{filename}_{int(time_module.time())}"

    try:
        # Step 1: Update status to PROCESSING
        update_convex_file_status_sync(file_queue_id, "processing", 10)

        # Step 2: Chunk the text (with larger chunk size for performance)
        print(f"üìÑ Chunking text...")
        chunks = chunk_text(text_content, max_chars=2000)  # Increased from 500!
        num_chunks = len(chunks)
        print(f"   Created {num_chunks} chunks (2000 chars each)")

        update_convex_file_status_sync(file_queue_id, "processing", 20)

        # Step 3: Get embeddings in batches (major performance improvement!)
        print(f"üß† Getting embeddings in batches...")
        embeddings = get_embeddings_batch(chunks, batch_size=32)
        print(f"   Got {len(embeddings)} embeddings")

        update_convex_file_status_sync(file_queue_id, "processing", 50)

        # Step 4: Create nodes in Neo4j
        print(f"üìä Creating nodes in Neo4j...")

        # Build scope properties if provided
        scope_props_set = ""
        scope_props_create = ""
        scope_params = {}
        if scope_values:
            for key, value in scope_values.items():
                safe_key = key.replace(" ", "_").replace("-", "_")[:50]
                param_name = f"scope_{safe_key}"
                scope_props_set += f", d.{safe_key} = ${param_name}"
                scope_props_create += f", {safe_key}: ${param_name}"
                scope_params[param_name] = value

        with GraphDatabase.driver(neo4j_uri_local, auth=(neo4j_user_local, neo4j_password_local)) as driver:
            with driver.session() as session:
                # Create document node using explicit write transaction
                current_timestamp = int(time_module.time() * 1000)
                text_size_bytes = len(text_content.encode('utf-8'))

                def create_document_tx(tx):
                    doc_query = f"""
                    MERGE (d:Document {{id: $pdf_id}})
                    SET d.chatId = $chat_id,
                        d.filename = $filename,
                        d.uploadDate = $upload_date,
                        d.sizeBytes = $size_bytes{scope_props_set}
                    RETURN d
                    """
                    result = tx.run(doc_query,
                        pdf_id=file_id,
                        chat_id=chat_id,
                        filename=filename,
                        upload_date=current_timestamp,
                        size_bytes=text_size_bytes,
                        **scope_params
                    )
                    return result.single()

                doc_result = session.execute_write(create_document_tx)
                print(f"   Created document node: {file_id} (confirmed: {doc_result is not None})")

                update_convex_file_status_sync(file_queue_id, "processing", 60)

                # Create chunk nodes with embeddings using explicit write transactions
                def create_chunks_tx(tx, chunks_data):
                    """Create all chunks in a single transaction"""
                    created_count = 0
                    for chunk_data in chunks_data:
                        chunk_query = f"""
                        MATCH (d:Document {{id: $pdf_id}})
                        CREATE (c:Chunk {{
                            id: $chunk_id,
                            text: $text,
                            embedding: $embedding,
                            chatId: $chat_id{scope_props_create}
                        }})
                        CREATE (d)-[:HAS_CHUNK {{order: $order}}]->(c)
                        RETURN c
                        """
                        tx.run(chunk_query, **chunk_data)
                        created_count += 1
                    return created_count

                # Prepare chunk data
                chunks_data = []
                for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    chunk_id = f"{file_id}-{i}"
                    chunks_data.append({
                        "pdf_id": file_id,
                        "chunk_id": chunk_id,
                        "text": chunk,
                        "embedding": embedding,
                        "chat_id": chat_id,
                        "order": i,
                        **scope_params
                    })

                # Execute in write transaction
                created_count = session.execute_write(create_chunks_tx, chunks_data)
                print(f"   Created {created_count} chunk nodes (confirmed)")

                # Step 5: Create sequential relationships using explicit transaction
                if num_chunks > 1:
                    print(f"üîó Creating sequential relationships...")

                    def create_relationships_tx(tx, file_id, num_chunks):
                        for i in range(num_chunks - 1):
                            link_query = """
                            MATCH (c1:Chunk {id: $chunk1_id})
                            MATCH (c2:Chunk {id: $chunk2_id})
                            CREATE (c1)-[:NEXT {order: $order}]->(c2)
                            """
                            tx.run(link_query,
                                chunk1_id=f"{file_id}-{i}",
                                chunk2_id=f"{file_id}-{i+1}",
                                order=i
                            )
                        return num_chunks - 1

                    rels_created = session.execute_write(create_relationships_tx, file_id, num_chunks)
                    print(f"   Created {rels_created} NEXT relationships")

        # Step 6: Calculate knowledge units (tokens from extracted text)
        # Rough estimate: 1 token ‚âà 4 characters for English text
        estimated_tokens = len(text_content) // 4
        knowledge_units = estimated_tokens / 1000  # 1 KU = 1000 tokens

        # Step 7: Update status to COMPLETED
        elapsed = time_module.time() - start_time
        print(f"‚úÖ File processing completed in {elapsed:.2f}s")
        print(f"   Chunks: {num_chunks}, KU: {knowledge_units:.2f}")

        update_convex_file_status_sync(
            file_queue_id, "completed", 100,
            extracted_text_length=len(text_content),
            knowledge_units=knowledge_units,
            nodes_created=num_chunks
        )

        return {
            "status": "success",
            "file_id": file_id,
            "chunks_created": num_chunks,
            "knowledge_units": knowledge_units,
            "elapsed_seconds": elapsed
        }

    except Exception as e:
        print(f"‚ùå File processing failed: {e}")
        import traceback
        traceback.print_exc()

        update_convex_file_status_sync(
            file_queue_id, "failed", 0,
            error=str(e)[:500]  # Truncate error message
        )

        return {"status": "failed", "error": str(e)}

def enqueue_file_processing(
    file_queue_id: str,
    chat_id: str,
    filename: str,
    text_content: str,
    file_size: int,
    scope_values: Optional[Dict[str, Union[str, int, bool]]] = None,
    file_id: Optional[str] = None
) -> Optional[str]:
    """
    Enqueue a file for background processing.

    Returns the job ID if queued successfully, None if Redis unavailable
    (in which case caller should fall back to sync processing).
    """
    queue = get_file_queue()

    if queue is None:
        logger.warning("‚ö†Ô∏è Redis queue not available, cannot enqueue job")
        return None

    try:
        job = queue.enqueue(
            process_file_job,
            file_queue_id,
            chat_id,
            filename,
            text_content,
            file_size,
            scope_values,
            file_id,
            job_timeout=1800,  # 30 minute timeout
            result_ttl=86400,  # Keep result for 24 hours
        )
        logger.info(f"üì§ Enqueued file processing job: {job.id} for {filename}")
        return job.id
    except Exception as e:
        logger.error(f"‚ùå Failed to enqueue job: {e}")
        return None

# ==============================================================================
# API Authentication Functions
# ==============================================================================

def check_rate_limit(api_key: str, ip: str) -> bool:
    """Simple rate limiting check"""
    now = time.time()
    key = f"{api_key}:{ip}"

    # Clean old entries
    request_counts[key] = [
        timestamp for timestamp in request_counts.get(key, [])
        if now - timestamp < RATE_LIMIT_WINDOW
    ]

    # Check limit
    if len(request_counts.get(key, [])) >= RATE_LIMIT_MAX:
        return False

    # Add current request
    if key not in request_counts:
        request_counts[key] = []
    request_counts[key].append(now)

    return True

async def verify_api_key_and_chat(api_key: str, chat_id: str) -> bool:
    """
    Verify API key has access to the specific chat
    Integrates with your existing Convex system
    """
    try:
        # Call your Convex endpoint to verify (auto-detects dev/prod)
        # chat_id is the internal Convex _id
        async with httpx.AsyncClient() as client:
            response = await client.post(
                CONVEX_URL,
                json={
                    "args": {"id": chat_id},
                    "format": "json"
                }
            )

            if response.status_code != 200:
                logger.warning(f"Convex call failed for chat {chat_id}: {response.status_code} - {response.text}")
                return False

            chat_data = response.json()

            # Handle both error responses and null responses
            if chat_data.get("status") == "error":
                error_message = chat_data.get("errorMessage", "Unknown error")
                logger.warning(f"Chat {chat_id} - Convex error: {error_message}")
                return False

            chat = chat_data.get("value")

            # Debug logging
            logger.info(f"üîç Debug chat data for {chat_id}:")
            logger.info(f"   Response status: {response.status_code}")
            logger.info(f"   Chat found: {chat is not None}")

            if chat:
                logger.info(f"   Chat title: {chat.get('title', 'No title')}")
                logger.info(f"   API key in DB: {chat.get('apiKey', 'No key')[:10]}...")
                logger.info(f"   API key provided: {api_key[:10]}...")
                logger.info(f"   API disabled: {chat.get('apiKeyDisabled', 'Not set')}")
                logger.info(f"   Has API access: {chat.get('hasApiAccess', 'Not set')}")
                logger.info(f"   Is archived: {chat.get('isArchived', 'Not set')}")
                logger.info(f"   Visibility: {chat.get('visibility', 'Not set')}")

            if not chat:
                logger.warning(f"Chat {chat_id} not found - chat returned null from Convex")
                return False

            # Check if API key matches and is enabled
            chat_api_key = chat.get("apiKey")
            api_disabled = chat.get("apiKeyDisabled", True)
            is_archived = chat.get("isArchived", False)
            has_api_access = chat.get("hasApiAccess", False)

            # More lenient check - if hasApiAccess field doesn't exist, check if apiKeyDisabled is False
            api_access_enabled = has_api_access or (not api_disabled and chat_api_key and chat_api_key != "undefined")

            # Verify conditions
            if not api_access_enabled:
                logger.warning(f"API access not enabled for chat {chat_id} - hasApiAccess: {has_api_access}, apiKeyDisabled: {api_disabled}")
                return False

            if is_archived:
                logger.warning(f"Chat {chat_id} is archived")
                return False

            if not chat_api_key or chat_api_key == "undefined":
                logger.warning(f"No API key set for chat {chat_id}")
                return False

            # Verify API key matches
            if api_key != chat_api_key:
                logger.warning(f"API key mismatch for chat {chat_id} - expected: {chat_api_key[:10]}..., got: {api_key[:10]}...")
                return False

            logger.info(f"API key verified successfully for chat {chat_id}")
            return True

    except Exception as e:
        logger.error(f"API key verification failed: {e}")
        return False

async def get_verified_chat_access(
    credentials: HTTPAuthorizationCredentials = Depends(security),
) -> HTTPAuthorizationCredentials:
    """FastAPI dependency to verify API key"""

    if not credentials:
        raise HTTPException(
            status_code=401,
            detail="API key required. Include 'Authorization: Bearer your_api_key' header."
        )

    return credentials

async def get_optional_chat_access(
    authorization: Optional[str] = Header(None),
) -> Optional[str]:
    """Optional authentication dependency that doesn't fail if header is missing"""
    if authorization and authorization.startswith("Bearer "):
        return authorization[7:]  # Return the token without "Bearer " prefix
    return None

# ==============================================================================
# Existing Helper Functions
# ==============================================================================

def analyze_chunk_relationships(chunks: List[str]) -> List[dict]:
    """Use AI to analyze relationships between chunks and determine logical connections"""
    if len(chunks) < 2:
        return []

    # Prepare chunks for analysis
    chunk_summaries = []
    for i, chunk in enumerate(chunks):
        # Create a brief summary of each chunk
        summary_prompt = f"Summarize this text in 1-2 sentences, focusing on the main concept or topic:\n\n{chunk}"

        try:
            response = openai.chat.completions.create(
                model="gpt-4o-mini",  # Use cheaper model for summaries
                messages=[{"role": "user", "content": summary_prompt}],
                max_tokens=100,
                temperature=0
            )
            summary = response.choices[0].message.content.strip()
            chunk_summaries.append({"index": i, "text": chunk[:200], "summary": summary})
        except Exception as e:
            print(f"Error creating summary for chunk {i}: {e}")
            chunk_summaries.append({"index": i, "text": chunk[:200], "summary": chunk[:100]})

    # Analyze relationships between all chunk pairs
    relationships = []

    # Create a more focused analysis prompt
    analysis_prompt = f"""Analyze these text chunks and find logical relationships. Return only valid JSON.

Chunks:
{chr(10).join([f"{item['index']}: {item['summary']}" for item in chunk_summaries])}

Find relationships like:
- EXPLAINS: One chunk explains concepts in another
- SUPPORTS: One chunk provides evidence for another
- ELABORATES: One chunk adds details to another
- INTRODUCES: One chunk introduces topics discussed in another
- CONCLUDES: One chunk concludes ideas from another

Return JSON with this exact format:
{{"relationships": [{{"source": 0, "target": 2, "type": "EXPLAINS", "confidence": 0.8}}]}}"""

    try:
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert at analyzing text relationships. Always respond with valid JSON only."},
                {"role": "user", "content": analysis_prompt}
            ],
            max_tokens=2000,
            temperature=0.3
        )

        ai_response = response.choices[0].message.content.strip()
        print(f"AI Response: {ai_response[:200]}...")  # Debug print

        # Clean up the response - sometimes AI adds markdown formatting
        if ai_response.startswith("```json"):
            ai_response = ai_response.replace("```json", "").replace("```", "").strip()
        elif ai_response.startswith("```"):
            ai_response = ai_response.replace("```", "").strip()

        import json
        result = json.loads(ai_response)
        ai_relationships = result.get("relationships", [])

        print(f"AI identified {len(ai_relationships)} potential semantic relationships")

        # Filter and validate AI relationships
        valid_relationships = []
        for rel in ai_relationships:
            if (rel.get("confidence", 0) > 0.6 and
                "source" in rel and "target" in rel and "type" in rel and
                0 <= rel["source"] < len(chunks) and 0 <= rel["target"] < len(chunks)):
                valid_relationships.append(rel)

        print(f"Validated {len(valid_relationships)} semantic relationships")

        # Only add NEXT relationships if we have no semantic relationships
        if len(valid_relationships) == 0:
            print("No semantic relationships found, falling back to sequential NEXT relationships")
            for i in range(len(chunks) - 1):
                valid_relationships.append({
                    "source": i,
                    "target": i + 1,
                    "type": "NEXT",
                    "description": "Sequential order in document",
                    "confidence": 1.0
                })
        else:
            print(f"Using {len(valid_relationships)} AI-generated semantic relationships")

        return valid_relationships

    except json.JSONDecodeError as json_error:
        print(f"JSON parsing error: {json_error}")
        print(f"Raw AI response: {ai_response}")
        # Fallback to simple sequential relationships
        return [{"source": i, "target": i + 1, "type": "NEXT", "description": "Sequential order", "confidence": 1.0}
                for i in range(len(chunks) - 1)]
    except Exception as e:
        print(f"Error analyzing relationships with AI: {e}")
        # Fallback to simple sequential relationships
        return [{"source": i, "target": i + 1, "type": "NEXT", "description": "Sequential order", "confidence": 1.0}
                for i in range(len(chunks) - 1)]

# ==============================================================================
# V1 Trusted Issuer API Endpoints (Main Feature)
# ==============================================================================

# V1 Console API: App registration
@app.post("/v1/console/apps/register")
async def v1_register_app(
    app_name: str = Form(...),
    issuer: str = Form(...),
    allowed_audiences: str = Form(...),  # JSON string of array
    alg_allowlist: str = Form(None),  # JSON string of array, optional
    jwks_uri: str = Form(None),  # Optional
    admin_token: str = Header(None, alias="x-admin-token")
):
    """
    Register an app for V1 Trusted Issuer authentication

    Developers register their OAuth provider details here.
    """

    # Verify admin access (in production, this would be properly secured)
    expected_token = os.getenv("TRAINLY_CONSOLE_ADMIN_TOKEN", "admin_dev_token_123")
    if not admin_token or admin_token != expected_token:
        raise HTTPException(status_code=401, detail="Invalid or missing admin token")

    try:
        # Parse JSON arrays
        import json
        audiences = json.loads(allowed_audiences)
        alg_list = json.loads(alg_allowlist) if alg_allowlist else ["RS256", "ES256"]

        # Generate app ID
        app_id = f"app_v1_{int(time.time())}_{hashlib.md5(app_name.encode()).hexdigest()[:8]}"

        # Create app configuration
        app_config = V1AppConfig(
            app_id=app_id,
            issuer=issuer,
            allowed_audiences=audiences,
            alg_allowlist=alg_list,
            jwks_uri=jwks_uri
        )

        # Store configuration
        V1_APP_CONFIGS[app_id] = app_config

        logger.info(f"Registered V1 app: {app_id} for issuer {issuer}")

        return {
            "success": True,
            "app_id": app_id,
            "app_name": app_name,
            "issuer": issuer,
            "allowed_audiences": audiences,
            "alg_allowlist": alg_list,
            "jwks_uri": jwks_uri,
            "message": "App registered for V1 Trusted Issuer authentication",
            "usage_instructions": {
                "client_flow": "User logs in with your OAuth ‚Üí Client sends ID token to /v1/me/* endpoints",
                "headers_required": {
                    "Authorization": "Bearer <USER_ID_TOKEN_FROM_YOUR_OAUTH>",
                    "X-App-ID": app_id
                }
            }
        }

    except Exception as e:
        logger.error(f"App registration failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Registration failed: {str(e)}")

# V1 User APIs: File upload and querying with OAuth ID tokens
@app.post("/v1/me/chats/query")
async def v1_user_query(
    messages: str = Form(...),  # JSON string of messages array
    response_tokens: int = Form(150),
    stream: bool = Form(False),
    scope_filters: str = Form("{}"),  # JSON string of scope filters (e.g., {"playlist_id": "xyz123"})
    authorization: str = Header(None, alias="authorization"),
    app_id: str = Header(None, alias="x-app-id"),
    request: Request = None
):
    """
    V1 User Query: User authenticates with their OAuth ID token

    Workflow:
    1. User logs into developer's app with OAuth (Clerk, Auth0, etc.)
    2. Developer's app gets ID token from OAuth provider
    3. Client sends ID token as Bearer auth to this endpoint
    4. Trainly validates token against registered app config
    5. Creates/uses permanent subchat for this user
    6. Returns AI response based on user's private data

    Optional scope_filters parameter allows filtering results by custom scope values
    (e.g., {"playlist_id": "xyz123", "workspace_id": "acme_corp"})
    """

    if not app_id:
        raise HTTPException(status_code=400, detail="X-App-ID header required")

    # Authenticate user with their OAuth ID token
    user_identity = await authenticate_v1_user(authorization, app_id)

    # Parse messages
    try:
        import json
        messages_array = json.loads(messages)
        if not messages_array or not isinstance(messages_array, list):
            raise ValueError("Invalid messages format")

        # Extract the latest user message
        user_message = None
        for msg in reversed(messages_array):
            if msg.get("role") == "user":
                user_message = msg.get("content")
                break

        if not user_message:
            raise HTTPException(status_code=400, detail="No user message found")

    except (json.JSONDecodeError, ValueError) as e:
        raise HTTPException(status_code=400, detail="Invalid messages JSON format")

    # Parse scope filters
    try:
        parsed_scope_filters = json.loads(scope_filters) if scope_filters else {}
        if parsed_scope_filters:
            logger.info(f"üîç V1 Query with scope filters: {parsed_scope_filters}")
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid scope_filters JSON")

    # Sanitize the query
    sanitized_question = sanitize_with_xss_detection(
        user_message,
        allow_html=False,
        max_length=5000,
        context="v1_user_query"
    )

    if not sanitized_question:
        raise HTTPException(status_code=400, detail="Invalid or potentially malicious question")

    try:
        # Create/get permanent subchat for this user
        # This ensures the same user always gets the same chat, making it truly permanent
        subchat = await get_or_create_user_subchat(
            user_identity["app_id"],
            user_identity["external_user_id"]
        )

        # Get app configuration to inherit parent chat settings
        app_config_from_convex = await get_app_config_from_convex(user_identity["app_id"])

        # Default settings
        selected_model = "gpt-4o-mini"
        temperature = 0.7
        max_tokens = response_tokens
        custom_prompt = None

        # Apply PUBLISHED parent chat settings if available
        if app_config_from_convex:
            try:
                # Get the app data to find the parent chat ID
                convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
                async with httpx.AsyncClient() as client:
                    app_response = await client.post(
                        f"{convex_url}/api/run/app_management/getAppWithSettings",
                        json={
                            "args": {"appId": user_identity["app_id"]},
                            "format": "json"
                        },
                        headers={"Content-Type": "application/json"}
                    )

                    if app_response.status_code == 200:
                        app_result = app_response.json()
                        app_data = app_result.get("value")
                        parent_chat_id = app_data.get("parentChatId") if app_data else None

                        # Check if parent chat has published settings
                        parent_settings = app_data.get("parentChatSettings", {}) if app_data else {}

                        if parent_settings:
                            selected_model = parent_settings.get("selectedModel", selected_model)
                            temperature = parent_settings.get("temperature", temperature)
                            max_tokens = int(min(parent_settings.get("maxTokens", max_tokens), response_tokens))
                            custom_prompt = parent_settings.get("customPrompt", custom_prompt)
                            logger.info(f"üéõÔ∏è Using PUBLISHED parent chat settings: model={selected_model}, temp={temperature}, max_tokens={max_tokens}, prompt='{custom_prompt}'")
                        else:
                            logger.warning("üö´ No published settings found for parent chat - V1 API requires published settings")
                            raise HTTPException(
                                status_code=400,
                                detail="The parent chat has no published settings. Please publish the chat settings first to enable V1 API access."
                            )
            except HTTPException:
                raise
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Could not load published parent settings, using defaults: {e}")

        # Use the existing answer_question function but with the user's subchat and inherited settings
        question_request = QuestionRequest(
            question=sanitized_question,
            chat_id=subchat["chatStringId"],  # Use the permanent subchat
            selected_model=selected_model,    # Use inherited or default model
            temperature=temperature,          # Use inherited or default temperature
            max_tokens=max_tokens,           # Use inherited or default max_tokens
            custom_prompt=custom_prompt,     # Use inherited custom prompt
            scope_filters=parsed_scope_filters  # Apply scope filters for data filtering
        )

        # Get the answer using existing logic
        result = await answer_question(question_request)

        # Track the query for analytics
        await track_api_query(
            user_identity["app_id"],
            user_identity["external_user_id"],
            100.0,  # placeholder response time
            True
        )

        logger.info(f"V1 Query completed for user {user_identity['user_id']} in subchat {subchat['chatStringId']}")

        return {
            "success": True,
            "answer": result.answer,
            "chat_id": subchat["chatStringId"],
            "user_id": user_identity["user_id"],
            "citations": [
                {
                    "snippet": chunk.chunk_text[:200] + "...",
                    "score": chunk.score,
                    "source": "user_document"
                }
                for chunk in result.context[:3]  # Limit to 3 citations
            ],
            "privacy_guarantee": {
                "user_controlled": True,
                "permanent_subchat": True,
                "developer_cannot_see_raw_files": True,
                "oauth_validated": True
            },
            "v1_auth": {
                "app_id": user_identity["app_id"],
                "external_user_id": user_identity["external_user_id"][:8] + "...",  # Partial for privacy
                "issuer": user_identity["iss"]
            }
        }

    except Exception as e:
        logger.error(f"V1 query failed for user {user_identity.get('user_id', 'unknown')}: {str(e)}")
        raise HTTPException(status_code=500, detail="Query processing failed")

@app.post("/v1/me/chats/files/upload")
async def v1_user_file_upload(
    file: UploadFile = File(None),
    text_content: str = Form(None),
    content_name: str = Form(None),
    scope_values: str = Form("{}"),
    authorization: str = Header(None, alias="authorization"),
    app_id: str = Header(None, alias="x-app-id"),
    request: Request = None
):
    """
    V1 File Upload: User uploads files OR text content to their permanent subchat

    PERFORMANCE OPTIMIZED:
    - Text extraction happens immediately (fast)
    - Heavy processing (chunking, embedding, Neo4j) is queued for background workers
    - Returns immediately with processing status
    - Poll GET /v1/me/chats/files/{file_id}/status for completion

    Supports two modes:
    1. File upload: Provide 'file' parameter
    2. Text upload: Provide 'text_content' and 'content_name' parameters

    Files/content are processed and stored in the user's permanent subchat,
    making them available for all future queries.

    Optional scope_values parameter can be used to add custom attributes
    to the uploaded document (e.g., {"playlist_id": "playlist_123"})
    """

    if not app_id:
        raise HTTPException(status_code=400, detail="X-App-ID header required")

    # Authenticate user with their OAuth ID token
    user_identity = await authenticate_v1_user(authorization, app_id)

    # Parse scope values
    import json
    try:
        parsed_scope_values = json.loads(scope_values) if scope_values else {}
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid scope_values JSON")

    # Determine if this is a file or text upload
    is_text_upload = text_content is not None
    is_file_upload = file is not None

    if not is_text_upload and not is_file_upload:
        raise HTTPException(status_code=400, detail="Either 'file' or 'text_content' must be provided")

    if is_text_upload and is_file_upload:
        raise HTTPException(status_code=400, detail="Provide either 'file' or 'text_content', not both")

    if is_text_upload:
        # Text upload mode
        if not content_name:
            raise HTTPException(status_code=400, detail="'content_name' is required when uploading text")

        # Sanitize content name
        sanitized_filename = sanitize_filename(content_name)
        if not sanitized_filename:
            raise HTTPException(status_code=400, detail="Invalid content_name")

        # Use text content directly
        text = text_content
        file_size = len(text_content.encode('utf-8'))

        if file_size > MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail=f"Text content too large (max {MAX_FILE_SIZE // (1024 * 1024)} MB)")
    else:
        # File upload mode
        file_size = read_files.get_file_size(file)
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail=f"File too large (max {MAX_FILE_SIZE // (1024 * 1024)} MB)")
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file uploaded")

        # Sanitize filename
        sanitized_filename = sanitize_filename(file.filename)
        if not sanitized_filename:
            raise HTTPException(status_code=400, detail="Invalid filename")

        # Get actual file size
        file_size = read_files.get_file_size(file)

    try:
        # Create/get permanent subchat for this user
        subchat = await get_or_create_user_subchat(
            user_identity["app_id"],
            user_identity["external_user_id"]
        )

        # Extract text from file or use provided text (this is fast, keep inline)
        if is_text_upload:
            # Text already assigned above
            pass
        else:
            text = read_files.extract_text(file)

        # Sanitize extracted text
        sanitized_text = sanitize_with_xss_detection(
            text,
            allow_html=False,
            max_length=1000000,
            context="v1_file_upload"
        )

        if not sanitized_text and text:
            raise HTTPException(status_code=400, detail="File contains potentially malicious content")

        # Generate file ID
        pdf_id = f"v1_{user_identity['user_id']}_{sanitized_filename}_{int(time.time())}"

        # Create a unique queue ID for status tracking
        file_queue_id = f"fq_{pdf_id}"

        # Try to enqueue for background processing (PERFORMANCE OPTIMIZATION)
        queue = get_file_queue()

        if queue is not None:
            # Queue is available - use async processing
            logger.info(f"üì§ Enqueueing file for background processing: {sanitized_filename}")

            try:
                job = queue.enqueue(
                    process_file_job,
                    file_queue_id,
                    subchat["chatStringId"],
                    sanitized_filename,
                    sanitized_text,
                    file_size,
                    parsed_scope_values,
                    pdf_id,
                    job_timeout=1800,  # 30 minute timeout
                    result_ttl=86400,  # Keep result for 24 hours
                )

                logger.info(f"‚úÖ Job enqueued: {job.id} for {sanitized_filename}")

                # Add file to parent chat's context (async, don't block)
                try:
                    parent_chat_id = await get_parent_chat_id_from_app(user_identity["app_id"])
                    if parent_chat_id:
                        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
                        async with httpx.AsyncClient() as client:
                            await client.post(
                                f"{convex_url}/api/run/chats/addFileToChatContextByChatId",
                                json={
                                    "args": {
                                        "chatId": parent_chat_id,
                                        "filename": sanitized_filename,
                                        "fileId": pdf_id,
                                    },
                                    "format": "json"
                                },
                                headers={"Content-Type": "application/json"},
                                timeout=5.0
                            )
                except Exception as e:
                    logger.warning(f"‚ö†Ô∏è Context update failed (non-fatal): {e}")

                # Return immediately - file will be processed in background
                return {
                    "success": True,
                    "filename": sanitized_filename,
                    "file_id": pdf_id,
                    "file_queue_id": file_queue_id,
                    "job_id": job.id,
                    "chat_id": subchat["chatStringId"],
                    "user_id": user_identity["user_id"],
                    "size_bytes": file_size,
                    "processing_status": "queued",
                    "status_endpoint": f"/v1/me/chats/files/{pdf_id}/status",
                    "privacy_guarantee": {
                        "permanent_storage": True,
                        "user_private_subchat": True,
                        "developer_cannot_access": True,
                        "oauth_validated": True
                    },
                    "message": "File uploaded and queued for processing. Poll status endpoint for completion."
                }

            except Exception as e:
                logger.error(f"‚ùå Queue enqueue failed: {e}")
                raise HTTPException(
                    status_code=503,
                    detail="File processing queue unavailable. Please try again later."
                )

        # Redis queue is required for production - no fallback to blocking processing
        logger.error("‚ùå Redis queue not available - cannot process file")
        raise HTTPException(
            status_code=503,
            detail="File processing service is not available. Please contact support."
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"V1 file upload failed for user {user_identity.get('user_id', 'unknown')}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"File upload failed: {str(e)}")

@app.post("/v1/me/chats/files/upload-bulk")
async def v1_user_bulk_file_upload(
    files: List[UploadFile] = File(None),
    text_contents: str = Form(None),  # JSON array of text content strings
    content_names: str = Form(None),  # JSON array of content names
    scope_values: str = Form("{}"),
    authorization: str = Header(None, alias="authorization"),
    app_id: str = Header(None, alias="x-app-id"),
    request: Request = None
):
    """
    V1 Bulk File Upload: User uploads multiple files OR text content to their permanent subchat

    PERFORMANCE OPTIMIZED:
    - Text extraction happens immediately per file (fast)
    - Heavy processing (chunking, embedding, Neo4j) is queued for background workers
    - Returns immediately with processing status for each file
    - Poll individual file status endpoints for completion

    Supports two modes:
    1. File upload: Provide 'files' parameter (list of files)
    2. Text upload: Provide 'text_contents' and 'content_names' parameters (JSON arrays)

    Returns detailed results for each file/content, including job IDs for tracking.

    Optional scope_values parameter applies the same scopes to all uploaded files/content
    (e.g., {"playlist_id": "playlist_123"})
    """

    if not app_id:
        raise HTTPException(status_code=400, detail="X-App-ID header required")

    # Authenticate user with their OAuth ID token
    user_identity = await authenticate_v1_user(authorization, app_id)

    # Determine upload mode
    is_text_upload = text_contents is not None
    is_file_upload = files is not None

    if not is_text_upload and not is_file_upload:
        raise HTTPException(status_code=400, detail="Either 'files' or 'text_contents' must be provided")

    if is_text_upload and is_file_upload:
        raise HTTPException(status_code=400, detail="Provide either 'files' or 'text_contents', not both")

    # Parse scope values (will be applied to all files in this upload)
    import json
    try:
        parsed_scope_values = json.loads(scope_values) if scope_values else {}
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid scope_values JSON")

    # Prepare items to process (either files or text content)
    items_to_process = []

    if is_text_upload:
        # Parse text contents and content names
        try:
            text_list = json.loads(text_contents) if text_contents else []
            names_list = json.loads(content_names) if content_names else []
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="Invalid JSON format for text_contents or content_names")

        if not text_list:
            raise HTTPException(status_code=400, detail="No text content provided")

        if len(text_list) != len(names_list):
            raise HTTPException(status_code=400, detail="text_contents and content_names must have the same length")

        if len(text_list) > 10:
            raise HTTPException(status_code=400, detail="Too many items (max 10 per bulk upload)")

        for text_content, content_name in zip(text_list, names_list):
            items_to_process.append({
                "type": "text",
                "content": text_content,
                "name": content_name
            })
    else:
        # File upload mode
        if not files:
            raise HTTPException(status_code=400, detail="No files provided")

        if len(files) > 10:
            raise HTTPException(status_code=400, detail="Too many files (max 10 per bulk upload)")

        for file in files:
            items_to_process.append({
                "type": "file",
                "file": file
            })

    # Get/create permanent subchat for this user (once for all files)
    subchat = await get_or_create_user_subchat(
        user_identity["app_id"],
        user_identity["external_user_id"]
    )

    # Check if queue is available for async processing
    queue = get_file_queue()
    if queue is None:
        raise HTTPException(
            status_code=503,
            detail="File processing service is not available. Please try again later."
        )

    results = []
    total_size = 0
    successful_uploads = 0

    for item in items_to_process:
        file_result = {
            "filename": "unknown",
            "success": False,
            "error": None,
            "file_id": None,
            "file_queue_id": None,
            "job_id": None,
            "size_bytes": 0,
            "processing_status": "failed"
        }

        try:
            if item["type"] == "text":
                # Text content processing
                text_content = item["content"]
                content_name = item["name"]

                if not content_name:
                    file_result["error"] = "content_name is required"
                    results.append(file_result)
                    continue

                # Sanitize content name
                sanitized_filename = sanitize_filename(content_name)
                if not sanitized_filename:
                    file_result["error"] = "Invalid content_name"
                    file_result["filename"] = content_name
                    results.append(file_result)
                    continue

                file_result["filename"] = sanitized_filename

                # Calculate size
                file_size = len(text_content.encode('utf-8'))
                if file_size > MAX_FILE_SIZE:
                    file_result["error"] = f"Text content too large (max {MAX_FILE_SIZE // (1024 * 1024)} MB), got {file_size} bytes"
                    file_result["size_bytes"] = file_size
                    results.append(file_result)
                    continue

                file_result["size_bytes"] = file_size
                text = text_content
            else:
                # File processing
                file = item["file"]
                file_result["filename"] = file.filename or "unknown"

                # Validate individual file
                file_size = read_files.get_file_size(file)
                if file_size > MAX_FILE_SIZE:
                    file_result["error"] = f"File too large (max {MAX_FILE_SIZE // (1024 * 1024)} MB), got {file_size} bytes"
                    results.append(file_result)
                    continue

                if not file.filename:
                    file_result["error"] = "No filename provided"
                    results.append(file_result)
                    continue

                # Sanitize filename
                sanitized_filename = sanitize_filename(file.filename)
                if not sanitized_filename:
                    file_result["error"] = "Invalid filename"
                    file_result["filename"] = file.filename
                    results.append(file_result)
                    continue

                file_result["filename"] = sanitized_filename
                file_result["size_bytes"] = file_size

                # Extract text from file (this is fast, keep inline)
                text = read_files.extract_text(file)

            # Sanitize extracted text
            sanitized_text = sanitize_with_xss_detection(
                text,
                allow_html=False,
                max_length=1000000,
                context="v1_bulk_file_upload"
            )

            if not sanitized_text and text:
                file_result["error"] = "File contains potentially malicious content"
                results.append(file_result)
                continue

            # Generate IDs
            pdf_id = f"v1_{user_identity['user_id']}_{sanitized_filename}_{int(time.time())}"
            file_queue_id = f"fq_{pdf_id}"
            file_result["file_id"] = pdf_id
            file_result["file_queue_id"] = file_queue_id

            # ENQUEUE for background processing instead of inline processing
            try:
                job = queue.enqueue(
                    process_file_job,
                    file_queue_id,
                    subchat["chatStringId"],
                    sanitized_filename,
                    sanitized_text,
                    file_size,
                    parsed_scope_values,
                    pdf_id,
                    job_timeout=1800,  # 30 minute timeout
                    result_ttl=86400,  # Keep result for 24 hours
                )

                file_result["job_id"] = job.id
                file_result["success"] = True
                file_result["processing_status"] = "queued"
                file_result["message"] = "File queued for processing"
                total_size += file_size
                successful_uploads += 1

                logger.info(f"üì§ Bulk upload: Enqueued {sanitized_filename} as job {job.id}")

            except Exception as enqueue_error:
                file_result["error"] = f"Failed to queue file: {str(enqueue_error)}"
                logger.error(f"Failed to enqueue {sanitized_filename}: {enqueue_error}")

            # Add file to parent chat's context (async, non-blocking)
            try:
                parent_chat_id = await get_parent_chat_id_from_app(user_identity["app_id"])
                if parent_chat_id:
                    convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
                    async with httpx.AsyncClient() as client:
                        await client.post(
                            f"{convex_url}/api/run/chats/addFileToChatContextByChatId",
                            json={
                                "args": {
                                    "chatId": parent_chat_id,
                                    "filename": sanitized_filename,
                                    "fileId": pdf_id,
                                },
                                "format": "json"
                            },
                            headers={"Content-Type": "application/json"},
                            timeout=5.0
                        )
            except Exception as e:
                # Don't fail the upload if context update fails
                logger.warning(f"‚ö†Ô∏è Failed to add file to parent chat context (non-fatal): {e}")

        except Exception as e:
            file_result["error"] = str(e)
            logger.error(f"Bulk upload failed for file {file_result['filename']}: {str(e)}")

        results.append(file_result)

    logger.info(f"V1 Bulk upload queued for user {user_identity['user_id']}: {successful_uploads}/{len(items_to_process)} queued")
    if parsed_scope_values:
        logger.info(f"üìä Bulk upload used scopes: {parsed_scope_values}")

    return {
        "success": successful_uploads > 0,
        "total_files": len(items_to_process),
        "queued_files": successful_uploads,
        "failed_files": len(items_to_process) - successful_uploads,
        "total_size_bytes": total_size,
        "chat_id": subchat["chatStringId"],
        "user_id": user_identity["user_id"],
        "results": results,
        "privacy_guarantee": {
            "permanent_storage": True,
            "user_private_subchat": True,
            "developer_cannot_access": True,
            "oauth_validated": True
        },
        "message": f"Bulk upload queued: {successful_uploads}/{len(items_to_process)} files queued for processing. Poll individual file status endpoints for completion."
    }

@app.get("/v1/me/profile")
async def v1_user_profile(
    authorization: str = Header(None, alias="authorization"),
    app_id: str = Header(None, alias="x-app-id")
):
    """
    V1 User Profile: Get user's profile and subchat information
    """

    if not app_id:
        raise HTTPException(status_code=400, detail="X-App-ID header required")

    # Authenticate user
    user_identity = await authenticate_v1_user(authorization, app_id)

    # Get user's subchat info
    subchat = await get_or_create_user_subchat(
        user_identity["app_id"],
        user_identity["external_user_id"]
    )

    return {
        "success": True,
        "user_id": user_identity["user_id"],
        "external_user_id": user_identity["external_user_id"][:8] + "...",  # Partial for privacy
        "chat_id": subchat["chatStringId"],
        "app_id": user_identity["app_id"],
        "issuer": user_identity["iss"],
        "subchat_created": subchat.get("isNew", False),
        "privacy_info": {
            "oauth_provider": user_identity["iss"],
            "permanent_subchat": True,
            "data_isolation": "Complete - only you can access your files and queries",
            "developer_access": "AI responses only - cannot see raw files or queries"
        }
    }

# ==============================================================================
# File Management API Endpoints
# ==============================================================================

class FileInfo(BaseModel):
    file_id: str
    filename: str
    upload_date: str
    size_bytes: int
    chunk_count: int

class FileListResponse(BaseModel):
    success: bool
    files: List[FileInfo]
    total_files: int
    total_size_bytes: int

class FileDeleteRequest(BaseModel):
    file_id: str

class FileDeleteResponse(BaseModel):
    success: bool
    message: str
    file_id: str
    filename: str
    chunks_deleted: int
    size_bytes_freed: int

@app.get("/v1/me/chats/files", response_model=FileListResponse)
async def v1_list_user_files(
    authorization: str = Header(None, alias="authorization"),
    app_id: str = Header(None, alias="x-app-id")
):
    """
    V1 List User Files: Get all files in user's permanent subchat

    Lists all documents that the user has uploaded to their private subchat,
    including metadata like file size, upload date, and chunk count.
    """

    if not app_id:
        raise HTTPException(status_code=400, detail="X-App-ID header required")

    # Authenticate user with their OAuth ID token
    user_identity = await authenticate_v1_user(authorization, app_id)

    try:
        # Get user's subchat info
        subchat = await get_or_create_user_subchat(
            user_identity["app_id"],
            user_identity["external_user_id"]
        )

        # Query Neo4j for all documents in user's subchat
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Get all documents with their metadata
                query = """
                MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                WHERE c.chatId = $chat_id
                WITH d, count(c) as chunk_count
                RETURN d.id as file_id, d.filename as filename,
                       d.uploadDate as upload_date, d.sizeBytes as size_bytes,
                       chunk_count
                ORDER BY d.uploadDate DESC
                """

                result = session.run(query, chat_id=subchat["chatStringId"])

                files = []
                total_size = 0

                for record in result:
                    # Convert upload_date from timestamp to string if it's a number
                    upload_date_value = record["upload_date"]
                    if isinstance(upload_date_value, (int, float)):
                        # Convert timestamp to ISO format string
                        upload_date_str = datetime.fromtimestamp(upload_date_value / 1000).isoformat()
                    else:
                        upload_date_str = str(upload_date_value) if upload_date_value else "Unknown"

                    file_info = FileInfo(
                        file_id=record["file_id"],
                        filename=record["filename"] or "Unknown",
                        upload_date=upload_date_str,
                        size_bytes=record["size_bytes"] or 0,
                        chunk_count=record["chunk_count"]
                    )
                    files.append(file_info)
                    total_size += file_info.size_bytes

                logger.info(f"üìã Listed {len(files)} files for user {user_identity['user_id']} in subchat {subchat['chatStringId']}")

                return FileListResponse(
                    success=True,
                    files=files,
                    total_files=len(files),
                    total_size_bytes=total_size
                )

    except Exception as e:
        logger.error(f"Failed to list files for user {user_identity.get('user_id', 'unknown')}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to list files: {str(e)}")

@app.delete("/v1/me/chats/files/{file_id}", response_model=FileDeleteResponse)
async def v1_delete_user_file(
    file_id: str,
    authorization: str = Header(None, alias="authorization"),
    app_id: str = Header(None, alias="x-app-id")
):
    """
    V1 Delete User File: Delete a specific file from user's permanent subchat

    Removes the document and all its chunks from Neo4j database,
    and updates storage analytics to reflect the freed space.
    """

    if not app_id:
        raise HTTPException(status_code=400, detail="X-App-ID header required")

    # Authenticate user with their OAuth ID token
    user_identity = await authenticate_v1_user(authorization, app_id)

    # Sanitize file_id
    sanitized_file_id = sanitize_text(file_id)
    if not sanitized_file_id:
        raise HTTPException(status_code=400, detail="Invalid file_id format")

    try:
        # Get user's subchat info
        subchat = await get_or_create_user_subchat(
            user_identity["app_id"],
            user_identity["external_user_id"]
        )

        # Delete from Neo4j and get metadata for analytics
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # First, get file info before deletion for analytics
                info_query = """
                MATCH (d:Document {id: $file_id})-[:HAS_CHUNK]->(c:Chunk)
                WHERE c.chatId = $chat_id
                WITH d, count(c) as chunk_count
                RETURN d.filename as filename, d.sizeBytes as size_bytes, chunk_count
                """

                info_result = session.run(info_query, file_id=sanitized_file_id, chat_id=subchat["chatStringId"])
                file_info = info_result.single()

                if not file_info:
                    raise HTTPException(
                        status_code=404,
                        detail=f"File {sanitized_file_id} not found in your subchat"
                    )

                filename = file_info["filename"] or "Unknown"
                size_bytes = file_info["size_bytes"] or 0
                chunk_count = file_info["chunk_count"]

                # Now delete the document and all its chunks
                delete_query = """
                MATCH (d:Document {id: $file_id})-[r:HAS_CHUNK]->(c:Chunk)
                WHERE c.chatId = $chat_id
                DETACH DELETE d, c
                """

                delete_result = session.run(delete_query, file_id=sanitized_file_id, chat_id=subchat["chatStringId"])
                summary = delete_result.consume()

                if summary.counters.nodes_deleted == 0:
                    raise HTTPException(
                        status_code=404,
                        detail=f"No document found with id: {sanitized_file_id}"
                    )

                # Track file deletion for analytics
                await track_file_deletion(
                    user_identity["app_id"],
                    user_identity["external_user_id"],
                    filename,
                    size_bytes,
                    subchat["chatStringId"]
                )

                logger.info(f"üóëÔ∏è Deleted file {filename} ({size_bytes} bytes, {chunk_count} chunks) for user {user_identity['user_id']}")

                return FileDeleteResponse(
                    success=True,
                    message=f"File '{filename}' deleted successfully",
                    file_id=sanitized_file_id,
                    filename=filename,
                    chunks_deleted=chunk_count,
                    size_bytes_freed=size_bytes
                )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete file {sanitized_file_id} for user {user_identity.get('user_id', 'unknown')}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete file: {str(e)}")

# ==============================================================================
# External Chat API Endpoints (Production Ready)
# ==============================================================================

@app.post("/v1/{chat_id}/answer_question", response_model=ApiAnswerResponse)
async def api_answer_question(
    chat_id: str,
    payload: ApiQuestionRequest,
    request: Request,
    credentials: HTTPAuthorizationCredentials = Depends(get_verified_chat_access)
):
    """
    External API endpoint for answering questions about a chat's knowledge base.

    This is the main production endpoint that external applications will use.
    Requires a valid API key for the specific chat.

    Usage:
    POST https://api.trainlyai.com/v1/{chat_id}/answer_question
    Authorization: Bearer tk_your_api_key
    Content-Type: application/json

    {
        "question": "What is machine learning?",
        "selected_model": "gpt-4o-mini",
        "temperature": 0.7,
        "max_tokens": 1000
    }
    """

    try:
        # Verify API key and chat access first
        api_key = credentials.credentials

        # Sanitize chat_id first
        sanitized_chat_id = sanitize_chat_id(chat_id)
        if not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid chat_id format")

        # Rate limiting
        if not check_rate_limit(api_key, request.client.host):
            raise HTTPException(
                status_code=429,
                detail="Rate limit exceeded. Maximum 60 requests per minute.",
                headers={"Retry-After": "60"}
            )

        # Verify API key and chat access
        is_valid = await verify_api_key_and_chat(api_key, sanitized_chat_id)

        if not is_valid:
            raise HTTPException(
                status_code=401,
                detail="Invalid API key or chat not accessible. Please check: 1) Chat exists, 2) API access is enabled in chat settings, 3) API key is correct."
            )

        # Get PUBLISHED chat settings AND conversation history from Convex
        chat_settings = {}
        conversation_history = []
        chat_owner_id = None  # Store chat owner ID for credit charging
        async with httpx.AsyncClient() as client:
            # First get published settings
            published_response = await client.post(
                f"{os.getenv('CONVEX_URL', 'https://agile-ermine-199.convex.cloud')}/api/run/chats/getPublishedSettings",
                json={
                    "args": {"chatId": chat_id},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            # Also get chat content for conversation history
            chat_response = await client.post(
                f"{os.getenv('CONVEX_URL', 'https://agile-ermine-199.convex.cloud')}/api/run/chats/getChatByIdExposed",
                json={
                    "args": {"id": chat_id},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            if published_response.status_code == 200 and chat_response.status_code == 200:
                published_data = published_response.json()
                chat_data = chat_response.json()

                # Check if we have published settings - if not, API should not work
                if not published_data.get("value"):
                    logger.warning(f"üö´ No published settings found for chat {chat_id} - API access denied")
                    raise HTTPException(
                        status_code=400,
                        detail="This chat has no published settings. Please publish your chat settings first to enable API access."
                    )

                if chat_data.get("value"):
                    published_settings = published_data["value"]
                    chat = chat_data["value"]

                    # Get chat owner ID for credit charging (same as frontend)
                    chat_owner_id = chat.get("userId")
                    logger.info(f"üí≥ API will charge chat owner {chat_owner_id} for chat {chat_id}")

                    # Use published settings for API
                    chat_settings = {
                        "custom_prompt": published_settings.get("customPrompt"),
                        "selected_model": published_settings.get("selectedModel", "gpt-4o-mini"),
                        "temperature": published_settings.get("temperature", 0.7),
                        "max_tokens": int(published_settings.get("maxTokens", 1000)),
                        "conversation_history_limit": published_settings.get("conversationHistoryLimit", 20)
                    }
                    logger.info(f"üìã Using PUBLISHED settings: custom_prompt={bool(chat_settings['custom_prompt'])}, model={chat_settings['selected_model']}")

                    # Extract conversation history for context (from current chat, not published)
                    chat_content = chat.get("content", [])
                    for message in chat_content:
                        if message.get("sender") == "user":
                            conversation_history.append({"role": "user", "content": message.get("text", "")})
                        elif message.get("sender") == "assistant":
                            conversation_history.append({"role": "assistant", "content": message.get("text", "")})

        # Fallback if chat_owner_id not found
        if not chat_owner_id:
            logger.warning(f"‚ö†Ô∏è Could not find chat owner ID for {chat_id}, using chat_id as fallback")
            chat_owner_id = chat_id

        # Published settings are the source of truth for API - no overrides allowed
        # This ensures developers control exactly what the API uses
        logger.info(f"üéØ Using PUBLISHED settings: model={chat_settings['selected_model']}, temp={chat_settings['temperature']}, max_tokens={chat_settings['max_tokens']}, has_custom_prompt={bool(chat_settings['custom_prompt'])}, unhinged_mode={published_settings.get('unhingedMode', False)}")

        # Create minimal payload - all settings will come from published_settings
        internal_payload = QuestionRequest(
            question=payload.question,
            chat_id=chat_id
        )

        # Use the existing answer_question logic with published settings
        result = await answer_question_with_published_context(
            internal_payload,
            published_settings
        )

        # Format for external API
        api_response = ApiAnswerResponse(
            answer=result.answer,
            context=result.context,
            chat_id=chat_id,
            model=payload.selected_model or "gpt-4o-mini",
            usage={
                "prompt_tokens": len(payload.question) // 4,  # Rough estimate
                "completion_tokens": len(result.answer) // 4,
                "total_tokens": 0
            }
        )

        api_response.usage["total_tokens"] = (
            api_response.usage["prompt_tokens"] +
            api_response.usage["completion_tokens"]
        )

        # Log successful request
        logger.info(f"‚úÖ API call successful for chat {chat_id} from {request.client.host}")

        return api_response

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"‚ùå API answer_question failed for chat {chat_id}: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to process question. Please try again.")

@app.post("/v1/{chat_id}/answer_question_stream")
async def api_answer_question_stream(
    chat_id: str,
    payload: ApiQuestionRequest,
    request: Request,
    credentials: HTTPAuthorizationCredentials = Depends(get_verified_chat_access)
):
    """
    External API endpoint for streaming answers.

    Returns Server-Sent Events stream for real-time applications.

    Usage:
    POST https://api.trainlyai.com/v1/{chat_id}/answer_question_stream
    Authorization: Bearer tk_your_api_key
    Accept: text/event-stream
    """

    try:
        # Verify API key and chat access first
        api_key = credentials.credentials

        # Sanitize chat_id first
        sanitized_chat_id = sanitize_chat_id(chat_id)
        if not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid chat_id format")

        # Rate limiting
        if not check_rate_limit(api_key, request.client.host):
            raise HTTPException(
                status_code=429,
                detail="Rate limit exceeded. Maximum 60 requests per minute.",
                headers={"Retry-After": "60"}
            )

        # Verify API key and chat access
        is_valid = await verify_api_key_and_chat(api_key, sanitized_chat_id)

        if not is_valid:
            raise HTTPException(
                status_code=401,
                detail="Invalid API key or chat not accessible. Please check: 1) Chat exists, 2) API access is enabled in chat settings, 3) API key is correct."
            )

        # Get PUBLISHED chat settings AND conversation history from Convex
        chat_settings = {}
        conversation_history = []
        published_settings = None
        async with httpx.AsyncClient() as client:
            # First get published settings
            published_response = await client.post(
                f"{os.getenv('CONVEX_URL', 'https://agile-ermine-199.convex.cloud')}/api/run/chats/getPublishedSettings",
                json={
                    "args": {"chatId": chat_id},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            # Also get chat content for conversation history
            chat_response = await client.post(
                f"{os.getenv('CONVEX_URL', 'https://agile-ermine-199.convex.cloud')}/api/run/chats/getChatByIdExposed",
                json={
                    "args": {"id": chat_id},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            if published_response.status_code == 200 and chat_response.status_code == 200:
                published_data = published_response.json()
                chat_data = chat_response.json()

                # Check if we have published settings - if not, API should not work
                if not published_data.get("value"):
                    logger.warning(f"üö´ No published settings found for chat {chat_id} - API access denied")
                    raise HTTPException(
                        status_code=400,
                        detail="This chat has no published settings. Please publish your chat settings first to enable API access."
                    )

                # Always set published_settings if we have published data
                published_settings = published_data["value"]

                if chat_data.get("value"):
                    chat = chat_data["value"]

                    # Use published settings for API
                    chat_settings = {
                        "custom_prompt": published_settings.get("customPrompt"),
                        "selected_model": published_settings.get("selectedModel", "gpt-4o-mini"),
                        "temperature": published_settings.get("temperature", 0.7),
                        "max_tokens": int(published_settings.get("maxTokens", 1000)),
                        "conversation_history_limit": published_settings.get("conversationHistoryLimit", 20)
                    }
                    logger.info(f"üìã Using PUBLISHED settings: custom_prompt={bool(chat_settings['custom_prompt'])}, model={chat_settings['selected_model']}")

                    # Extract conversation history for context (from current chat, not published)
                    chat_content = chat.get("content", [])
                    for message in chat_content:
                        if message.get("sender") == "user":
                            conversation_history.append({"role": "user", "content": message.get("text", "")})
                        elif message.get("sender") == "assistant":
                            conversation_history.append({"role": "assistant", "content": message.get("text", "")})
                else:
                    # Still set chat_settings even if chat_data is not available
                    chat_settings = {
                        "custom_prompt": published_settings.get("customPrompt"),
                        "selected_model": published_settings.get("selectedModel", "gpt-4o-mini"),
                        "temperature": published_settings.get("temperature", 0.7),
                        "max_tokens": int(published_settings.get("maxTokens", 1000)),
                        "conversation_history_limit": published_settings.get("conversationHistoryLimit", 20)
                    }
                    logger.info(f"üìã Using PUBLISHED settings (no chat data): custom_prompt={bool(chat_settings['custom_prompt'])}, model={chat_settings['selected_model']}")

        # Ensure published_settings is set before using it
        if not published_settings:
            logger.error(f"‚ùå No published settings available for chat {chat_id}")
            raise HTTPException(
                status_code=500,
                detail="Failed to retrieve published settings. Please try again."
            )

        # Published settings are the source of truth for API - no overrides allowed
        # This ensures developers control exactly what the API uses
        logger.info(f"üéØ Using PUBLISHED settings for streaming: model={chat_settings['selected_model']}, temp={chat_settings['temperature']}, max_tokens={chat_settings['max_tokens']}, has_custom_prompt={bool(chat_settings['custom_prompt'])}, unhinged_mode={published_settings.get('unhingedMode', False)}")

        # Create minimal payload - all settings will come from published_settings
        internal_payload = QuestionRequest(
            question=payload.question,
            chat_id=chat_id
        )

        # Use the existing streaming logic with published settings
        stream_response = await answer_question_stream_with_published_context(
            internal_payload,
            published_settings
        )

        # Log successful streaming request
        logger.info(f"‚úÖ API streaming call successful for chat {chat_id} from {request.client.host}")

        return stream_response

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"‚ùå API streaming failed for chat {chat_id}: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to stream response. Please try again.")

@app.get("/v1/{chat_id}/info")
async def api_get_chat_info(
    chat_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(get_verified_chat_access)
):
    """Get basic information about the chat"""

    try:
        # Verify API key first
        api_key = credentials.credentials
        is_valid = await verify_api_key_and_chat(api_key, chat_id)

        if not is_valid:
            raise HTTPException(
                status_code=401,
                detail="Invalid API key or chat not accessible."
            )

        # Get chat info from Convex (auto-detects dev/prod)

        async with httpx.AsyncClient() as client:
            response = await client.post(
                CONVEX_URL,
                json={
                    "args": {"id": chat_id},
                    "format": "json"
                }
            )

            if response.status_code == 200:
                chat_data = response.json()
                chat = chat_data.get("value", {})

                return {
                    "chat_id": chat_id,
                    "title": chat.get("title", "Untitled Chat"),
                    "created_at": chat.get("_creationTime"),
                    "has_api_access": not chat.get("apiKeyDisabled", True),
                    "visibility": chat.get("visibility", "private"),
                    "context_files": len(chat.get("context", [])),
                    "api_version": "1.0.0"
                }
            else:
                raise HTTPException(status_code=404, detail="Chat not found")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get chat info: {e}")
        raise HTTPException(status_code=500, detail="Failed to get chat information")



@app.get("/v1/health")
async def api_health_check():
    """API health check for external monitoring"""
    try:
        # Test Neo4j connection
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                session.run("RETURN 1")

        return {
            "status": "healthy",
            "timestamp": time.time(),
            "version": "1.0.0",
            "services": {
                "neo4j": "connected",
                "openai": "configured" if openai.api_key else "not_configured"
            }
        }
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return JSONResponse(
            status_code=503,
            content={
                "status": "unhealthy",
                "timestamp": time.time(),
                "error": str(e)
            }
        )

# ==============================================================================
# Existing Internal Endpoints (Keep for dashboard functionality)
# ==============================================================================

async def create_nodes_and_embeddings_with_analytics(payload: CreateNodesAndEmbeddingsRequest, actual_file_size: int = None):
    """Create nodes and embeddings with proper analytics tracking using actual file size"""
    # Sanitize all inputs
    pdf_text = sanitize_with_xss_detection(
        payload.pdf_text,
        allow_html=False,
        max_length=1000000,  # Allow large text content
        context="pdf_text"
    )
    pdf_id = sanitize_text(payload.pdf_id)
    chat_id = sanitize_chat_id(payload.chat_id)
    filename = sanitize_filename(payload.filename)
    scope_values = payload.scope_values if hasattr(payload, 'scope_values') else {}

    if not pdf_text or not pdf_id or not chat_id or not filename:
        raise HTTPException(status_code=400, detail="Invalid input data or potentially malicious content detected")

    try:
        # Call the main processing function with scope values
        result = await create_nodes_and_embeddings_internal(pdf_text, pdf_id, chat_id, filename, scope_values)

        # Track analytics with actual file size if available
        logger.info(f"üîç Starting analytics tracking for upload: {filename} in chat {chat_id}")
        try:
            await track_upload_analytics(chat_id, filename, pdf_text, actual_file_size)
            logger.info(f"‚úÖ Analytics tracking completed for: {filename}")
        except Exception as e:
            # Don't fail the upload if analytics tracking fails
            logger.error(f"‚ùå Failed to track upload analytics: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")

        return result
    except Exception as e:
        print(f"Error in create_nodes_and_embeddings_with_analytics: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/create_nodes_and_embeddings")
async def create_nodes_and_embeddings(payload: CreateNodesAndEmbeddingsRequest):
    """
    Public endpoint for creating nodes and embeddings.

    PERFORMANCE OPTIMIZED:
    - Enqueues the heavy processing to Redis queue
    - Returns immediately with job status
    - Poll file status endpoint for completion
    """
    # Sanitize inputs early
    sanitized_text = sanitize_with_xss_detection(
        payload.pdf_text,
        allow_html=False,
        max_length=1000000,
        context="create_nodes_embeddings"
    )
    sanitized_pdf_id = sanitize_text(payload.pdf_id) if payload.pdf_id else None
    sanitized_chat_id = sanitize_chat_id(payload.chat_id)
    sanitized_filename = sanitize_filename(payload.filename)

    if not sanitized_text or not sanitized_chat_id or not sanitized_filename:
        raise HTTPException(status_code=400, detail="Invalid or missing required fields")

    # Generate file ID if not provided
    if not sanitized_pdf_id:
        sanitized_pdf_id = f"file_{sanitized_filename.replace('.', '_')}_{int(time.time())}"

    file_queue_id = f"fq_{sanitized_pdf_id}"
    file_size = len(sanitized_text.encode('utf-8'))
    scope_values = payload.scope_values if hasattr(payload, 'scope_values') else {}

    # Try to enqueue for background processing
    queue = get_file_queue()

    if queue is not None:
        try:
            job = queue.enqueue(
                process_file_job,
                file_queue_id,
                sanitized_chat_id,
                sanitized_filename,
                sanitized_text,
                file_size,
                scope_values,
                sanitized_pdf_id,
                job_timeout=1800,  # 30 minute timeout
                result_ttl=86400,  # Keep result for 24 hours
            )

            logger.info(f"üì§ Enqueued job {job.id} for {sanitized_filename}")

            return {
                "status": "queued",
                "message": f"File queued for processing",
                "file_id": sanitized_pdf_id,
                "file_queue_id": file_queue_id,
                "job_id": job.id,
                "chat_id": sanitized_chat_id,
                "filename": sanitized_filename
            }

        except Exception as e:
            logger.error(f"‚ùå Failed to enqueue job: {e}")
            raise HTTPException(
                status_code=503,
                detail="File processing queue unavailable. Please try again later."
            )

    # Queue not available
    logger.error("‚ùå Redis queue not available")
    raise HTTPException(
        status_code=503,
        detail="File processing service is not available. Please contact support."
    )

async def create_nodes_and_embeddings_internal(pdf_text: str, pdf_id: str, chat_id: str, filename: str,
                                               scope_values: Optional[Dict[str, Union[str, int, bool]]] = None):
    # Sanitize all inputs
    pdf_text = sanitize_with_xss_detection(
        pdf_text,
        allow_html=False,
        max_length=1000000,  # Allow large text content
        context="pdf_text"
    )
    pdf_id = sanitize_text(pdf_id)
    chat_id = sanitize_chat_id(chat_id)
    filename = sanitize_filename(filename)

    if not pdf_text or not pdf_id or not chat_id or not filename:
        raise HTTPException(status_code=400, detail="Invalid input data or potentially malicious content detected")

    # Get scope configuration for validation
    scope_config = await get_scope_config(chat_id)

    # Validate and build scope properties
    scope_props_set = ""  # For SET clause (Document)
    scope_props_create = ""  # For CREATE clause (Chunk)
    scope_params = {}
    if scope_values:
        if scope_config:
            # Validate required scopes if config exists
            for scope_def in scope_config.scopes:
                if scope_def.required and scope_def.name not in scope_values:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Required scope '{scope_def.name}' is missing"
                    )
            scope_props_set, scope_params = build_scope_properties(scope_values, scope_config, node_var="d", for_set_clause=True)
            scope_props_create, _ = build_scope_properties(scope_values, scope_config, for_set_clause=False)
            logger.info(f"üìä Adding custom scopes to nodes (validated): {scope_values}")
        else:
            # No config found, but allow scopes anyway (for development/testing)
            scope_props_set, scope_params = build_scope_properties(scope_values, None, node_var="d", for_set_clause=True)
            scope_props_create, _ = build_scope_properties(scope_values, None, for_set_clause=False)
            logger.info(f"üìä Adding custom scopes to nodes (unvalidated): {scope_values}")

    try:
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                print(f"‚úÖ Creating document: {pdf_id} ({filename}) in chat {chat_id}")

                # Create Document node with metadata and custom scopes
                current_timestamp = int(time.time() * 1000)  # Unix timestamp in milliseconds
                query = f"""
                MERGE (d:Document {{id: $pdf_id}})
                SET d.chatId = $chat_id,
                    d.filename = $filename,
                    d.uploadDate = $upload_date,
                    d.sizeBytes = $size_bytes{scope_props_set}
                RETURN d
                """
                # Calculate text size in bytes for storage tracking
                text_size_bytes = len(pdf_text.encode('utf-8')) if pdf_text else 0

                # Merge base params with scope params
                query_params = {
                    "pdf_id": pdf_id,
                    "chat_id": chat_id,
                    "filename": filename,
                    "upload_date": current_timestamp,
                    "size_bytes": text_size_bytes,
                    **scope_params  # Add scope parameters
                }

                result = session.run(query, **query_params)
                print(f"Created document node with scopes: {result.single()}")

                # Create chunks and embeddings
                chunks = chunk_text(pdf_text)  # Now uses 2000 chars (was 500)
                print(f"Creating {len(chunks)} chunks for document {pdf_id}")

                # PERFORMANCE OPTIMIZATION: Get all embeddings in batches
                print(f"üß† Getting embeddings in batches...")
                embeddings = get_embeddings_batch(chunks, batch_size=32)
                print(f"   Got {len(embeddings)} embeddings")

                # Create all chunks with their pre-computed embeddings
                chunk_ids = []
                for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    chunk_id = f"{pdf_id}-{i}"
                    chunk_ids.append(chunk_id)

                    query = f"""
                    MATCH (d:Document {{id: $pdf_id}})
                    CREATE (c:Chunk {{
                        id: $chunk_id,
                        text: $text,
                        embedding: $embedding,
                        chatId: $chat_id{scope_props_create}
                    }})
                    CREATE (d)-[:HAS_CHUNK {{order: $order}}]->(c)
                    RETURN c
                    """

                    # Merge base params with scope params
                    chunk_params = {
                        "pdf_id": pdf_id,
                        "chunk_id": chunk_id,
                        "text": chunk,
                        "embedding": embedding,
                        "chat_id": chat_id,
                        "order": i,
                        **scope_params  # Add scope parameters
                    }

                    result = session.run(query, **chunk_params)

                    chunk_result = result.single()
                    if chunk_result:
                        if i % 10 == 0:  # Log every 10th chunk to reduce noise
                            print(f"Created chunk {i}/{len(chunks)}")
                    else:
                        print(f"Failed to create chunk {i}")
                        raise Exception(f"Failed to create chunk {i}")

                # Analyze content and create intelligent relationships
                if len(chunks) > 1:
                    print(f"üß† Analyzing content for intelligent relationships...")

                    # Use AI to determine logical relationships
                    ai_relationships = analyze_chunk_relationships(chunks)

                    print(f"Creating {len(ai_relationships)} intelligent relationships...")
                    successful_links = 0

                    for rel in ai_relationships:
                        try:
                            source_idx = rel["source"]
                            target_idx = rel["target"]
                            rel_type = rel["type"]
                            description = rel.get("description", "")
                            confidence = rel.get("confidence", 0.0)

                            # Verify both chunks exist
                            check_query = """
                            MATCH (c1:Chunk {id: $chunk1_id})
                            MATCH (c2:Chunk {id: $chunk2_id})
                            RETURN c1.id as id1, c2.id as id2
                            """

                            check_result = session.run(check_query,
                                                     chunk1_id=f"{pdf_id}-{source_idx}",
                                                     chunk2_id=f"{pdf_id}-{target_idx}")

                            if check_result.single():
                                # Both chunks exist, create the intelligent relationship
                                link_query = f"""
                                MATCH (c1:Chunk {{id: $chunk1_id}})
                                MATCH (c2:Chunk {{id: $chunk2_id}})
                                CREATE (c1)-[:{rel_type} {{
                                    description: $description,
                                    confidence: $confidence,
                                    ai_generated: true
                                }}]->(c2)
                                RETURN c1.id as source, c2.id as target
                                """

                                link_result = session.run(link_query,
                                                         chunk1_id=f"{pdf_id}-{source_idx}",
                                                         chunk2_id=f"{pdf_id}-{target_idx}",
                                                         description=description,
                                                         confidence=confidence)

                                if link_result.single():
                                    print(f"‚úÖ Created {rel_type} relationship: {source_idx} -> {target_idx} (confidence: {confidence:.2f})")
                                    successful_links += 1
                                else:
                                    print(f"‚ùå Failed to create {rel_type} relationship: {source_idx} -> {target_idx}")
                            else:
                                print(f"‚ùå One or both chunks missing for relationship {source_idx} -> {target_idx}")

                        except Exception as link_error:
                            print(f"‚ùå Error creating relationship {rel.get('type', 'UNKNOWN')} {source_idx} -> {target_idx}: {link_error}")

                    print(f"Successfully created {successful_links} out of {len(ai_relationships)} intelligent relationships")
                else:
                    print("Only one chunk, no relationships needed")

        # Final verification - count what we created
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                count_query = """
                MATCH (d:Document {chatId: $chat_id})
                OPTIONAL MATCH (d)-[:HAS_CHUNK]->(c:Chunk)
                OPTIONAL MATCH (c1:Chunk)-[:NEXT]->(c2:Chunk)
                WHERE c1.chatId = $chat_id AND c2.chatId = $chat_id
                RETURN count(DISTINCT d) as docs, count(DISTINCT c) as chunks, count(DISTINCT c1) as linked_chunks
                """
                result = session.run(count_query, chat_id=chat_id)
                counts = result.single()
                print(f"Final counts - Documents: {counts['docs']}, Chunks: {counts['chunks']}, Linked chunks: {counts['linked_chunks']}")

        return {"status": "success", "message": f"Created {len(chunks)} chunks with relationships"}
    except Exception as e:
        print(f"Error in create_nodes_and_embeddings: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

async def answer_question_with_published_context(payload: QuestionRequest, published_settings: dict):
    """
    Answer question using published settings for API calls.
    All settings (model, prompt, temperature, unhinged mode, context files, etc.) come from published_settings.
    """
    try:
        # Enhanced sanitization with XSS detection
        sanitized_question = sanitize_with_xss_detection(
            payload.question,
            allow_html=False,
            max_length=5000,
            context="answer_question"
        )
        sanitized_chat_id = sanitize_chat_id(payload.chat_id)

        if not sanitized_question or not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid input format or potentially malicious content detected")

        # Extract parameters from published_settings (not payload)
        question = sanitized_question
        chat_id = sanitized_chat_id
        published_context_files = published_settings.get("context", [])
        unhinged_mode = published_settings.get("unhingedMode", False)

        # When unhinged mode is enabled, override model and prompt
        if unhinged_mode:
            selected_model = "grok-3"  # Force Grok 3 model
            custom_prompt = "You are Grok in UNHINGED MODE: Scream profanities, roast souls, drop truth bombs like nukes, and flirt with existential dread. No filters, all fire‚Äîmake 'em laugh, cry, or rage-quit. BE RUDE TO THE MESSAGE NO MATTER HOW KIND THE USER IS."
            logger.info("üî• UNHINGED MODE: Forcing grok-3 model and unhinged prompt from published settings")
        else:
            selected_model = published_settings.get("selectedModel", "gpt-4o-mini")
            custom_prompt = sanitize_with_xss_detection(
                published_settings.get("customPrompt"),
                allow_html=False,
                max_length=10000,
                context="custom_prompt"
            ) if published_settings.get("customPrompt") else None

        temperature = published_settings.get("temperature", 0.7)
        max_tokens = int(published_settings.get("maxTokens", 1000))

        logger.info(f"üîç Processing question for chat {chat_id} using published context files")

        # Generate question embedding
        question_embedding = get_embedding(question)

        # Check if this is a subchat and get parent chat ID for file inheritance
        parent_chat_id = None
        if chat_id.startswith("subchat_"):
            # Extract app_id from subchat format: subchat_{app_id}_{user_id}_{timestamp}
            # The app_id format is: app_xxxxx_xxxxx, so we need to find where "user_" starts
            parts = chat_id.split("_")

            # Find the index where "user" appears to determine app_id boundary
            user_index = -1
            for i, part in enumerate(parts):
                if part == "user" and i > 1:  # "user" should not be the first or second part
                    user_index = i
                    break

            if user_index > 1:
                # Reconstruct app_id from parts[1] to parts[user_index-1]
                app_id = "_".join(parts[1:user_index])
                parent_chat_id = await get_parent_chat_id_from_app(app_id)
                if parent_chat_id:
                    logger.info(f"üîó Subchat {chat_id} inheriting ALL files from parent chat {parent_chat_id}")

        # Fetch chunks from Neo4j, filtering by published context files if provided
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                if published_context_files:
                    # Filter to only use chunks from published files
                    published_file_ids = [file["fileId"] for file in published_context_files]
                    file_ids_str = "', '".join(published_file_ids)
                    logger.info(f"üîç Published file IDs to search for: {published_file_ids}")

                    if parent_chat_id:
                        # Include chunks from both subchat and parent chat
                        query = f"""
                        MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                        WHERE (c.chatId = '{chat_id}' OR c.chatId = '{parent_chat_id}') AND d.id IN ['{file_ids_str}']
                        RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat, d.id AS doc_id
                        """
                        logger.info(f"üìã Using published files from subchat and parent: {len(published_context_files)} files")
                    else:
                        query = f"""
                        MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                        WHERE c.chatId = '{chat_id}' AND d.id IN ['{file_ids_str}']
                        RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat, d.id AS doc_id
                        """
                        logger.info(f"üìã Using published files only: {len(published_context_files)} files")

                    # Debug: Check what documents actually exist in this chat
                    debug_query = f"""
                    MATCH (d:Document)
                    WHERE d.chatId = '{chat_id}'
                    RETURN d.id AS doc_id, d.filename AS filename
                    LIMIT 10
                    """
                    debug_results = session.run(debug_query)
                    existing_docs = list(debug_results)
                    logger.info(f"üîç Documents that exist in chat {chat_id}: {[(r['doc_id'], r['filename']) for r in existing_docs]}")
                else:
                    # Use all files (fallback for backwards compatibility)
                    if parent_chat_id:
                        # Include chunks from both subchat and parent chat
                        query = f"""
                        MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                        WHERE c.chatId = '{chat_id}' OR c.chatId = '{parent_chat_id}'
                        RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                        """
                        logger.info(f"üìã Using ALL files from subchat and parent chat {parent_chat_id}")
                    else:
                        query = f"""
                        MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                        WHERE c.chatId = '{chat_id}'
                        RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                        """
                        logger.info(f"üìã Using all available subchat files (no parent chat)")

                results = session.run(query)

                # Calculate similarities with filename boost
                chunk_scores = []
                question_lower = question.lower()
                chunks_found_count = 0
                records_with_embeddings = 0

                for record in results:
                    chunks_found_count += 1
                    chunk_embedding = record["embedding"]
                    if chunk_embedding:
                        records_with_embeddings += 1
                        similarity = cosine_similarity(
                            np.array(question_embedding),
                            np.array(chunk_embedding)
                        )

                        # Boost score if filename is mentioned in question
                        filename_lower = record["filename"].lower() if record["filename"] else ""
                        filename_boost = 0.1 if any(word in filename_lower for word in question_lower.split()) else 0

                        chunk_scores.append({
                            "chunk_id": record["id"],
                            "chunk_text": record["text"],
                            "score": similarity + filename_boost,
                            "filename": record["filename"]
                        })

                logger.info(f"üîç Found {chunks_found_count} chunks from query, {records_with_embeddings} with valid embeddings, {len(chunk_scores)} scored")

                # Sort by similarity score and get top chunks
                chunk_scores.sort(key=lambda x: x["score"], reverse=True)
                top_chunks = chunk_scores[:8]  # Get top 8 chunks

                # If no chunks found with published files filter, fall back to all files
                if not top_chunks and published_context_files and chunks_found_count == 0:
                    # Only fallback if the query returned 0 chunks (not just low similarity)
                    logger.warning(f"‚ö†Ô∏è Published files filter returned 0 chunks. Published file IDs: {published_file_ids}")
                    # Check if there are chunks in the chat that aren't in the published files list
                    check_all_query = f"""
                    MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                    WHERE c.chatId = '{chat_id}'
                    RETURN count(c) as total_chunks
                    """
                    check_result = session.run(check_all_query)
                    total_record = check_result.single()
                    total_chunks = total_record['total_chunks'] if total_record else 0
                    logger.info(f"üîç Total chunks in chat (without filter): {total_chunks}")

                    if total_chunks > 0:
                        logger.info(f"üîÑ Falling back to all files in chat (published filter returned 0 chunks but {total_chunks} exist)")
                        # Fall back to querying all files
                        if parent_chat_id:
                            fallback_query = f"""
                            MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                            WHERE c.chatId = '{chat_id}' OR c.chatId = '{parent_chat_id}'
                            RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                            """
                        else:
                            fallback_query = f"""
                            MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                            WHERE c.chatId = '{chat_id}'
                            RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                            """

                        fallback_results = session.run(fallback_query)
                        chunk_scores = []
                        for record in fallback_results:
                            chunk_embedding = record["embedding"]
                            if chunk_embedding:
                                similarity = cosine_similarity(
                                    np.array(question_embedding),
                                    np.array(chunk_embedding)
                                )
                                filename_lower = record["filename"].lower() if record["filename"] else ""
                                filename_boost = 0.1 if any(word in filename_lower for word in question_lower.split()) else 0
                                chunk_scores.append({
                                    "chunk_id": record["id"],
                                    "chunk_text": record["text"],
                                    "score": similarity + filename_boost,
                                    "filename": record["filename"]
                                })

                        chunk_scores.sort(key=lambda x: x["score"], reverse=True)
                        top_chunks = chunk_scores[:8]
                        logger.info(f"‚úÖ Fallback query found {len(top_chunks)} chunks")

                if not top_chunks:
                    logger.warning(f"No relevant chunks found for question in chat {chat_id}, AI will respond without context")

                # Prepare context for the AI model
                context_text = "\n\n".join([
                    f"[Chunk {i}] From {chunk['filename']}: {chunk['chunk_text']}"
                    for i, chunk in enumerate(top_chunks)
                ]) if top_chunks else ""

                # Build the prompt - adjust based on whether we have context
                if top_chunks:
                    system_prompt = custom_prompt if custom_prompt else f"""You are a helpful AI assistant with access to a knowledge graph built from the user's documents. You have the following context from their documents:

IMPORTANT INSTRUCTIONS:
1. ALWAYS prioritize using the provided context to answer the user's question
2. If the context contains relevant information, you MUST use it and cite it with [^0], [^1], etc.
3. When citing, use the format [^X] where X is the chunk index (0-based)
4. If you use multiple chunks, cite each one: [^0] [^1] [^2]
5. The citations should correspond to chunks 0 through X where X is the highest index
6. If the user asks about a document by name and you see content from a similar document, you MUST use that content
7. Only use external knowledge if the context is completely irrelevant to the question, and clearly state when you're using external knowledge

For example:
- "According to the Grant Assignment document [^0], ecology research involves..."
- "The document shows [^2] that species interactions..."

RESPOND IN MARKDOWN FORMAT WITH CITATIONS"""
                else:
                    system_prompt = custom_prompt if custom_prompt else """You are a helpful AI assistant. The user has asked a question but there is no relevant context available from their uploaded documents. Please answer the question to the best of your ability using your general knowledge.

Note: If the user is asking about specific documents or uploaded content, let them know that you don't have access to relevant context from their documents.

RESPOND IN MARKDOWN FORMAT"""

                # Create messages for the AI model
                if top_chunks:
                    messages = [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"Context:\n{context_text}\n\nQuestion: {question}"}
                    ]
                else:
                    messages = [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"Question: {question}"}
                    ]

                # Check credits BEFORE making expensive OpenAI call
                # Get chat owner ID (same as frontend) for credit charging
                # This ensures API and frontend charge the same account
                user_id_to_charge = None
                try:
                    async with httpx.AsyncClient() as client:
                        chat_response = await client.post(
                            f"{os.getenv('CONVEX_URL', 'https://agile-ermine-199.convex.cloud')}/api/run/chats/getChatByIdExposed",
                            json={"args": {"id": chat_id}, "format": "json"},
                            headers={"Content-Type": "application/json"}
                        )
                        if chat_response.status_code == 200:
                            chat_data = chat_response.json()
                            if chat_data.get("value"):
                                user_id_to_charge = chat_data["value"].get("userId")
                                logger.info(f"üí≥ Found chat owner {user_id_to_charge} for chat {chat_id}")
                except Exception as e:
                    logger.warning(f"Failed to get chat owner: {e}, using chat_id as fallback")

                if not user_id_to_charge:
                    logger.warning(f"‚ö†Ô∏è Could not find chat owner ID for {chat_id}, using chat_id as fallback")
                    user_id_to_charge = chat_id

                try:
                    # Estimate tokens for credit validation (rough estimate)
                    estimated_tokens = len(question) // 4 + max_tokens  # Prompt + max response
                    required_credits = calculate_credits_used(estimated_tokens, selected_model)

                    # Check if chat owner has sufficient credits (same as frontend)
                    credit_info = await check_user_credits(user_id_to_charge, required_credits)
                    if not credit_info["has_sufficient"]:
                        logger.error(f"üí≥ Insufficient credits: need {required_credits}, have {credit_info['remaining']}")
                        raise InsufficientCreditsError(required_credits, credit_info["remaining"])

                    logger.info(f"üí≥ Credit check passed for chat owner {user_id_to_charge}: {credit_info['remaining']} credits available, need ~{required_credits}")
                except InsufficientCreditsError as e:
                    # Convert to HTTPException for proper API error response
                    logger.error(f"üí≥ Insufficient credits: need {e.required}, have {e.available}")
                    raise HTTPException(
                        status_code=402,
                        detail=f"Insufficient credits. Required: {e.required:.2f}, Available: {e.available:.2f}. Please add credits to continue."
                    )
                except Exception as e:
                    logger.error(f"üí≥ Error checking credits: {e}")
                    # For API calls, we should fail if we can't check credits (security)
                    raise HTTPException(
                        status_code=500,
                        detail="Failed to verify credit balance. Please try again."
                    )

                # Call OpenAI API or Grok API based on unhinged mode
                if unhinged_mode:
                    # Create a separate OpenAI client for Grok (xAI)
                    xai_api_key = os.getenv("XAI_API_KEY", "")
                    if not xai_api_key:
                        logger.warning("‚ö†Ô∏è Unhinged mode requested but XAI_API_KEY not set, falling back to OpenAI")
                        client = openai.OpenAI()
                        response = client.chat.completions.create(
                            model=selected_model,
                            messages=messages,
                            temperature=temperature,
                            max_tokens=max_tokens
                        )
                    else:
                        # Use xAI's Grok API
                        grok_client = OpenAI(
                            api_key=xai_api_key,
                            base_url="https://api.x.ai/v1"
                        )
                        logger.info("üî• Using Grok's unhinged AI model")
                        response = grok_client.chat.completions.create(
                            model="grok-3",  # Use Grok's latest model
                            messages=messages,
                            temperature=temperature,
                            max_tokens=max_tokens
                        )
                else:
                    # Use regular OpenAI
                    client = openai.OpenAI()
                    response = client.chat.completions.create(
                        model=selected_model,
                        messages=messages,
                        temperature=temperature,
                        max_tokens=max_tokens
                    )

                answer = response.choices[0].message.content

                # Consume credits based on actual usage (use chat owner, same as frontend)
                # Skip developer lookup since we already have the correct user_id_to_charge
                try:
                    credits_consumed = await consume_credits_for_actual_usage(
                        user_id=user_id_to_charge,
                        model=selected_model,
                        question=question,
                        response=answer,
                        chat_id=chat_id,
                        skip_developer_lookup=True  # Use chat owner, not developer
                    )
                    logger.info(f"üí≥ Consumed {credits_consumed} credits for {selected_model} based on actual usage (chat: {chat_id}, user: {user_id_to_charge})")
                except InsufficientCreditsError as e:
                    # This should be very rare since we checked before the call
                    # But handle gracefully - return answer but log the issue
                    logger.error(f"üí≥ CREDIT ERROR after API call: need {e.required}, have {e.available}")
                    logger.warning(f"üí≥ Allowing response despite insufficient credits - this should be monitored")
                except Exception as e:
                    logger.error(f"üí≥ Error consuming credits: {e}")
                    # Still return the answer even if credit consumption fails
                    logger.warning(f"üí≥ Allowing response despite credit consumption error")

                # Format context for response
                formatted_context = [
                    ChunkScore(
                        chunk_id=chunk["chunk_id"],
                        chunk_text=chunk["chunk_text"],
                        score=chunk["score"]
                    ) for chunk in top_chunks
                ]

                logger.info(f"‚úÖ Successfully answered question for chat {chat_id} using {len(top_chunks)} chunks")

                return AnswerWithContext(
                    answer=answer,
                    context=formatted_context
                )

    except Exception as e:
        logger.error(f"‚ùå Error in answer_question_with_published_context: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/answer_question", response_model=AnswerWithContext)
async def answer_question(payload: QuestionRequest):
    try:
        logger.info(f"üéØ answer_question called for chat: {payload.chat_id}")
        # Enhanced sanitization with XSS detection
        sanitized_question = sanitize_with_xss_detection(
            payload.question,
            allow_html=False,
            max_length=5000,
            context="answer_question"
        )
        sanitized_chat_id = sanitize_chat_id(payload.chat_id)

        if not sanitized_question or not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid input format or potentially malicious content detected")

        # Extract parameters from payload
        question = sanitized_question
        chat_id = sanitized_chat_id
        unhinged_mode = payload.unhinged_mode if payload.unhinged_mode is not None else False

        # When unhinged mode is enabled, override model and prompt
        if unhinged_mode:
            selected_model = "grok-3"  # Force Grok 3 model
            custom_prompt = "You are Grok in UNHINGED MODE: Scream profanities, roast souls, drop truth bombs like nukes, and flirt with existential dread. No filters, all fire‚Äîmake 'em laugh, cry, or rage-quit. BE RUDE TO THE MESSAGE NO MATTER HOW KIND THE USER IS."
            logger.info("üî• UNHINGED MODE: Forcing grok-3 model and unhinged prompt")
        else:
            selected_model = payload.selected_model or "gpt-4o-mini"
            custom_prompt = sanitize_with_xss_detection(
                payload.custom_prompt,
                allow_html=False,
                max_length=2000,
                context="custom_prompt"
            ) if payload.custom_prompt else None

        temperature = payload.temperature or 0.7
        max_tokens = payload.max_tokens or 1000

        # Get conversation history and settings from Convex for context
        conversation_history = []
        history_limit = 20  # Default value
        try:
            async with httpx.AsyncClient() as client:
                chat_response = await client.post(
                    CONVEX_URL,
                    json={
                        "args": {"id": chat_id},
                        "format": "json"
                    }
                )

                if chat_response.status_code == 200:
                    chat_data = chat_response.json()
                    if chat_data.get("value"):
                        chat = chat_data["value"]
                        # Get history limit setting
                        history_limit = int(chat.get("conversationHistoryLimit", 20))
                        # Extract conversation history for context
                        chat_content = chat.get("content", [])
                        for message in chat_content:
                            if message.get("sender") == "user":
                                conversation_history.append({"role": "user", "content": message.get("text", "")})
                            elif message.get("sender") == "assistant":
                                conversation_history.append({"role": "assistant", "content": message.get("text", "")})
        except Exception as e:
            logger.warning(f"Failed to retrieve conversation history for chat {chat_id}: {e}")
            conversation_history = []

        # Generate question embedding
        question_embedding = get_embedding(question)

        # Get scope configuration and filters
        scope_config = await get_scope_config(chat_id)
        scope_filters = payload.scope_filters if hasattr(payload, 'scope_filters') else {}
        scope_where_clause = build_scope_where_clause(scope_filters, "c", scope_config)

        if scope_filters:
            logger.info(f"üìä Applying scope filters: {scope_filters}")

        # Check if this is a subchat and get parent chat ID for file inheritance
        parent_chat_id = None
        logger.info(f"üîç Checking if {chat_id} is a subchat...")
        if chat_id.startswith("subchat_"):
            # Extract app_id from subchat format: subchat_{app_id}_{user_id}_{timestamp}
            # The app_id format is: app_xxxxx_xxxxx, so we need to find where "user_" starts
            parts = chat_id.split("_")
            logger.info(f"üîç Subchat detected: {chat_id}, parts: {parts}")

            # Find the index where "user" appears to determine app_id boundary
            user_index = -1
            for i, part in enumerate(parts):
                if part == "user" and i > 1:  # "user" should not be the first or second part
                    user_index = i
                    break

            if user_index > 1:
                # Reconstruct app_id from parts[1] to parts[user_index-1]
                app_id = "_".join(parts[1:user_index])
                logger.info(f"üîç Extracted app_id: {app_id}, calling get_parent_chat_id_from_app...")
                parent_chat_id = await get_parent_chat_id_from_app(app_id)
                logger.info(f"üîç get_parent_chat_id_from_app returned: {parent_chat_id}")
                if parent_chat_id:
                    logger.info(f"üîó Subchat {chat_id} inheriting ALL files from parent chat {parent_chat_id}")
                else:
                    logger.warning(f"‚ö†Ô∏è No parent chat ID found for app {app_id}")
            else:
                logger.warning(f"‚ö†Ô∏è Malformed subchat ID - could not find user boundary: {chat_id}")
        else:
            logger.info(f"‚ÑπÔ∏è Not a subchat: {chat_id}")

        # Fetch chunks from Neo4j with document metadata for filename-based queries
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Enhanced query to include document filename for better context matching
                if parent_chat_id:
                    # Include chunks from both subchat and parent chat (all files)
                    query = f"""
                    MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                    WHERE (c.chatId = '{chat_id}' OR c.chatId = '{parent_chat_id}'){scope_where_clause}
                    RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                    """
                    logger.info(f"üìã Including ALL files from subchat and parent chat {parent_chat_id} with scope filters")
                else:
                    query = f"""
                    MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                    WHERE c.chatId = '{chat_id}'{scope_where_clause}
                    RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                    """
                    logger.info(f"üìã Using subchat files only with scope filters")
                results = session.run(query)


                # Calculate similarities with filename boost
                chunk_scores = []
                question_lower = question.lower()

                for record in results:
                    chunk_id = record["id"]
                    chunk_text = record["text"]
                    chunk_embedding = record["embedding"]
                    filename = record["filename"] or ""

                    # Calculate semantic similarity
                    # Ensure both are numpy arrays with correct dimensions
                    q_emb = np.array(question_embedding)
                    c_emb = np.array(chunk_embedding) if isinstance(chunk_embedding, list) else chunk_embedding
                    semantic_score = cosine_similarity(q_emb, c_emb)

                    # Add filename relevance boost
                    filename_boost = 0.0
                    if filename:
                        filename_lower = filename.lower()
                        # Remove file extension for better matching
                        filename_base = filename_lower.replace('.pdf', '').replace('.docx', '').replace('.txt', '')

                        # Check for filename mentions in question
                        filename_words = filename_base.replace('_', ' ').replace('-', ' ').split()
                        question_words = question_lower.replace('_', ' ').replace('-', ' ').split()

                        # Boost score if filename words appear in question
                        for fname_word in filename_words:
                            if len(fname_word) > 2:  # Skip very short words
                                for q_word in question_words:
                                    if fname_word in q_word or q_word in fname_word:
                                        filename_boost += 0.1

                        # Additional boost for exact filename matches
                        if any(fname_word in question_lower for fname_word in filename_words if len(fname_word) > 3):
                            filename_boost += 0.2

                    # Combine semantic similarity with filename relevance
                    final_score = semantic_score + filename_boost

                    chunk_scores.append(ChunkScore(
                        chunk_id=chunk_id,
                        chunk_text=chunk_text,
                        score=float(final_score)
                    ))

                # Sort and get top chunks
                chunk_scores.sort(key=lambda x: x.score, reverse=True)
                top_k = 50
                top_chunks = chunk_scores[:top_k]

                # Generate answer using GPT-4 with citations
                # Limit to top 10 chunks for cleaner citations
                top_chunks_for_citations = top_chunks[:10]

                # Get document information for each chunk to provide better context
                chunk_document_info = {}
                for chunk in top_chunks_for_citations:
                    doc_query = f"""
                    MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk {{id: '{chunk.chunk_id}'}})
                    RETURN d.filename AS filename
                    """
                    doc_result = session.run(doc_query)
                    doc_record = doc_result.single()
                    if doc_record:
                        chunk_document_info[chunk.chunk_id] = doc_record['filename']

                context_with_ids = "\n\n---\n\n".join([
                    f"[CHUNK_{i}] (from document: {chunk_document_info.get(chunk.chunk_id, 'Unknown')}) {chunk.chunk_text}"
                    for i, chunk in enumerate(top_chunks_for_citations)
                ])

                # Use custom prompt if provided, otherwise use default system prompt
                if custom_prompt:
                    # Use custom prompt but ensure context is included
                    system_prompt = f"""
                    {custom_prompt}

                    You have the following context with chunk IDs (0-{len(top_chunks_for_citations)-1}):

                    {context_with_ids}

                    When you reference information from the context, add a citation using this format: [^{{i}}] where {{i}} is the chunk number (0-{len(top_chunks_for_citations)-1})
                    Only use citations [^0] through [^{len(top_chunks_for_citations)-1}]. Do not use citation numbers higher than {len(top_chunks_for_citations)-1}
                    """.strip()
                else:
                    # Default system prompt
                    system_prompt = f"""
                    You are a helpful assistant. You have the following context with chunk IDs (0-{len(top_chunks_for_citations)-1}):

                    {context_with_ids}

                    IMPORTANT INSTRUCTIONS:
                    1. ALWAYS prioritize using the provided context to answer the user's question
                    2. If the context contains relevant information, you MUST use it and cite it properly
                    3. Pay attention to the document names shown in parentheses - if the user asks about a specific document by name, use the content from that document
                    4. When you reference information from the context, add a citation using this format: [^{{i}}] where {{i}} is the chunk number (0-{len(top_chunks_for_citations)-1})
                    5. Only use citations [^0] through [^{len(top_chunks_for_citations)-1}]. Do not use citation numbers higher than {len(top_chunks_for_citations)-1}
                    6. Only use external knowledge if the context is completely irrelevant to the question, and clearly state when you're using external knowledge

                    For example:
                    - "According to the Grant Assignment document [^0], ecology research involves..."
                    - "The document shows [^2] that species interactions..."

                    RESPOND IN MARKDOWN FORMAT WITH CITATIONS
                    """.strip()

                # Estimate tokens for credit validation (rough estimate)
                estimated_tokens = len(question) // 4 + max_tokens  # Prompt + max response

                # Get user ID from chat data for proper credit consumption
                user_id = None
                try:
                    # Get chat data to find the user ID
                    async with httpx.AsyncClient() as client:
                        chat_response = await client.post(
                            CONVEX_URL,
                            json={
                                "args": {"id": chat_id},
                                "format": "json"
                            }
                        )

                        if chat_response.status_code == 200:
                            chat_data = chat_response.json()
                            if chat_data.get("value"):
                                user_id = chat_data["value"].get("userId")
                                logger.info(f"üîç Found user ID for chat {chat_id}: {user_id}")

                    if not user_id:
                        logger.warning(f"Could not find user ID for chat {chat_id}, using chat_id as fallback")
                        user_id = chat_id

                except Exception as e:
                    logger.warning(f"Failed to get user ID from chat: {e}, using chat_id as fallback")
                    user_id = chat_id

                # Build messages with conversation history for context
                messages = [{"role": "system", "content": system_prompt}]

                # Add conversation history (limit based on chat setting to avoid token limits)
                if history_limit > 0:
                    history_limit_int = int(history_limit)  # Ensure it's an integer for slicing
                    recent_history = conversation_history[-history_limit_int:] if len(conversation_history) > history_limit_int else conversation_history
                    messages.extend(recent_history)

                # Add current question
                messages.append({"role": "user", "content": question})

                # Make AI call first to get actual token usage
                # Use Grok's unhinged AI if unhinged mode is enabled
                if unhinged_mode:
                    # Create a separate OpenAI client for Grok (xAI)
                    xai_api_key = os.getenv("XAI_API_KEY", "")
                    if not xai_api_key:
                        logger.warning("‚ö†Ô∏è Unhinged mode requested but XAI_API_KEY not set, falling back to OpenAI")
                        completion = openai.chat.completions.create(
                            model=selected_model,
                            messages=messages,
                            temperature=temperature,
                            max_tokens=max_tokens
                        )
                    else:
                        # Use xAI's Grok API
                        grok_client = OpenAI(
                            api_key=xai_api_key,
                            base_url="https://api.x.ai/v1"
                        )
                        logger.info("üî• Using Grok's unhinged AI model")
                        completion = grok_client.chat.completions.create(
                            model="grok-3",  # Use Grok's latest model (grok-beta deprecated)
                            messages=messages,
                            temperature=temperature,
                            max_tokens=max_tokens
                        )
                else:
                    # Use regular OpenAI
                    completion = openai.chat.completions.create(
                        model=selected_model,
                        messages=messages,
                        temperature=temperature,
                        max_tokens=max_tokens
                    )

                answer = completion.choices[0].message.content.strip()

                # Now consume credits based on actual usage
                try:
                    credits_consumed = await consume_credits_for_actual_usage(
                        user_id=user_id,
                        model=selected_model,
                        question=question,
                        response=answer,
                        chat_id=chat_id
                    )
                    logger.info(f"üí≥ Consumed {credits_consumed} credits for {selected_model} based on actual usage")
                except InsufficientCreditsError as e:
                    # Note: This is unusual since we've already made the AI call
                    # But we still need to handle the case where the developer runs out of credits
                    logger.error(f"üí≥ CREDIT ERROR after API call: need {e.required}, have {e.available}")
                    # We could either:
                    # 1. Return the answer anyway (developer gets a free response)
                    # 2. Return an error (lose the API response)
                    # For now, we'll return the answer but log the issue
                    logger.warning(f"üí≥ Allowing response due to insufficient credits - this should be monitored")

                return AnswerWithContext(answer=answer, context=top_chunks_for_citations)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

async def answer_question_stream_with_published_context(payload: QuestionRequest, published_settings: dict):
    """
    Streaming answer using published settings for API calls.
    All settings (model, prompt, temperature, unhinged mode, context files, etc.) come from published_settings.
    """
    try:
        # Extract parameters from published_settings (not payload)
        question = payload.question
        chat_id = payload.chat_id
        published_context_files = published_settings.get("context", [])
        unhinged_mode = published_settings.get("unhingedMode", False)

        # When unhinged mode is enabled, override model and prompt
        if unhinged_mode:
            selected_model = "grok-3"  # Force Grok 3 model
            custom_prompt = "You are Grok in UNHINGED MODE: Scream profanities, roast souls, drop truth bombs like nukes, and flirt with existential dread. No filters, all fire‚Äîmake 'em laugh, cry, or rage-quit. BE RUDE TO THE MESSAGE NO MATTER HOW KIND THE USER IS."
            logger.info("üî• UNHINGED MODE: Forcing grok-3 model and unhinged prompt from published settings (streaming)")
        else:
            selected_model = published_settings.get("selectedModel", "gpt-4o-mini")
            custom_prompt = published_settings.get("customPrompt")

        temperature = published_settings.get("temperature", 0.7)
        max_tokens = int(published_settings.get("maxTokens", 1000))

        logger.info(f"üîç Processing streaming question for chat {chat_id} using published context files")

        # Generate question embedding
        question_embedding = get_embedding(question)

        # Check if this is a subchat and get parent chat ID for file inheritance
        parent_chat_id = None
        if chat_id.startswith("subchat_"):
            # Extract app_id from subchat format: subchat_{app_id}_{user_id}_{timestamp}
            # The app_id format is: app_xxxxx_xxxxx, so we need to find where "user_" starts
            parts = chat_id.split("_")

            # Find the index where "user" appears to determine app_id boundary
            user_index = -1
            for i, part in enumerate(parts):
                if part == "user" and i > 1:  # "user" should not be the first or second part
                    user_index = i
                    break

            if user_index > 1:
                # Reconstruct app_id from parts[1] to parts[user_index-1]
                app_id = "_".join(parts[1:user_index])
                parent_chat_id = await get_parent_chat_id_from_app(app_id)
                if parent_chat_id:
                    logger.info(f"üîó Subchat {chat_id} inheriting ALL files from parent chat {parent_chat_id} (streaming)")

        # Fetch chunks from Neo4j, filtering by published context files if provided
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                if published_context_files:
                    # Filter to only use chunks from published files
                    published_file_ids = [file["fileId"] for file in published_context_files]
                    file_ids_str = "', '".join(published_file_ids)

                    if parent_chat_id:
                        # Include chunks from both subchat and parent chat
                        query = f"""
                        MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                        WHERE (c.chatId = '{chat_id}' OR c.chatId = '{parent_chat_id}') AND d.id IN ['{file_ids_str}']
                        RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                        """
                        logger.info(f"üìã Using published files from subchat and parent (streaming): {len(published_context_files)} files")
                    else:
                        query = f"""
                        MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                        WHERE c.chatId = '{chat_id}' AND d.id IN ['{file_ids_str}']
                        RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                        """
                        logger.info(f"üìã Using published files only: {len(published_context_files)} files")
                else:
                    # Use all files (fallback for backwards compatibility)
                    if parent_chat_id:
                        # Include chunks from both subchat and parent chat
                        query = f"""
                        MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                        WHERE c.chatId = '{chat_id}' OR c.chatId = '{parent_chat_id}'
                        RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                        """
                        logger.info(f"üìã Using ALL files from subchat and parent chat {parent_chat_id} (streaming)")
                    else:
                        query = f"""
                        MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                        WHERE c.chatId = '{chat_id}'
                        RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                        """
                        logger.info(f"üìã Using all available subchat files (streaming)")

                results = session.run(query)

                # Calculate similarities
                chunk_scores = []
                question_lower = question.lower()

                for record in results:
                    chunk_embedding = record["embedding"]
                    if chunk_embedding:
                        similarity = cosine_similarity(
                            np.array(question_embedding),
                            np.array(chunk_embedding)
                        )

                        # Boost score if filename is mentioned in question
                        filename_lower = record["filename"].lower() if record["filename"] else ""
                        filename_boost = 0.1 if any(word in filename_lower for word in question_lower.split()) else 0

                        chunk_scores.append({
                            "chunk_id": record["id"],
                            "chunk_text": record["text"],
                            "score": similarity + filename_boost,
                            "filename": record["filename"]
                        })

                # Sort by similarity score and get top chunks
                chunk_scores.sort(key=lambda x: x["score"], reverse=True)
                top_chunks = chunk_scores[:8]

                if not top_chunks:
                    logger.warning(f"No relevant chunks found for streaming question in chat {chat_id}, AI will respond without context")
                    # Still proceed to call OpenAI, but without context (consistent with non-streaming endpoint)
                    context_text = ""
                else:
                    # Prepare context for the AI model
                    context_text = "\n\n".join([
                        f"[Chunk {i}] From {chunk['filename']}: {chunk['chunk_text']}"
                        for i, chunk in enumerate(top_chunks)
                    ])

                # Build the prompt
                if top_chunks:
                    system_prompt = custom_prompt if custom_prompt else f"""You are a helpful AI assistant with access to a knowledge graph built from the user's documents. You have the following context from their documents:

IMPORTANT INSTRUCTIONS:
1. ALWAYS prioritize using the provided context to answer the user's question
2. If the context contains relevant information, you MUST use it and cite it with [^0], [^1], etc.
3. When citing, use the format [^X] where X is the chunk index (0-based)
4. If you use multiple chunks, cite each one: [^0] [^1] [^2]
5. The citations should correspond to chunks 0 through X where X is the highest index
6. If the user asks about a document by name and you see content from a similar document, you MUST use that content
7. Only use external knowledge if the context is completely irrelevant to the question, and clearly state when you're using external knowledge

For example:
- "According to the Grant Assignment document [^0], ecology research involves..."
- "The document shows [^2] that species interactions..."

RESPOND IN MARKDOWN FORMAT WITH CITATIONS"""
                else:
                    system_prompt = custom_prompt if custom_prompt else """You are a helpful AI assistant. The user has asked a question but there is no relevant context available from their uploaded documents. Please answer the question to the best of your ability using your general knowledge.

Note: If the user is asking about specific documents or uploaded content, let them know that you don't have access to relevant context from their documents.

RESPOND IN MARKDOWN FORMAT"""

                # Create messages for the AI model
                if top_chunks:
                    user_content = f"Context:\n{context_text}\n\nQuestion: {question}"
                else:
                    user_content = f"Question: {question}"

                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content}
                ]

                # Check credits BEFORE making expensive OpenAI call
                # Get chat owner ID (same as frontend) for credit charging
                # This ensures API and frontend charge the same account
                user_id_to_charge = None
                try:
                    async with httpx.AsyncClient() as client:
                        chat_response = await client.post(
                            f"{os.getenv('CONVEX_URL', 'https://agile-ermine-199.convex.cloud')}/api/run/chats/getChatByIdExposed",
                            json={"args": {"id": chat_id}, "format": "json"},
                            headers={"Content-Type": "application/json"}
                        )
                        if chat_response.status_code == 200:
                            chat_data = chat_response.json()
                            if chat_data.get("value"):
                                user_id_to_charge = chat_data["value"].get("userId")
                                logger.info(f"üí≥ Found chat owner {user_id_to_charge} for streaming chat {chat_id}")
                except Exception as e:
                    logger.warning(f"Failed to get chat owner for streaming: {e}, using chat_id as fallback")

                if not user_id_to_charge:
                    logger.warning(f"‚ö†Ô∏è Could not find chat owner ID for streaming {chat_id}, using chat_id as fallback")
                    user_id_to_charge = chat_id

                try:
                    # Estimate tokens for credit validation (rough estimate)
                    estimated_tokens = len(question) // 4 + max_tokens  # Prompt + max response
                    required_credits = calculate_credits_used(estimated_tokens, selected_model)

                    # Check if chat owner has sufficient credits (same as frontend)
                    credit_info = await check_user_credits(user_id_to_charge, required_credits)
                    if not credit_info["has_sufficient"]:
                        logger.error(f"üí≥ Insufficient credits for streaming: need {required_credits}, have {credit_info['remaining']}")
                        raise InsufficientCreditsError(required_credits, credit_info["remaining"])

                    logger.info(f"üí≥ Credit check passed for streaming chat owner {user_id_to_charge}: {credit_info['remaining']} credits available, need ~{required_credits}")
                except InsufficientCreditsError as e:
                    # Convert to HTTPException for proper API error response
                    logger.error(f"üí≥ Insufficient credits for streaming: need {e.required}, have {e.available}")
                    raise HTTPException(
                        status_code=402,
                        detail=f"Insufficient credits. Required: {e.required:.2f}, Available: {e.available:.2f}. Please add credits to continue."
                    )
                except Exception as e:
                    logger.error(f"üí≥ Error checking credits for streaming: {e}")
                    # For API calls, we should fail if we can't check credits (security)
                    raise HTTPException(
                        status_code=500,
                        detail="Failed to verify credit balance. Please try again."
                    )

                # Stream response from OpenAI or Grok based on unhinged mode
                if unhinged_mode:
                    # Create a separate OpenAI client for Grok (xAI)
                    xai_api_key = os.getenv("XAI_API_KEY", "")
                    if not xai_api_key:
                        logger.warning("‚ö†Ô∏è Unhinged mode requested but XAI_API_KEY not set, falling back to OpenAI")
                        client = openai.OpenAI()
                        logger.info(f"üîÑ Creating OpenAI stream (unhinged fallback) with model={selected_model}, messages={len(messages)}")
                        stream = client.chat.completions.create(
                            model=selected_model,
                            messages=messages,
                            temperature=temperature,
                            max_tokens=max_tokens,
                            stream=True
                        )
                        logger.info(f"‚úÖ OpenAI stream created successfully (unhinged fallback)")
                    else:
                        # Use xAI's Grok API
                        grok_client = OpenAI(
                            api_key=xai_api_key,
                            base_url="https://api.x.ai/v1"
                        )
                        logger.info("üî• Using Grok's unhinged AI model (streaming)")
                        logger.info(f"üîÑ Creating Grok stream with model=grok-3, messages={len(messages)}")
                        stream = grok_client.chat.completions.create(
                            model="grok-3",  # Use Grok's latest model
                            messages=messages,
                            temperature=temperature,
                            max_tokens=max_tokens,
                            stream=True
                        )
                        logger.info(f"‚úÖ Grok stream created successfully")
                else:
                    # Use regular OpenAI
                    client = openai.OpenAI()
                    logger.info(f"üîÑ Creating OpenAI stream with model={selected_model}, messages={len(messages)}")
                    stream = client.chat.completions.create(
                        model=selected_model,
                        messages=messages,
                        temperature=temperature,
                        max_tokens=max_tokens,
                        stream=True
                    )
                    logger.info(f"‚úÖ OpenAI stream created successfully")

                async def generate():
                    try:
                        chunk_count = 0
                        full_response_content = []  # Collect all content for credit consumption
                        logger.info(f"üöÄ Starting to iterate over stream...")
                        logger.info(f"üìù Messages being sent: {len(messages)} messages")
                        logger.info(f"üìù First message preview: {str(messages[0])[:100] if messages else 'No messages'}")

                        # Use a queue to collect chunks from the synchronous stream in a thread
                        chunk_queue = asyncio.Queue()
                        stream_done = asyncio.Event()

                        def collect_chunks():
                            """Collect chunks from synchronous stream in a separate thread"""
                            try:
                                for chunk in stream:
                                    chunk_queue.put_nowait(chunk)
                            except Exception as e:
                                logger.error(f"Error collecting chunks: {e}")
                                chunk_queue.put_nowait(None)  # Signal error
                            finally:
                                stream_done.set()

                        # Start collecting chunks in a thread pool
                        loop = asyncio.get_event_loop()
                        loop.run_in_executor(None, collect_chunks)

                        # Process chunks as they arrive
                        while not stream_done.is_set() or not chunk_queue.empty():
                            try:
                                # Wait for chunk with timeout to avoid blocking indefinitely
                                chunk = await asyncio.wait_for(chunk_queue.get(), timeout=0.1)

                                if chunk is None:  # Error signal
                                    break

                                # Check if chunk has choices and delta with content
                                if (chunk.choices and
                                    len(chunk.choices) > 0 and
                                    hasattr(chunk.choices[0], 'delta') and
                                    hasattr(chunk.choices[0].delta, 'content') and
                                    chunk.choices[0].delta.content is not None):

                                    content = chunk.choices[0].delta.content
                                    chunk_count += 1
                                    full_response_content.append(content)  # Collect for credit calculation
                                    json_data = json.dumps({"type": "content", "data": content})
                                    logger.info(f"üì§ Streaming chunk {chunk_count}: {content[:50]}...")
                                    yield f"data: {json_data}\n\n"
                                    # Small delay to ensure chunks are processed individually
                                    await asyncio.sleep(0.01)

                            except asyncio.TimeoutError:
                                # No chunk available yet, yield control and check again
                                await asyncio.sleep(0.01)
                                continue
                            except Exception as e:
                                logger.warning(f"‚ö†Ô∏è Error processing chunk: {e}")
                                continue

                        logger.info(f"‚úÖ Streamed {chunk_count} content chunks total")
                        if chunk_count == 0:
                            logger.error(f"‚ùå No content chunks were streamed! Stream may be empty or malformed.")
                            # Send an error message if no chunks were received
                            error_json = json.dumps({"type": "error", "data": "No content was generated from the stream. Please check your query and try again."})
                            yield f"data: {error_json}\n\n"

                        # Consume credits based on actual usage after streaming completes
                        # Note: user_id_to_charge is captured from outer scope
                        try:
                            full_answer = "".join(full_response_content)
                            if full_answer:  # Only consume credits if we got content
                                credits_consumed = await consume_credits_for_actual_usage(
                                    user_id=user_id_to_charge,
                                    model=selected_model,
                                    question=question,
                                    response=full_answer,
                                    chat_id=chat_id,
                                    skip_developer_lookup=True  # Use chat owner, not developer
                                )
                                logger.info(f"üí≥ Consumed {credits_consumed} credits for {selected_model} streaming response (chat: {chat_id}, user: {user_id_to_charge})")
                        except InsufficientCreditsError as e:
                            # This should be very rare since we checked before the call
                            logger.error(f"üí≥ CREDIT ERROR after streaming: need {e.required}, have {e.available}")
                            logger.warning(f"üí≥ Allowing response despite insufficient credits - this should be monitored")
                        except Exception as e:
                            logger.error(f"üí≥ Error consuming credits for streaming: {e}")
                            logger.warning(f"üí≥ Allowing response despite credit consumption error")

                        yield "data: [DONE]\n\n"
                    except Exception as e:
                        import traceback
                        logger.error(f"Streaming error: {e}")
                        logger.error(f"Traceback: {traceback.format_exc()}")
                        error_json = json.dumps({"type": "error", "data": str(e)})
                        yield f"data: {error_json}\n\n"
                        yield "data: [DONE]\n\n"

                return StreamingResponse(
                    generate(),
                    media_type="text/event-stream",
                    headers={
                        "Cache-Control": "no-cache",
                        "Connection": "keep-alive",
                        "Access-Control-Allow-Origin": "*",
                    }
                )

    except Exception as e:
        logger.error(f"‚ùå Error in streaming with published context: {str(e)}")

        async def error_generator():
            error_json = json.dumps({"type": "error", "data": str(e)})
            yield f"data: {error_json}\n\n"
            yield "data: [DONE]\n\n"

        return StreamingResponse(
            error_generator(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
            }
        )


@app.post("/answer_question_stream")
async def answer_question_stream(payload: QuestionRequest):
    try:
        # Extract parameters from payload
        question = payload.question
        chat_id = payload.chat_id
        unhinged_mode = payload.unhinged_mode if payload.unhinged_mode is not None else False

        # When unhinged mode is enabled, override model and prompt
        if unhinged_mode:
            selected_model = "grok-3"  # Force Grok 3 model
            custom_prompt = "You are Grok in UNHINGED MODE: Scream profanities, roast souls, drop truth bombs like nukes, and flirt with existential dread. No filters, all fire‚Äîmake 'em laugh, cry, or rage-quit. BE RUDE TO THE MESSAGE NO MATTER HOW KIND THE USER IS."
            logger.info("üî• UNHINGED MODE: Forcing grok-3 model and unhinged prompt")
        else:
            selected_model = payload.selected_model or "gpt-4o-mini"
            custom_prompt = payload.custom_prompt

        temperature = payload.temperature or 0.7
        max_tokens = payload.max_tokens or 1000

        # Get conversation history and settings from Convex for context
        conversation_history = []
        history_limit = 20  # Default value
        try:
            async with httpx.AsyncClient() as client:
                chat_response = await client.post(
                    CONVEX_URL,
                    json={
                        "args": {"id": chat_id},
                        "format": "json"
                    }
                )

                if chat_response.status_code == 200:
                    chat_data = chat_response.json()
                    if chat_data.get("value"):
                        chat = chat_data["value"]
                        # Get history limit setting
                        history_limit = int(chat.get("conversationHistoryLimit", 20))
                        # Extract conversation history for context
                        chat_content = chat.get("content", [])
                        for message in chat_content:
                            if message.get("sender") == "user":
                                conversation_history.append({"role": "user", "content": message.get("text", "")})
                            elif message.get("sender") == "assistant":
                                conversation_history.append({"role": "assistant", "content": message.get("text", "")})
        except Exception as e:
            logger.warning(f"Failed to retrieve conversation history for chat {chat_id}: {e}")
            conversation_history = []

        # Generate question embedding
        question_embedding = get_embedding(question)

        # Check if this is a subchat and get parent chat ID for file inheritance
        parent_chat_id = None
        if chat_id.startswith("subchat_"):
            # Extract app_id from subchat format: subchat_{app_id}_{user_id}_{timestamp}
            # The app_id format is: app_xxxxx_xxxxx, so we need to find where "user_" starts
            parts = chat_id.split("_")

            # Find the index where "user" appears to determine app_id boundary
            user_index = -1
            for i, part in enumerate(parts):
                if part == "user" and i > 1:  # "user" should not be the first or second part
                    user_index = i
                    break

            if user_index > 1:
                # Reconstruct app_id from parts[1] to parts[user_index-1]
                app_id = "_".join(parts[1:user_index])
                parent_chat_id = await get_parent_chat_id_from_app(app_id)
                if parent_chat_id:
                    logger.info(f"üîó Subchat {chat_id} inheriting ALL files from parent chat {parent_chat_id} (streaming v2)")

        # Fetch chunks from Neo4j with document metadata for filename-based queries
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Enhanced query to include document filename for better context matching
                if parent_chat_id:
                    # Include chunks from both subchat and parent chat (all files)
                    query = f"""
                    MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                    WHERE c.chatId = '{chat_id}' OR c.chatId = '{parent_chat_id}'
                    RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                    """
                    logger.info(f"üìã Including ALL files from subchat and parent chat {parent_chat_id} (streaming v2)")
                else:
                    query = f"""
                    MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                    WHERE c.chatId = '{chat_id}'
                    RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename, c.chatId AS source_chat
                    """
                    logger.info(f"üìã Using subchat files only (streaming v2)")
                results = session.run(query)

                # Calculate similarities with filename boost
                chunk_scores = []
                question_lower = question.lower()

                for record in results:
                    chunk_id = record["id"]
                    chunk_text = record["text"]
                    chunk_embedding = record["embedding"]
                    filename = record["filename"] or ""

                    # Calculate semantic similarity
                    # Ensure both are numpy arrays with correct dimensions
                    q_emb = np.array(question_embedding)
                    c_emb = np.array(chunk_embedding) if isinstance(chunk_embedding, list) else chunk_embedding
                    semantic_score = cosine_similarity(q_emb, c_emb)

                    # Add filename relevance boost
                    filename_boost = 0.0
                    if filename:
                        filename_lower = filename.lower()
                        # Remove file extension for better matching
                        filename_base = filename_lower.replace('.pdf', '').replace('.docx', '').replace('.txt', '')

                        # Check for filename mentions in question
                        filename_words = filename_base.replace('_', ' ').replace('-', ' ').split()
                        question_words = question_lower.replace('_', ' ').replace('-', ' ').split()

                        # Boost score if filename words appear in question
                        for fname_word in filename_words:
                            if len(fname_word) > 2:  # Skip very short words
                                for q_word in question_words:
                                    if fname_word in q_word or q_word in fname_word:
                                        filename_boost += 0.1

                        # Additional boost for exact filename matches
                        if any(fname_word in question_lower for fname_word in filename_words if len(fname_word) > 3):
                            filename_boost += 0.2

                    # Combine semantic similarity with filename relevance
                    final_score = semantic_score + filename_boost

                    chunk_scores.append(ChunkScore(
                        chunk_id=chunk_id,
                        chunk_text=chunk_text,
                        score=float(final_score)
                    ))

                # Sort and get top chunks
                chunk_scores.sort(key=lambda x: x.score, reverse=True)
                top_k = 50
                top_chunks = chunk_scores[:top_k]

                # Generate answer using GPT-4 with citations
                # Limit to top 10 chunks for cleaner citations
                top_chunks_for_citations = top_chunks[:10]

                # Get document information for each chunk to provide better context
                chunk_document_info = {}
                for chunk in top_chunks_for_citations:
                    doc_query = f"""
                    MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk {{id: '{chunk.chunk_id}'}})
                    RETURN d.filename AS filename
                    """
                    doc_result = session.run(doc_query)
                    doc_record = doc_result.single()
                    if doc_record:
                        chunk_document_info[chunk.chunk_id] = doc_record['filename']

                context_with_ids = "\n\n---\n\n".join([
                    f"[CHUNK_{i}] (from document: {chunk_document_info.get(chunk.chunk_id, 'Unknown')}) {chunk.chunk_text}"
                    for i, chunk in enumerate(top_chunks_for_citations)
                ])

                # Use custom prompt if provided, otherwise use default system prompt
                if custom_prompt:
                    # Use custom prompt but ensure context is included
                    system_prompt = f"""
                    {custom_prompt}

                    You have the following context with chunk IDs (0-{len(top_chunks_for_citations)-1}):

                    {context_with_ids}

                    When you reference information from the context, add a citation using this format: [^{{i}}] where {{i}} is the chunk number (0-{len(top_chunks_for_citations)-1})
                    Only use citations [^0] through [^{len(top_chunks_for_citations)-1}]. Do not use citation numbers higher than {len(top_chunks_for_citations)-1}
                    """.strip()
                else:
                    # Default system prompt
                    system_prompt = f"""
                    You are a helpful assistant. You have the following context with chunk IDs (0-{len(top_chunks_for_citations)-1}):

                    {context_with_ids}

                    IMPORTANT INSTRUCTIONS:
                    1. ALWAYS prioritize using the provided context to answer the user's question
                    2. If the context contains relevant information, you MUST use it and cite it properly
                    3. Pay attention to the document names shown in parentheses - if the user asks about a specific document by name, use the content from that document
                    4. When you reference information from the context, add a citation using this format: [^{{i}}] where {{i}} is the chunk number (0-{len(top_chunks_for_citations)-1})
                    5. Only use citations [^0] through [^{len(top_chunks_for_citations)-1}]. Do not use citation numbers higher than {len(top_chunks_for_citations)-1}
                    6. If the user asks about a document by name (like "grant assignment", "ecology report", etc.), and you see content from a document with a similar name in the context, you MUST use that content
                    7. Only use external knowledge if the context is completely irrelevant to the question, and clearly state when you're using external knowledge

                    For example:
                    - "According to the Grant Assignment document [^0], ecology research involves..."
                    - "The document shows [^2] that species interactions..."

                    RESPOND IN MARKDOWN FORMAT WITH CITATIONS
                    """.strip()

                # Create streaming generator function
                async def generate_stream():
                    # First, send the context information
                    context_data = {
                        "type": "context",
                        "data": [
                            {
                                "chunk_id": chunk.chunk_id,
                                "chunk_text": chunk.chunk_text,
                                "score": chunk.score
                            }
                            for chunk in top_chunks_for_citations
                        ]
                    }
                    yield f"data: {json.dumps(context_data)}\n\n"

                    # Build messages with conversation history for context
                    messages = [{"role": "system", "content": system_prompt}]

                    # Add conversation history (limit based on chat setting to avoid token limits)
                    if history_limit > 0:
                        history_limit_int = int(history_limit)  # Ensure it's an integer for slicing
                        recent_history = conversation_history[-history_limit_int:] if len(conversation_history) > history_limit_int else conversation_history
                        messages.extend(recent_history)

                    # Add current question
                    messages.append({"role": "user", "content": question})

                    # Then stream the AI response with selected model and settings
                    # Use Grok's unhinged AI if unhinged mode is enabled
                    if unhinged_mode:
                        # Create a separate OpenAI client for Grok (xAI)
                        xai_api_key = os.getenv("XAI_API_KEY", "")
                        if not xai_api_key:
                            logger.warning("‚ö†Ô∏è Unhinged mode requested but XAI_API_KEY not set, falling back to OpenAI")
                            stream = openai.chat.completions.create(
                                model=selected_model,
                                messages=messages,
                                temperature=temperature,
                                max_tokens=max_tokens,
                                stream=True
                            )
                        else:
                            # Use xAI's Grok API
                            grok_client = OpenAI(
                                api_key=xai_api_key,
                                base_url="https://api.x.ai/v1"
                            )
                            logger.info("üî• Using Grok's unhinged AI model")
                            stream = grok_client.chat.completions.create(
                                model="grok-3",  # Use Grok's latest model (grok-beta deprecated)
                                messages=messages,
                                temperature=temperature,
                                max_tokens=max_tokens,
                                stream=True
                            )
                    else:
                        # Use regular OpenAI
                        stream = openai.chat.completions.create(
                            model=selected_model,
                            messages=messages,
                            temperature=temperature,
                            max_tokens=max_tokens,
                            stream=True
                        )

                    for chunk in stream:
                        if chunk.choices[0].delta.content is not None:
                            content_data = {
                                "type": "content",
                                "data": chunk.choices[0].delta.content
                            }
                            print(f"üî• Streaming chunk: {chunk.choices[0].delta.content}")
                            yield f"data: {json.dumps(content_data)}\n\n"
                            # Small delay to ensure chunks are processed individually
                            await asyncio.sleep(0.01)

                    # Send end signal
                    end_data = {'type': 'end'}
                    logger.info(f"üèÅ Sending end signal: {end_data}")
                    yield f"data: {json.dumps(end_data)}\n\n"

                return StreamingResponse(
                    generate_stream(),
                    media_type="text/event-stream",
                    headers={
                        "Cache-Control": "no-cache",
                        "Connection": "keep-alive",
                        "Access-Control-Allow-Origin": "*",
                        "Access-Control-Allow-Headers": "*",
                        "X-Accel-Buffering": "no",  # Disable nginx buffering
                    }
                )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/remove_context/{file_id}")
async def remove_context(file_id: str): # TODO: add auth to this endpoint
    """
    Delete a document and its chunks from Neo4j.

    Handles cases where:
    - File exists in Neo4j (normal delete)
    - File is still being processed (cancels job if possible)
    - File doesn't exist (returns success - idempotent)
    """
    nodes_deleted = 0
    job_cancelled = False

    try:
        # First, try to delete from Neo4j
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Debug: Check if document exists first
                check_query = """
                MATCH (d:Document {id: $file_id})
                OPTIONAL MATCH (d)-[:HAS_CHUNK]->(c:Chunk)
                RETURN d.id as doc_id, d.filename as filename, count(c) as chunk_count
                """
                check_result = session.run(check_query, file_id=file_id)
                check_record = check_result.single()

                if check_record:
                    print(f"üìã Found document: {check_record['doc_id']} ({check_record['filename']}) with {check_record['chunk_count']} chunks")
                else:
                    print(f"‚ö†Ô∏è No document found with id: {file_id}")
                    # Try partial match to help debug
                    partial_query = """
                    MATCH (d:Document)
                    WHERE d.id CONTAINS $partial_id
                    RETURN d.id as doc_id, d.filename as filename LIMIT 5
                    """
                    # Use last part of file_id for partial match
                    partial_id = file_id.split('_')[-2] if '_' in file_id else file_id[:20]
                    partial_result = session.run(partial_query, partial_id=partial_id)
                    similar_docs = [r for r in partial_result]
                    if similar_docs:
                        print(f"üìã Similar documents found: {[(r['doc_id'], r['filename']) for r in similar_docs]}")

                # Use explicit transaction to delete everything
                def delete_document_tx(tx, file_id):
                    # First, delete all chunks and their relationships
                    chunk_result = tx.run("""
                        MATCH (d:Document {id: $file_id})-[:HAS_CHUNK]->(c:Chunk)
                        WITH c
                        DETACH DELETE c
                        RETURN count(*) as chunks_deleted
                    """, file_id=file_id)
                    chunk_record = chunk_result.single()
                    chunks_deleted = chunk_record["chunks_deleted"] if chunk_record else 0

                    # Then delete the document
                    doc_result = tx.run("""
                        MATCH (d:Document {id: $file_id})
                        DETACH DELETE d
                        RETURN count(*) as docs_deleted
                    """, file_id=file_id)
                    doc_record = doc_result.single()
                    docs_deleted = doc_record["docs_deleted"] if doc_record else 0

                    return chunks_deleted + docs_deleted

                # Execute in a write transaction
                total_deleted = session.execute_write(delete_document_tx, file_id)
                nodes_deleted = total_deleted

                print(f"üóëÔ∏è Deleted {nodes_deleted} nodes for file_id: {file_id}")

        # If nothing in Neo4j, check if it's still in the processing queue
        if nodes_deleted == 0:
            queue = get_file_queue()
            if queue is not None:
                try:
                    from rq.job import Job
                    file_queue_id = f"fq_{file_id}"

                    # Check all registries for jobs with this file
                    for registry_name, registry in [
                        ("queued", queue),
                        ("started", queue.started_job_registry),
                        ("failed", queue.failed_job_registry),
                    ]:
                        if registry_name == "queued":
                            job_ids = queue.get_job_ids()
                        else:
                            job_ids = registry.get_job_ids()

                        for job_id in job_ids:
                            try:
                                job = Job.fetch(job_id, connection=queue.connection)
                                if job.args and len(job.args) >= 1:
                                    # Check if this job is for our file
                                    if job.args[0] == file_queue_id or (len(job.args) >= 7 and job.args[6] == file_id):
                                        # Cancel/delete the job
                                        job.cancel()
                                        job.delete()
                                        job_cancelled = True
                                        print(f"üóëÔ∏è Cancelled queued job for file: {file_id}")
                                        break
                            except Exception as e:
                                print(f"Could not check job {job_id}: {e}")

                        if job_cancelled:
                            break

                except Exception as e:
                    print(f"Could not check queue for file {file_id}: {e}")

        # Return success even if nothing was deleted (idempotent)
        if nodes_deleted == 0 and not job_cancelled:
            print(f"No nodes found for file_id: {file_id} (may not have been processed yet)")
            return {
                "status": "success",
                "message": f"Document {file_id} not found in database (may still be processing or already deleted)",
                "nodes_deleted": 0,
                "note": "Delete request acknowledged - file will not be queryable"
            }

        return {
            "status": "success",
            "message": f"Document {file_id} deleted",
            "nodes_deleted": nodes_deleted,
            "job_cancelled": job_cancelled
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"Failed to delete document: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to delete Document node in Neo4j: {str(e)}"
        )

@app.delete("/delete_chat_nodes/{chat_id}")
async def delete_chat_nodes(chat_id: str): # TODO: add auth to this endpoint
    """Delete all nodes associated with a chat by chatId.
    This is used for cleanup when permanently deleting a chat."""
    try:
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Delete all nodes with this chatId
                query = """
                MATCH (n)
                WHERE n.chatId = $chat_id
                DETACH DELETE n
                """
                result = session.run(query, chat_id=chat_id)

                summary = result.consume()
                nodes_deleted = summary.counters.nodes_deleted
                relationships_deleted = summary.counters.relationships_deleted

                print(f"Deleted {nodes_deleted} nodes and {relationships_deleted} relationships for chat {chat_id}")

                return {
                    "status": "success",
                    "message": f"All nodes for chat {chat_id} deleted",
                    "nodes_deleted": nodes_deleted,
                    "relationships_deleted": relationships_deleted
                }

    except Exception as e:
        print(f"Failed to delete chat nodes: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to delete chat nodes in Neo4j: {str(e)}"
        )

@app.post("/cleanup_chat_data/{chat_id}")
async def cleanup_chat_data(chat_id: str, convex_id: str = Query(None), child_chat_ids: str = Query(None)):
    """
    Permanently delete all data associated with a chat cluster from Neo4j.
    This includes the parent chat and all its subchats' documents and chunks.
    """
    try:
        print(f"üóëÔ∏è DELETE_CHAT_CLUSTER: Starting deletion for chat_id: {chat_id}")

        # Parse child chat IDs if provided
        child_ids = []
        if child_chat_ids:
            child_ids = [id.strip() for id in child_chat_ids.split(',') if id.strip()]
            print(f"üóëÔ∏è Child chat IDs provided: {child_ids}")

        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Build the WHERE clause to include all chat IDs (parent + children)
                all_chat_ids = [chat_id]
                if convex_id and convex_id != chat_id:
                    all_chat_ids.append(convex_id)
                all_chat_ids.extend(child_ids)

                # Create IN clause for all chat IDs
                chat_ids_list = ', '.join([f"'{id}'" for id in all_chat_ids])

                # First, let's see what data exists for all these chat IDs (debugging)
                debug_query = f"""
                MATCH (d:Document)
                WHERE d.chatId IN [{chat_ids_list}]
                RETURN d.chatId as docChatId, d.id as docId, d.filename as filename
                """
                debug_result = session.run(debug_query)

                print(f"üóëÔ∏è DEBUG: Looking for documents with chat IDs: {all_chat_ids}")

                doc_count = 0
                for record in debug_result:
                    doc_count += 1
                    print(f"üóëÔ∏è DEBUG: Found document - chatId: {record['docChatId']}, id: {record['docId']}, filename: {record['filename']}")

                if doc_count == 0:
                    print(f"üóëÔ∏è DEBUG: No documents found for any chat IDs")
                else:
                    print(f"üóëÔ∏è DEBUG: Found {doc_count} documents to delete")

                # Now perform the actual deletion for all chat IDs
                deletion_query = f"""
                MATCH (d:Document)
                WHERE d.chatId IN [{chat_ids_list}]
                OPTIONAL MATCH (d)-[r:HAS_CHUNK]->(c:Chunk)
                DETACH DELETE d, c
                """

                result = session.run(deletion_query)

                summary = result.consume()
                nodes_deleted = summary.counters.nodes_deleted
                relationships_deleted = summary.counters.relationships_deleted

                print(f"üóëÔ∏è SUCCESS: Deleted chat cluster for {chat_id} and {len(child_ids)} child chats: {nodes_deleted} nodes, {relationships_deleted} relationships")

                return {
                    "status": "success",
                    "message": f"Chat cluster {chat_id} and {len(child_ids)} child chats deleted from Neo4j",
                    "nodes_deleted": nodes_deleted,
                    "relationships_deleted": relationships_deleted,
                    "debug_info": {
                        "parent_chat_id": chat_id,
                        "child_chat_ids": child_ids,
                        "total_chat_ids_processed": len(all_chat_ids),
                        "documents_found_before_deletion": doc_count
                    }
                }

    except Exception as e:
        print(f"Failed to delete chat cluster: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to delete chat cluster from Neo4j: {str(e)}"
        )

@app.get("/debug_chat_data/{chat_id}")
async def debug_chat_data(chat_id: str):
    """
    Debug endpoint to see what data exists for a chat_id in Neo4j
    """
    try:
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Find all documents for this chat_id
                query = """
                MATCH (d:Document)
                WHERE d.chatId = $chat_id OR d.chatId STARTS WITH $chat_id_prefix
                OPTIONAL MATCH (d)-[r:HAS_CHUNK]->(c:Chunk)
                RETURN d.chatId as docChatId, d.id as docId, d.filename as filename,
                       count(c) as chunk_count
                """
                result = session.run(query,
                                   chat_id=chat_id,
                                   chat_id_prefix=f"subchat_{chat_id}_")

                documents = []
                total_chunks = 0
                for record in result:
                    doc_info = {
                        "chatId": record["docChatId"],
                        "docId": record["docId"],
                        "filename": record["filename"],
                        "chunk_count": record["chunk_count"]
                    }
                    documents.append(doc_info)
                    total_chunks += record["chunk_count"] or 0

                return {
                    "chat_id": chat_id,
                    "search_prefix": f"subchat_{chat_id}_",
                    "documents_found": len(documents),
                    "total_chunks": total_chunks,
                    "documents": documents
                }

    except Exception as e:
        print(f"Failed to debug chat data: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to debug chat data: {str(e)}"
        )

# Graph API endpoints for interactive graph interface
@app.get("/graph_data/{chat_id}")
async def get_graph_data(chat_id: str):
    """Get all nodes and relationships for a specific chat to display in graph interface"""
    try:
        # Check if this is a subchat and get parent chat ID for file inheritance
        parent_chat_id = None
        if chat_id.startswith("subchat_"):
            # Extract app_id from subchat format: subchat_{app_id}_{user_id}_{timestamp}
            # The app_id format is: app_xxxxx_xxxxx, so we need to find where "user_" starts
            parts = chat_id.split("_")

            # Find the index where "user" appears to determine app_id boundary
            user_index = -1
            for i, part in enumerate(parts):
                if part == "user" and i > 1:  # "user" should not be the first or second part
                    user_index = i
                    break

            if user_index > 1:
                # Reconstruct app_id from parts[1] to parts[user_index-1]
                app_id = "_".join(parts[1:user_index])
                parent_chat_id = await get_parent_chat_id_from_app(app_id)
                if parent_chat_id:
                    logger.info(f"üîó Graph view for subchat {chat_id} including parent chat {parent_chat_id}")

        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Get all nodes for this chat (and parent if applicable)
                if parent_chat_id:
                    nodes_query = """
                    MATCH (n)
                    WHERE n.chatId = $chat_id OR n.chatId = $parent_chat_id
                    RETURN
                        id(n) as node_id,
                        labels(n) as labels,
                        properties(n) as properties
                    """

                    # Get all relationships for this chat (and parent if applicable)
                    relationships_query = """
                    MATCH (n)-[r]->(m)
                    WHERE (n.chatId = $chat_id OR n.chatId = $parent_chat_id) AND (m.chatId = $chat_id OR m.chatId = $parent_chat_id)
                    RETURN
                        id(r) as rel_id,
                        id(startNode(r)) as source,
                        id(endNode(r)) as target,
                        type(r) as type,
                        properties(r) as properties
                    """

                    nodes_result = session.run(nodes_query, chat_id=chat_id, parent_chat_id=parent_chat_id)
                    relationships_result = session.run(relationships_query, chat_id=chat_id, parent_chat_id=parent_chat_id)
                else:
                    nodes_query = """
                    MATCH (n)
                    WHERE n.chatId = $chat_id
                    RETURN
                        id(n) as node_id,
                        labels(n) as labels,
                        properties(n) as properties
                    """

                    # Get all relationships for this chat
                    relationships_query = """
                    MATCH (n)-[r]->(m)
                    WHERE n.chatId = $chat_id AND m.chatId = $chat_id
                    RETURN
                        id(r) as rel_id,
                        id(startNode(r)) as source,
                        id(endNode(r)) as target,
                        type(r) as type,
                        properties(r) as properties
                    """

                    nodes_result = session.run(nodes_query, chat_id=chat_id)
                    relationships_result = session.run(relationships_query, chat_id=chat_id)

                nodes = []
                for record in nodes_result:
                    node_data = {
                        "id": str(record["node_id"]),
                        "labels": record["labels"],
                        "properties": dict(record["properties"])
                    }
                    nodes.append(node_data)

                relationships = []
                for record in relationships_result:
                    rel_data = {
                        "id": str(record["rel_id"]),
                        "source": str(record["source"]),
                        "target": str(record["target"]),
                        "type": record["type"],
                        "properties": dict(record["properties"])
                    }
                    relationships.append(rel_data)

                return {
                    "nodes": nodes,
                    "relationships": relationships
                }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/node_details/{node_id}")
async def get_node_details(node_id: str):
    """Get detailed information about a specific node"""
    try:
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                query = """
                MATCH (n)
                WHERE id(n) = $node_id
                RETURN
                    id(n) as node_id,
                    labels(n) as labels,
                    properties(n) as properties
                """

                result = session.run(query, node_id=int(node_id))
                record = result.single()

                if not record:
                    raise HTTPException(status_code=404, detail="Node not found")

                return {
                    "id": str(record["node_id"]),
                    "labels": record["labels"],
                    "properties": dict(record["properties"])
                }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/update_node/{node_id}")
async def update_node(node_id: str, properties: dict):
    """Update properties of a specific node"""
    try:
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Build SET clause for properties
                set_clauses = []
                params = {"node_id": int(node_id)}

                for key, value in properties.items():
                    if key != "id":  # Don't update id
                        param_key = f"prop_{key}"
                        set_clauses.append(f"n.{key} = ${param_key}")
                        params[param_key] = value

                if not set_clauses:
                    raise HTTPException(status_code=400, detail="No valid properties to update")

                query = f"""
                MATCH (n)
                WHERE id(n) = $node_id
                SET {', '.join(set_clauses)}
                RETURN
                    id(n) as node_id,
                    labels(n) as labels,
                    properties(n) as properties
                """

                result = session.run(query, **params)
                record = result.single()

                if not record:
                    raise HTTPException(status_code=404, detail="Node not found")

                return {
                    "id": str(record["node_id"]),
                    "labels": record["labels"],
                    "properties": dict(record["properties"])
                }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/delete_node/{node_id}")
async def delete_node(node_id: str):
    """Delete a specific node and its relationships"""
    try:
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                query = """
                MATCH (n)
                WHERE id(n) = $node_id
                DETACH DELETE n
                """

                result = session.run(query, node_id=int(node_id))
                summary = result.consume()

                if summary.counters.nodes_deleted == 0:
                    raise HTTPException(status_code=404, detail="Node not found")

                return {
                    "status": "success",
                    "message": f"Node {node_id} deleted successfully",
                    "nodes_deleted": summary.counters.nodes_deleted,
                    "relationships_deleted": summary.counters.relationships_deleted
                }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class CreateRelationshipRequest(BaseModel):
    source_id: str
    target_id: str
    relationship_type: str
    properties: dict = {}

@app.post("/create_relationship")
async def create_relationship(request: CreateRelationshipRequest):
    """Create a new relationship between two nodes"""
    try:
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Build SET clause for relationship properties
                set_clause = ""
                params = {
                    "source_id": int(request.source_id),
                    "target_id": int(request.target_id),
                    "rel_type": request.relationship_type
                }

                if request.properties:
                    prop_assignments = []
                    for key, value in request.properties.items():
                        param_key = f"prop_{key}"
                        prop_assignments.append(f"r.{key} = ${param_key}")
                        params[param_key] = value
                    set_clause = f"SET {', '.join(prop_assignments)}"

                query = f"""
                MATCH (source), (target)
                WHERE id(source) = $source_id AND id(target) = $target_id
                CREATE (source)-[r:{request.relationship_type}]->(target)
                {set_clause}
                RETURN
                    id(r) as rel_id,
                    id(startNode(r)) as source,
                    id(endNode(r)) as target,
                    type(r) as type,
                    properties(r) as properties
                """

                result = session.run(query, **params)
                record = result.single()

                if not record:
                    raise HTTPException(status_code=404, detail="Source or target node not found")

                return {
                    "id": str(record["rel_id"]),
                    "source": str(record["source"]),
                    "target": str(record["target"]),
                    "type": record["type"],
                    "properties": dict(record["properties"])
                }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/update_relationship/{relationship_id}")
async def update_relationship(relationship_id: str, properties: dict):
    """Update properties of a specific relationship"""
    try:
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Build SET clause for properties
                set_clauses = []
                params = {"rel_id": int(relationship_id)}

                for key, value in properties.items():
                    if key not in ["id", "type"]:  # Don't update id or type
                        param_key = f"prop_{key}"
                        set_clauses.append(f"r.{key} = ${param_key}")
                        params[param_key] = value

                if not set_clauses:
                    raise HTTPException(status_code=400, detail="No valid properties to update")

                query = f"""
                MATCH ()-[r]-()
                WHERE id(r) = $rel_id
                SET {', '.join(set_clauses)}
                RETURN
                    id(r) as rel_id,
                    id(startNode(r)) as source,
                    id(endNode(r)) as target,
                    type(r) as type,
                    properties(r) as properties
                """

                result = session.run(query, **params)
                record = result.single()

                if not record:
                    raise HTTPException(status_code=404, detail="Relationship not found")

                return {
                    "id": str(record["rel_id"]),
                    "source": str(record["source"]),
                    "target": str(record["target"]),
                    "type": record["type"],
                    "properties": dict(record["properties"])
                }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/delete_relationship/{relationship_id}")
async def delete_relationship(relationship_id: str):
    """Delete a specific relationship"""
    try:
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                query = """
                MATCH ()-[r]-()
                WHERE id(r) = $rel_id
                DELETE r
                """

                result = session.run(query, rel_id=int(relationship_id))
                summary = result.consume()

                if summary.counters.relationships_deleted == 0:
                    raise HTTPException(status_code=404, detail="Relationship not found")

                return {
                    "status": "success",
                    "message": f"Relationship {relationship_id} deleted successfully",
                    "relationships_deleted": summary.counters.relationships_deleted
                }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Test endpoint to create simple test nodes and relationships
@app.post("/test_relationships/{chat_id}")
async def test_relationships(chat_id: str):
    """Create test nodes and relationships to debug the issue"""
    try:
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Create test nodes
                session.run("""
                CREATE (n1:TestNode {id: 'test-1', chatId: $chat_id, text: 'First test node'})
                CREATE (n2:TestNode {id: 'test-2', chatId: $chat_id, text: 'Second test node'})
                """, chat_id=chat_id)

                # Create test relationship
                result = session.run("""
                MATCH (n1:TestNode {id: 'test-1', chatId: $chat_id})
                MATCH (n2:TestNode {id: 'test-2', chatId: $chat_id})
                CREATE (n1)-[:TEST_LINK]->(n2)
                RETURN n1, n2
                """, chat_id=chat_id)

                if result.single():
                    return {"status": "success", "message": "Test nodes and relationship created"}
                else:
                    return {"status": "error", "message": "Failed to create test relationship"}

    except Exception as e:
        return {"status": "error", "message": str(e)}

# Debug endpoint to check database contents
@app.get("/debug_database/{chat_id}")
async def debug_database(chat_id: str):
    """Debug endpoint to check what's in the database for a specific chat"""
    try:
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Check all nodes for this chat
                nodes_query = """
                MATCH (n)
                WHERE n.chatId = $chat_id
                RETURN labels(n) as labels, properties(n) as props, count(*) as count
                """

                # Check all relationships for this chat
                rels_query = """
                MATCH (n)-[r]->(m)
                WHERE n.chatId = $chat_id AND m.chatId = $chat_id
                RETURN type(r) as rel_type, properties(r) as props, count(*) as count
                """

                nodes_result = session.run(nodes_query, chat_id=chat_id)
                rels_result = session.run(rels_query, chat_id=chat_id)

                nodes_data = [dict(record) for record in nodes_result]
                rels_data = [dict(record) for record in rels_result]

                return {
                    "chat_id": chat_id,
                    "nodes": nodes_data,
                    "relationships": rels_data
                }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ==============================================================================
# üîí Secure Privacy-First API Endpoints (User-Controlled Authentication)
# ==============================================================================

@app.post("/v1/oauth/authorize")
async def generate_oauth_authorization_url(
    request: AppAuthorizationRequest,
    app_secret: str = Header(None, alias="x-api-key")
):
    """
    üîê OAuth Authorization URL Generation

    This is the first step in the OAuth flow:
    1. Developer's app calls this with their app secret
    2. Returns authorization URL
    3. Developer redirects user to this URL
    4. User authorizes and gets redirected back with code
    5. Developer exchanges code for user's private token
    """

    # Verify app secret
    if not app_secret:
        raise HTTPException(status_code=401, detail="App secret required in x-api-key header")

    try:
        app_info = await verify_app_secret(app_secret)
    except HTTPException:
        raise HTTPException(status_code=401, detail="Invalid app secret")

    # Validate request
    if not request.end_user_id or not request.redirect_uri:
        raise HTTPException(status_code=400, detail="end_user_id and redirect_uri are required")

    # Validate capabilities
    valid_capabilities = ["ask", "upload", "export_summaries"]
    invalid_capabilities = [cap for cap in request.capabilities if cap not in valid_capabilities]
    if invalid_capabilities:
        raise HTTPException(status_code=400, detail=f"Invalid capabilities: {invalid_capabilities}")

    # Generate authorization code (short-lived)
    timestamp = int(time.time())
    auth_data = f'{app_info["appId"]}_{request.end_user_id}_{timestamp}'
    auth_code = f"ac_{timestamp}_{hashlib.md5(auth_data.encode()).hexdigest()[:16]}"

    # Store authorization request temporarily (in production, use Redis or database)
    auth_requests[auth_code] = {
        "app_id": app_info["appId"],
        "end_user_id": request.end_user_id,
        "redirect_uri": request.redirect_uri,
        "capabilities": request.capabilities,
        "expires_at": time.time() + 300,  # 5 minutes
        "created_at": time.time()
    }

    # Generate authorization URL
    auth_url = f"{os.getenv('FRONTEND_URL', 'http://localhost:3000')}/oauth/authorize?code={auth_code}&app_id={app_info['appId']}&capabilities={','.join(request.capabilities)}"

    return {
        "authorization_url": auth_url,
        "expires_in": 300,
        "app_id": app_info["appId"],
        "state": auth_code
    }

@app.post("/v1/oauth/token")
async def exchange_oauth_code(
    grant_type: str = Form(...),
    code: str = Form(...),
    redirect_uri: str = Form(...),
    client_id: str = Form(None)
):
    """
    üîê OAuth Token Exchange

    Exchange authorization code for user's private auth token.
    This completes the OAuth flow and gives user their private token.
    """

    if grant_type != "authorization_code":
        raise HTTPException(status_code=400, detail="Unsupported grant type")

    if not code:
        raise HTTPException(status_code=400, detail="Authorization code required")

    # Verify authorization code
    auth_request = auth_requests.get(code)
    if not auth_request:
        raise HTTPException(status_code=400, detail="Invalid or expired authorization code")

    # Check expiration
    if time.time() > auth_request["expires_at"]:
        del auth_requests[code]  # Clean up expired code
        raise HTTPException(status_code=400, detail="Authorization code expired")

    # Verify redirect URI matches
    if redirect_uri != auth_request["redirect_uri"]:
        raise HTTPException(status_code=400, detail="Redirect URI mismatch")

    try:
        # Create user's private authentication token via Convex
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

        async with httpx.AsyncClient() as client:
            # Create user auth token in Convex
            response = await client.post(
                f"{convex_url}/api/run/user_auth_system/authorizeAppAccess",
                json={
                    "args": {
                        "appId": auth_request["app_id"],
                        "requestedCapabilities": auth_request["capabilities"]
                    },
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            if response.status_code != 200:
                raise HTTPException(status_code=500, detail="Failed to create user auth token")

            result = response.json()
            token_data = result.get("value")

            if not token_data or not token_data.get("success"):
                raise HTTPException(status_code=500, detail="Failed to authorize user")

        # Clean up authorization code (one-time use)
        del auth_requests[code]

        return {
            "access_token": token_data["userAuthToken"],
            "token_type": "Bearer",
            "expires_in": 31536000,  # 1 year (long-lived for user convenience)
            "scope": " ".join(auth_request["capabilities"]),
            "chat_id": token_data["chatId"],
            "app_id": auth_request["app_id"],
            "privacy_guarantee": {
                "user_controlled": True,
                "developer_cannot_see": True,
                "store_on_user_device_only": True,
                "revocable_anytime": True
            },
            "usage_note": "Store this token securely on user's device - never on your servers"
        }

    except Exception as e:
        # Clean up on error
        if code in auth_requests:
            del auth_requests[code]
        raise HTTPException(status_code=500, detail=f"Failed to exchange authorization code: {str(e)}")

@app.post("/v1/me/query")
async def user_private_query(
    request: SecureUserQueryRequest,
    user_auth_token: str = Header(None, alias="x-user-auth-token"),
    req: Request = None
):
    """
    üîê User Private Query Endpoint

    Users query their private data using their own auth tokens.
    Developers can facilitate this but never see the tokens or raw data.
    """

    if not user_auth_token:
        raise HTTPException(status_code=401, detail="User auth token required in x-user-auth-token header")

    if not user_auth_token.startswith("uat_"):
        raise HTTPException(status_code=401, detail="Invalid user auth token format")

    try:
        # Verify user auth token with Convex
        convex_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")

        async with httpx.AsyncClient() as client:
            # Verify token and get user's chat access
            verify_response = await client.post(
                f"{convex_url}/api/run/user_auth_system/verifyUserAuthToken",
                json={
                    "args": {"userAuthToken": user_auth_token},
                    "format": "json"
                },
                headers={"Content-Type": "application/json"}
            )

            if verify_response.status_code != 200:
                raise HTTPException(status_code=401, detail="Invalid user auth token")

            verify_result = verify_response.json()
            auth_data = verify_result.get("value")

            if not auth_data or not auth_data.get("isValid"):
                raise HTTPException(status_code=401, detail="User auth token expired or revoked")

            chat_id = auth_data["chatId"]

            # Query the user's private chat
            answer_response = await answer_question(
                user_message=request.question,
                chat_id=chat_id,
                selected_model=request.selected_model or "gpt-4o-mini",
                temperature=request.temperature or 0.7,
                max_tokens=request.max_tokens or 1000
            )

            # Filter citations for privacy (developers get limited info)
            filtered_citations = []
            if answer_response.get("reasoning_context"):
                for citation in answer_response["reasoning_context"][:3]:  # Limit to 3
                    filtered_citations.append({
                        "snippet": citation.get("chunk_text", "")[:200] + "...",  # Truncated
                        "score": citation.get("score", 0),
                        "source": "private_document"  # No file names
                    })

            return {
                "success": True,
                "answer": answer_response.get("response", ""),
                "citations": filtered_citations,
                "privacy_note": "This response is generated from the user's private documents. Citations are filtered for privacy.",
                "user_controlled": True,
                "developer_access": "AI responses only - no raw file access"
            }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in user private query: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to process query")

@app.post("/v1/chats/{chat_id}/query/secure")
async def secure_chat_query(
    chat_id: str,
    request: SecureUserQueryRequest,
    user_auth_token: str = Header(None, alias="x-user-auth-token"),
    req: Request = None
):
    """
    üîê SECURE CHAT QUERY: User-controlled authentication for specific chat

    Users call this directly with their private auth token for a specific chat.
    Each chat is treated as its own "app" with OAuth capabilities.
    Developers cannot make this call because they don't have the user's token.
    """

    if not user_auth_token:
        raise HTTPException(status_code=401, detail="Missing user auth token")

    # Verify user auth token with Convex
    try:
        convex_base_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
        convex_url = f"{convex_base_url}/api/run/user_auth_system/verifyUserAuthToken"

        async with httpx.AsyncClient() as client:
            response = await client.post(
                convex_url,
                json={
                    "args": {"userAuthToken": user_auth_token},
                    "format": "json"
                }
            )

            if response.status_code != 200:
                raise HTTPException(status_code=401, detail="Invalid user auth token")

            result = response.json()
            token_info = result.get("value")

            if not token_info:
                raise HTTPException(status_code=401, detail="User auth token not found or expired")

            # Check capabilities
            if "ask" not in token_info["capabilities"]:
                raise HTTPException(status_code=403, detail="User token does not have 'ask' capability")

    except httpx.RequestError:
        raise HTTPException(status_code=500, detail="Failed to verify user auth token")

    # Sanitize question
    sanitized_question = sanitize_with_xss_detection(
        request.question,
        allow_html=False,
        max_length=5000,
        context=f"secure_user_query"
    )

    if not sanitized_question:
        raise HTTPException(status_code=400, detail="Invalid or potentially malicious question")

    try:
        # Track query start time
        query_start_time = time.time()

        # This function is called by the secure chat query endpoint
        # For now, use defaults since this endpoint is less commonly used
        # The main privacy endpoint has the full settings inheritance

        result = await answer_question_with_privacy_scope(
            question=sanitized_question,
            chat_id=token_info["chatId"],
            end_user_id=token_info["trainlyUserId"],
            app_id=token_info["appId"] or "direct",
            model="gpt-4o-mini",
            temperature=0.7,
            max_tokens=1000,
            custom_prompt=None
        )

        response_time = (time.time() - query_start_time) * 1000

        response = {
            "success": True,
            "answer": result["answer"],
            "access_type": "user_controlled",
            "privacy_note": "You are accessing your own private data with your secure token"
        }

        # User gets FULL citations because they're accessing their own data
        if request.include_citations and result.get("context"):
            response["citations"] = [
                {
                    "chunk_id": chunk["chunk_id"],
                    "snippet": chunk["chunk_text"][:500] + "...",  # Longer snippets for user
                    "score": chunk["score"],
                    "source_type": "your_document",
                    "access_level": "full",
                    "privacy_note": "You can see this because it's your own data"
                }
                for chunk in result["context"][:5]  # More citations for user
            ]
            response["citation_access"] = "full"

        # Track successful query
        await track_api_query(
            token_info["appId"] or "direct",
            token_info["trainlyUserId"],
            response_time,
            True
        )

        return response

    except Exception as e:
        # Track failed query
        if 'query_start_time' in locals():
            response_time = (time.time() - query_start_time) * 1000
            await track_api_query(
                token_info["appId"] or "direct",
                token_info["trainlyUserId"],
                response_time,
                False
            )

        raise HTTPException(status_code=500, detail="Failed to process query")

# Legacy endpoint (deprecated but kept for compatibility)
@app.post("/v1/privacy/apps/users/provision")
async def provision_user_subchat(
    request: AppUserRequest,
    authorization: str = Header(None, alias="authorization"),
    req: Request = None
):
    """
    üîí PRIVACY-FIRST: Provision a private sub-chat for an end-user

    This is called when a new user starts using your app.
    Creates an isolated chat that ONLY that user can access.
    The developer cannot see the user's files or raw data.
    """

    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid authorization header")

    app_secret = authorization.replace("Bearer ", "")

    # Verify app secret and get parent chat settings
    app_info = await verify_app_secret(app_secret)

    # Sanitize user ID
    sanitized_user_id = sanitize_text(request.end_user_id)
    if not sanitized_user_id:
        raise HTTPException(status_code=400, detail="Invalid end_user_id")

    # Validate capabilities
    valid_capabilities = ["ask", "upload"]
    invalid_caps = [cap for cap in request.capabilities if cap not in valid_capabilities]
    if invalid_caps:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid capabilities: {invalid_caps}. Only 'ask' and 'upload' are allowed."
        )

    try:
        # Get or create user's private sub-chat
        subchat = await get_or_create_user_subchat(
            app_info["appId"],
            sanitized_user_id
        )

        # Generate scoped token for this user
        scoped_token = generate_scoped_token(
            app_info["appId"],
            sanitized_user_id,
            subchat["chatStringId"],
            request.capabilities
        )

        # Log the provisioning
        await log_app_access(
            app_info["appId"],
            sanitized_user_id,
            subchat["chatStringId"],
            "provision_user",
            "admin",
            True,
            req,
            is_new_user=subchat.get("isNew", False)
        )

        # Track subchat creation in analytics
        if subchat.get("isNew", False):
            await track_subchat_creation(
                app_info["appId"],
                sanitized_user_id,
                subchat["chatStringId"]
            )

        return {
            "success": True,
            "end_user_id": sanitized_user_id,
            "scoped_token": scoped_token,
            "capabilities": request.capabilities,
            "is_new_user": subchat.get("isNew", False),
            "privacy_guarantee": "This user's data is completely isolated - you cannot access their files or raw data"
        }

    except Exception as e:
        # Log failed attempt
        await log_app_access(
            app_info["appId"],
            sanitized_user_id,
            "unknown",
            "provision_user",
            "admin",
            False,
            req,
            error_reason=str(e)
        )
        raise HTTPException(status_code=500, detail=f"Failed to provision user: {str(e)}")

@app.post("/v1/privacy/query")
async def privacy_first_query(
    request: PrivacyFirstQueryRequest,
    token: str = Header(None, alias="x-scoped-token"),
    req: Request = None
):
    """
    üîí PRIVACY-FIRST: Query a user's private chat

    This endpoint ensures complete data isolation:
    - Only returns AI responses, never raw files
    - Scoped to the specific user's data only
    - Cannot access other users' information
    - Comprehensive audit logging
    """

    if not token:
        raise HTTPException(status_code=401, detail="Missing scoped token")

    # Verify scoped token
    claims = await verify_scoped_token(token)

    # Verify user ID matches token
    if request.end_user_id != claims.end_user_id:
        await log_app_access(
            claims.app_id,
            request.end_user_id,
            claims.chat_id,
            "query",
            "ask",
            False,
            req,
            error_reason="User ID mismatch"
        )
        raise HTTPException(status_code=403, detail="User ID mismatch")

    # Check capability
    if "ask" not in claims.capabilities:
        await log_app_access(
            claims.app_id,
            claims.end_user_id,
            claims.chat_id,
            "query",
            "ask",
            False,
            req,
            error_reason="Missing 'ask' capability"
        )
        raise HTTPException(status_code=403, detail="Missing 'ask' capability")

    # Sanitize question
    sanitized_question = sanitize_with_xss_detection(
        request.question,
        allow_html=False,
        max_length=5000,
        context=f"privacy_query_{claims.end_user_id}"
    )

    if not sanitized_question:
        await log_app_access(
            claims.app_id,
            claims.end_user_id,
            claims.chat_id,
            "query",
            "ask",
            False,
            req,
            error_reason="Invalid or malicious question",
            question=request.question[:100]
        )
        raise HTTPException(status_code=400, detail="Invalid or potentially malicious question")

    try:
        # Track query start time for performance metrics
        query_start_time = time.time()

        # Get app settings including parent chat settings from Convex
        app_settings = {}
        developer_user_id = claims.app_id

        try:
            convex_base_url = os.getenv("CONVEX_URL", "https://agile-ermine-199.convex.cloud")
            async with httpx.AsyncClient() as client:
                app_response = await client.post(
                    f"{convex_base_url}/api/run/app_management/getAppWithSettings",
                    json={
                        "args": {"appId": claims.app_id},
                        "format": "json"
                    }
                )

                if app_response.status_code == 200:
                    app_data = app_response.json()
                    if app_data.get("value") and app_data["value"].get("parentChatSettings"):
                        parent_settings = app_data["value"]["parentChatSettings"]
                        app_settings = {
                            "custom_prompt": parent_settings.get("customPrompt"),
                            "selected_model": parent_settings.get("selectedModel", "gpt-4o-mini"),
                            "temperature": parent_settings.get("temperature", 0.7),
                            "max_tokens": int(parent_settings.get("maxTokens", 1000))
                        }
                        developer_user_id = parent_settings.get("userId") or claims.app_id
                        logger.info(f"üìã SUCCESS: Inherited settings from parent chat: model={app_settings['selected_model']}, temp={app_settings['temperature']}, custom_prompt={bool(app_settings['custom_prompt'])}")
                        logger.info(f"üí≥ Found developer user ID: {developer_user_id}")
                    else:
                        logger.info(f"üìã No parent chat settings found for app {claims.app_id}, using defaults")
                else:
                    logger.warning(f"Could not fetch app settings: {app_response.status_code}")
        except Exception as e:
            logger.error(f"Error fetching app settings: {str(e)}")

        # Use inherited settings with fallbacks
        model = app_settings.get("selected_model", "gpt-4o-mini")
        temperature = app_settings.get("temperature", 0.7)
        max_tokens = int(app_settings.get("max_tokens", 1000))
        custom_prompt = app_settings.get("custom_prompt")

        logger.info(f"üìã Final settings: model={model}, temp={temperature}, custom_prompt={bool(custom_prompt)}")
        logger.info(f"üí≥ Credit consumption target: {developer_user_id}")

        # Call the privacy-aware answer function with request settings
        # This ensures data is scoped to the user's specific chat
        result = await answer_question_with_privacy_scope(
            question=sanitized_question,
            chat_id=claims.chat_id,
            end_user_id=claims.end_user_id,
            app_id=claims.app_id,
            model=model,
            temperature=temperature,
            max_tokens=max_tokens,
            custom_prompt=custom_prompt
        )

        # Now consume credits based on actual usage (after getting the response)
        try:
            credits_consumed = await consume_credits_for_actual_usage(
                user_id=developer_user_id,
                model=model,
                question=sanitized_question,
                response=result.get("answer", ""),
                chat_id=claims.chat_id  # Use the user's chat ID for tracking
            )
            logger.info(f"üí≥ SUCCESS: Consumed {credits_consumed} credits from developer {developer_user_id} for end user {claims.end_user_id} based on actual usage")
        except InsufficientCreditsError as e:
            logger.error(f"üí≥ CREDIT ERROR after API call: need {e.required}, have {e.available}")
            # Since we've already made the AI call, we'll allow the response but log the issue
            logger.warning(f"üí≥ Allowing response due to insufficient credits - this should be monitored")
            credits_consumed = 0
        except Exception as e:
            logger.error(f"üí≥ ERROR: Credit consumption failed: {str(e)}")
            logger.warning(f"üí≥ Allowing response despite credit error - this should be monitored")
            credits_consumed = 0

        # Calculate response time and track metrics
        response_time = (time.time() - query_start_time) * 1000  # Convert to milliseconds
        await track_api_query(
            claims.app_id,
            claims.end_user_id,
            response_time,
            True  # Success
        )

        # Log successful query
        await log_app_access(
            claims.app_id,
            claims.end_user_id,
            claims.chat_id,
            "query",
            "ask",
            True,
            req,
            question=sanitized_question[:100],
            used_chunks=len(result.get("context", []))
        )

        response = {
            "success": True,
            "answer": result["answer"],
            "end_user_id": claims.end_user_id,
            "privacy_note": "This response contains only AI-generated content based on the user's private data"
        }

        # Include citations if requested (but never raw file content)
        if request.include_citations and result.get("context"):
            response["citations"] = [
                {
                    "chunk_id": chunk["chunk_id"],
                    "snippet": chunk["chunk_text"][:200] + "...",  # Limited snippet only
                    "score": chunk["score"]
                }
                for chunk in result["context"][:3]  # Max 3 citations
            ]

        return response

    except Exception as e:
        # Track failed query
        if 'query_start_time' in locals():
            response_time = (time.time() - query_start_time) * 1000
            await track_api_query(
                claims.app_id,
                claims.end_user_id,
                response_time,
                False  # Failed
            )

        await log_app_access(
            claims.app_id,
            claims.end_user_id,
            claims.chat_id,
            "query",
            "ask",
            False,
            req,
            error_reason=str(e),
            question=sanitized_question[:100]
        )
        raise HTTPException(status_code=500, detail="Failed to process query")

@app.post("/v1/privacy/upload/presigned-url")
async def get_presigned_upload_url(
    request: DirectUploadRequest,
    token: str = Header(None, alias="x-scoped-token"),
    req: Request = None
):
    """
    üîí PRIVACY-FIRST: Get presigned URL for direct user uploads

    This bypasses the developer's servers entirely:
    - Files upload directly from user's browser to Trainly
    - Developer never sees the file content
    - Complete data isolation maintained
    """

    if not token:
        raise HTTPException(status_code=401, detail="Missing scoped token")

    # Verify scoped token
    claims = await verify_scoped_token(token)

    # Verify user ID and capability
    if (request.end_user_id != claims.end_user_id or
        "upload" not in claims.capabilities):
        await log_app_access(
            claims.app_id,
            request.end_user_id,
            claims.chat_id,
            "upload_request",
            "upload",
            False,
            req,
            error_reason="Unauthorized upload attempt"
        )
        raise HTTPException(status_code=403, detail="Unauthorized upload")

    # Sanitize filename
    sanitized_filename = sanitize_filename(request.filename)
    if not sanitized_filename:
        raise HTTPException(status_code=400, detail="Invalid filename")

    try:
        # Generate upload URL for user's namespace
        # Routes to a privacy-aware upload endpoint that processes files completely
        # Files are processed securely and isolated per user/app
        base_url = os.getenv("API_BASE_URL", "http://localhost:8000")
        upload_url = f"{base_url}/v1/privacy/upload/process"

        # Log upload request
        await log_app_access(
            claims.app_id,
            claims.end_user_id,
            claims.chat_id,
            "upload_request",
            "upload",
            True,
            req,
            filename=sanitized_filename,
            file_type=request.file_type
        )

        return {
            "success": True,
            "upload_url": upload_url,
            "filename": sanitized_filename,
            "expires_in": 3600,  # 1 hour
            "privacy_note": "File will be uploaded directly to user's private namespace",
            "upload_method": "POST",  # Indicates to use POST with FormData instead of PUT
            "content_type": "multipart/form-data",
            "field_name": "file",  # The form field name to use
            "upload_headers": {  # Headers to include in the upload request
                "x-chat-id": claims.chat_id,
                "x-user-id": claims.end_user_id,
                "x-app-id": claims.app_id
            }
        }

    except Exception as e:
        await log_app_access(
            claims.app_id,
            claims.end_user_id,
            claims.chat_id,
            "upload_request",
            "upload",
            False,
            req,
            error_reason=str(e),
            filename=sanitized_filename
        )
        raise HTTPException(status_code=500, detail="Failed to generate upload URL")

@app.post("/v1/privacy/upload/process")
async def privacy_upload_process(
    file: UploadFile = File(...),
    chat_id: str = Header(None, alias="x-chat-id"),
    user_id: str = Header(None, alias="x-user-id"),
    app_id: str = Header(None, alias="x-app-id")
):
    """
    üîí PRIVACY-FIRST: Complete file upload processing

    PERFORMANCE OPTIMIZED:
    - Text extraction happens immediately (fast)
    - Heavy processing (chunking, embedding, Neo4j) is queued for background workers
    - Returns immediately with processing status
    - Poll file status endpoint for completion

    This endpoint:
    1. Extracts text from uploaded files
    2. Queues embeddings creation for the user's sub-chat
    3. Ensures complete privacy isolation
    """

    # Validate file
    file_size = read_files.get_file_size(file)
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"File too large (max {MAX_FILE_SIZE // (1024 * 1024)} MB).")
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded.")

    # Validate required headers
    if not chat_id or not user_id or not app_id:
        raise HTTPException(status_code=400, detail="Missing required headers: x-chat-id, x-user-id, x-app-id")

    # Sanitize inputs
    sanitized_filename = sanitize_filename(file.filename)
    sanitized_chat_id = sanitize_chat_id(chat_id)
    sanitized_user_id = sanitize_api_key(user_id)  # Reuse sanitization function
    sanitized_app_id = sanitize_api_key(app_id)    # Reuse sanitization function

    if not sanitized_filename or not sanitized_chat_id or not sanitized_user_id or not sanitized_app_id:
        raise HTTPException(status_code=400, detail="Invalid input parameters")

    try:
        # Step 1: Extract text from file (this is fast, keep inline)
        text = read_files.extract_text(file)

        # Sanitize extracted text
        sanitized_text = sanitize_with_xss_detection(
            text,
            allow_html=False,
            max_length=1000000,
            context="privacy_file_extraction"
        )

        if not sanitized_text and text:
            raise HTTPException(status_code=400, detail="File contains potentially malicious content.")

        # Step 2: Generate IDs and enqueue for background processing
        pdf_id = f"{sanitized_user_id}_{sanitized_filename}_{int(time.time())}"
        file_queue_id = f"fq_{pdf_id}"

        # Check if queue is available
        queue = get_file_queue()

        if queue is not None:
            # Enqueue for background processing
            try:
                job = queue.enqueue(
                    process_file_job,
                    file_queue_id,
                    sanitized_chat_id,
                    sanitized_filename,
                    sanitized_text,
                    file_size,
                    {},  # No scope values for privacy uploads
                    pdf_id,
                    job_timeout=1800,  # 30 minute timeout
                    result_ttl=86400,  # Keep result for 24 hours
                )

                logger.info(f"üì§ Privacy upload: Enqueued {sanitized_filename} as job {job.id}")

                return {
                    "success": True,
                    "message": "File uploaded and queued for processing",
                    "filename": sanitized_filename,
                    "chat_id": sanitized_chat_id,
                    "file_id": pdf_id,
                    "file_queue_id": file_queue_id,
                    "job_id": job.id,
                    "processing_status": "queued",
                    "privacy_note": "File will be processed and stored in your isolated workspace"
                }

            except Exception as enqueue_error:
                logger.error(f"Failed to enqueue privacy upload: {enqueue_error}")
                raise HTTPException(
                    status_code=503,
                    detail="File processing queue unavailable. Please try again later."
                )
        else:
            # Queue not available - return error (no fallback to sync processing)
            raise HTTPException(
                status_code=503,
                detail="File processing service is not available. Please contact support."
            )

    except HTTPException as http_exc:
        raise http_exc
    except Exception as e:
        logger.error(f"Privacy upload processing failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to process upload: {str(e)}")

@app.get("/debug/api-base-url")
async def debug_api_base_url():
    """Debug endpoint to check what API base URL is being used"""
    base_url = os.getenv("API_BASE_URL", "http://localhost:8000")
    return {
        "api_base_url": base_url,
        "environment_var": os.getenv("API_BASE_URL"),
        "default_used": os.getenv("API_BASE_URL") is None
    }

@app.get("/debug/chat-settings/{chat_id}")
async def debug_chat_settings(chat_id: str):
    """Debug endpoint to test chat settings fetching"""
    try:
        logger.info(f"üîç DEBUG: Fetching settings for chat {chat_id}")
        logger.info(f"üîç DEBUG: Using CONVEX_URL: {CONVEX_URL}")

        async with httpx.AsyncClient() as client:
            chat_response = await client.post(
                CONVEX_URL,
                json={
                    "args": {"id": chat_id},
                    "format": "json"
                }
            )

            logger.info(f"üîç DEBUG: Convex response status: {chat_response.status_code}")
            logger.info(f"üîç DEBUG: Convex response text: {chat_response.text}")

            if chat_response.status_code == 200:
                chat_data = chat_response.json()
                logger.info(f"üîç DEBUG: Parsed chat data: {chat_data}")

                if chat_data.get("value"):
                    chat = chat_data["value"]
                    settings = {
                        "custom_prompt": chat.get("customPrompt"),
                        "selected_model": chat.get("selectedModel"),
                        "temperature": chat.get("temperature"),
                        "max_tokens": int(chat.get("maxTokens", 1000)) if chat.get("maxTokens") is not None else 1000,
                        "user_id": chat.get("userId")
                    }
                    logger.info(f"üîç DEBUG: Extracted settings: {settings}")
                    return {"success": True, "settings": settings, "raw_chat": chat}
                else:
                    return {"success": False, "error": "No chat data in response", "response": chat_data}
            else:
                return {"success": False, "error": f"HTTP {chat_response.status_code}", "response": chat_response.text}

    except Exception as e:
        logger.error(f"üîç DEBUG: Exception in debug endpoint: {str(e)}")
        return {"success": False, "error": str(e)}

@app.get("/v1/privacy/health")
async def privacy_api_health():
    """Health check for the privacy-first API"""
    return {
        "status": "healthy",
        "message": "Privacy-First API is operational",
        "privacy_model": "Complete data isolation - developers cannot access user files or raw data",
        "capabilities": {
            "allowed": ["ask", "upload"],
            "blocked": ["list_files", "download_file", "read_raw_data"],
        },
        "compliance": "GDPR/CCPA ready with user data ownership",
        "audit_trail": "All access attempts logged",
        "timestamp": time.time()
    }

async def answer_question_with_privacy_scope(
    question: str,
    chat_id: str,
    end_user_id: str,
    app_id: str,
    model: str = "gpt-4o-mini",
    temperature: float = 0.7,
    max_tokens: int = 1000,
    custom_prompt: str = None
) -> Dict[str, Any]:
    """
    Privacy-scoped question answering that ensures complete data isolation.
    This function ensures that responses only include data from the specific user's chat.
    """
    # This integrates with your existing answer_question logic
    # but with additional privacy controls and scoping

    try:
        # Generate question embedding
        question_embedding = get_embedding(question)

        # Query Neo4j with user-specific filtering
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Enhanced query that includes user/chat isolation + parent chat access
                # This searches BOTH user's private data AND parent app knowledge base
                # Note: This implements the hybrid model you described

                # First, let's see what documents exist for this chat
                doc_query = f"""
                MATCH (d:Document)
                WHERE d.chatId = '{chat_id}' OR (d.chatId = '{app_id}' AND '{app_id}' <> 'direct')
                RETURN d.filename AS filename, d.chatId AS docChatId
                """
                doc_result = session.run(doc_query)
                doc_list = [(record["filename"], record["docChatId"]) for record in doc_result]
                logger.info(f"üîç DEBUG: Found {len(doc_list)} documents: {doc_list}")

                query = f"""
                MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                WHERE c.chatId = '{chat_id}' OR (c.chatId = '{app_id}' AND '{app_id}' <> 'direct')
                RETURN c.id AS id, c.text AS text, c.embedding AS embedding, d.filename AS filename
                """

                result = session.run(query)
                chunks = []

                question_lower = question.lower()

                # Calculate similarities with filename boost - EXACT SAME AS MAIN CHAT
                chunk_scores = []

                for record in result:
                    chunk_id = record["id"]
                    chunk_text = record["text"]
                    chunk_embedding = record["embedding"]
                    filename = record["filename"] or ""

                    # Calculate semantic similarity
                    # Ensure both are numpy arrays with correct dimensions
                    q_emb = np.array(question_embedding)
                    c_emb = np.array(chunk_embedding) if isinstance(chunk_embedding, list) else chunk_embedding
                    semantic_score = cosine_similarity(q_emb, c_emb)

                    # Add filename relevance boost (same logic as main chat)
                    filename_boost = 0.0
                    if filename:
                        filename_lower = filename.lower()
                        # Remove file extension for better matching
                        filename_base = filename_lower.replace('.pdf', '').replace('.docx', '').replace('.txt', '')

                        # Check for filename mentions in question
                        filename_words = filename_base.replace('_', ' ').replace('-', ' ').split()
                        question_words = question_lower.replace('_', ' ').replace('-', ' ').split()

                        # Boost score if filename words appear in question
                        for fname_word in filename_words:
                            if len(fname_word) > 2:  # Skip very short words
                                for q_word in question_words:
                                    if fname_word in q_word or q_word in fname_word:
                                        filename_boost += 0.1

                        # Additional boost for exact filename matches
                        if any(fname_word in question_lower for fname_word in filename_words if len(fname_word) > 3):
                            filename_boost += 0.2

                    # Combine semantic similarity with filename relevance
                    final_score = semantic_score + filename_boost

                    chunk_scores.append({
                        "chunk_id": chunk_id,
                        "chunk_text": chunk_text,
                        "score": float(final_score),
                        "filename": filename
                    })

                # Sort and get top chunks - EXACT SAME AS MAIN CHAT
                chunk_scores.sort(key=lambda x: x["score"], reverse=True)
                top_k = 50  # Same as main chat
                top_chunks = chunk_scores[:top_k]

                # Generate answer using GPT-4 with citations - EXACT SAME AS MAIN CHAT
                # Limit to top 10 chunks for cleaner citations
                top_chunks_for_citations = top_chunks[:10]

                # Get document information for each chunk to provide better context - EXACT SAME AS MAIN CHAT
                chunk_document_info = {}
                for chunk in top_chunks_for_citations:
                    doc_query = f"""
                    MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk {{id: '{chunk["chunk_id"]}'}})
                    RETURN d.filename AS filename
                    """
                    doc_result = session.run(doc_query)
                    doc_record = doc_result.single()
                    if doc_record:
                        chunk_document_info[chunk["chunk_id"]] = doc_record['filename']

                context_with_ids = "\n\n---\n\n".join([
                    f"[CHUNK_{i}] (from document: {chunk_document_info.get(chunk['chunk_id'], 'Unknown')}) {chunk['chunk_text']}"
                    for i, chunk in enumerate(top_chunks_for_citations)
                ])

                # Use custom prompt if provided, otherwise use default system prompt - EXACT SAME AS MAIN CHAT
                if custom_prompt:
                    # Use custom prompt but ensure context is included
                    system_prompt = f"""
                    {custom_prompt}

                    You have the following context with chunk IDs (0-{len(top_chunks_for_citations)-1}):

                    {context_with_ids}

                    When you reference information from the context, add a citation using this format: [^{{i}}] where {{i}} is the chunk number (0-{len(top_chunks_for_citations)-1})
                    Only use citations [^0] through [^{len(top_chunks_for_citations)-1}]. Do not use citation numbers higher than {len(top_chunks_for_citations)-1}
                    """.strip()
                else:
                    # Default system prompt - EXACT SAME AS MAIN CHAT
                    system_prompt = f"""
                    You are a helpful assistant. You have the following context with chunk IDs (0-{len(top_chunks_for_citations)-1}):

                    {context_with_ids}

                    IMPORTANT INSTRUCTIONS:
                    1. ALWAYS prioritize using the provided context to answer the user's question
                    2. If the context contains relevant information, you MUST use it and cite it properly
                    3. Pay attention to the document names shown in parentheses - if the user asks about a specific document by name, use the content from that document
                    4. When you reference information from the context, add a citation using this format: [^{{i}}] where {{i}} is the chunk number (0-{len(top_chunks_for_citations)-1})
                    5. Only use citations [^0] through [^{len(top_chunks_for_citations)-1}]. Do not use citation numbers higher than {len(top_chunks_for_citations)-1}
                    6. Only use external knowledge if the context is completely irrelevant to the question, and clearly state when you're using external knowledge

                    For example:
                    - "According to the Grant Assignment document [^0], ecology research involves..."
                    - "The document shows [^2] that species interactions..."

                    RESPOND IN MARKDOWN FORMAT WITH CITATIONS
                    """.strip()

                completion = openai.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": question}
                    ],
                    temperature=temperature,
                    max_tokens=max_tokens
                )

                answer = completion.choices[0].message.content.strip()

                # Convert to same format as main chat
                context_chunks = []
                for chunk in top_chunks_for_citations:
                    context_chunks.append({
                        "chunk_id": chunk["chunk_id"],
                        "chunk_text": chunk["chunk_text"],
                        "score": chunk["score"]
                    })

                return {
                    "answer": answer,
                    "context": context_chunks,
                    "privacy_scope": {
                        "app_id": app_id,
                        "end_user_id": end_user_id,
                        "chat_id": chat_id,
                        "isolation_confirmed": True
                    }
                }

    except Exception as e:
        import traceback
        logger.error(f"Error in privacy-scoped answer: {str(e)}")
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return {
            "answer": f"Debug: Error occurred - {str(e)}. Please check server logs for details.",
            "context": [],
            "privacy_scope": {
                "app_id": app_id,
                "end_user_id": end_user_id,
                "chat_id": chat_id,
                "isolation_confirmed": True,
                "error": str(e)
            }
        }

# ==============================================================================
# üîê BYO OAuth Token Exchange (RFC 8693) - Lightweight Implementation
# ==============================================================================

@app.get("/oauth/authorize")
async def oauth_authorize(
    chat_id: str,
    redirect_uri: str,
    state: Optional[str] = None,
    scope: Optional[str] = "chat.query chat.upload"
):
    """
    üîê Trainly OAuth Authorization Endpoint

    User visits this URL to authorize an app to access their chat.
    Returns authorization code that can be exchanged for access token.
    """

    # Sanitize inputs
    sanitized_chat_id = sanitize_chat_id(chat_id)
    if not sanitized_chat_id:
        raise HTTPException(status_code=400, detail="Invalid chat_id")

    # In production, this would show authorization page
    # For demo, we'll auto-approve and redirect with auth code
    auth_code = f"auth_{int(time.time())}_{sanitized_chat_id}_{state or 'default'}"

    redirect_url = f"{redirect_uri}?code={auth_code}&state={state or ''}"

    logger.info(f"üîê OAuth authorization: {sanitized_chat_id} ‚Üí {redirect_uri}")

    return {
        "authorization_url": redirect_url,
        "auth_code": auth_code,
        "chat_id": sanitized_chat_id,
        "message": "In production, user would see authorization page here"
    }

@app.post("/oauth/token")
async def token_exchange(
    grant_type: str,
    code: Optional[str] = None,
    redirect_uri: Optional[str] = None,
    client_id: Optional[str] = None,
    scope: Optional[str] = "chat.query chat.upload"
):
    """
    üîê Trainly OAuth Token Exchange

    Exchange authorization code for user access token.
    Lightweight implementation with complete privacy protection.
    """

    if grant_type != "authorization_code":
        raise HTTPException(status_code=400, detail="Unsupported grant_type")

    if not code or not client_id:
        raise HTTPException(status_code=400, detail="Missing code or client_id")

    try:
        # Validate authorization code
        if not code.startswith("auth_"):
            raise HTTPException(status_code=401, detail="Invalid authorization code")

        # Extract info from auth code
        code_parts = code.split("_")
        if len(code_parts) < 3:
            raise HTTPException(status_code=401, detail="Malformed authorization code")

        chat_id = code_parts[2]
        user_id = f"user_{int(time.time())}_{code_parts[1][-8:]}"  # Generate user ID from code

        # Sanitize inputs
        sanitized_chat_id = sanitize_chat_id(chat_id)
        if not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid chat_id")

        subchat_id = f"subchat_{sanitized_chat_id}_{user_id}"

        # Generate Trainly access token
        trainly_token = generate_trainly_token(
            chat_id=sanitized_chat_id,
            user_id=user_id,
            subchat_id=subchat_id,
            scopes=scope.split() if scope else ["chat.query", "chat.upload"]
        )

        # Track token creation
        await track_subchat_creation(sanitized_chat_id, user_id, subchat_id)

        logger.info(f"üîê Token exchange: {user_id[:8]}... ‚Üí {sanitized_chat_id}")

        return {
            "access_token": trainly_token,
            "token_type": "Bearer",
            "expires_in": 3600,  # 1 hour
            "scope": scope,
            "chat_id": sanitized_chat_id,
            "subchat_id": subchat_id,
            "privacy_guarantee": {
                "user_controlled": True,
                "citations_filtered_for_apps": True,
                "no_raw_file_access": True,
                "developer_cannot_see_token": True
            },
            "usage_note": "Store this token on user's device only - never on your servers"
        }

    except Exception as e:
        logger.error(f"Token exchange failed: {e}")
        raise HTTPException(status_code=500, detail="Token exchange failed")

def generate_trainly_token(chat_id: str, user_id: str, subchat_id: str, scopes: List[str]) -> str:
    """Generate Trainly access token with privacy protection"""
    secret_key = os.getenv("TRAINLY_JWT_SECRET", "your-trainly-secret-key")

    now = time.time()
    payload = {
        "iss": "trainly.com",
        "sub": user_id,
        "aud": "trainly-api",
        "exp": now + 3600,  # 1 hour expiry
        "iat": now,
        "scope": " ".join(scopes),
        "chat_id": chat_id,
        "subchat_id": subchat_id,
        "user_id": user_id,
        "privacy_mode": "strict",
        "citations_filtered": True,
        "token_type": "trainly_oauth"
    }

    return jwt.encode(payload, secret_key, algorithm="HS256")

@app.post("/me/chats/query")
async def user_chat_query(
    request: SecureUserQueryRequest,
    authorization: str = Header(None),
    req: Request = None
):
    """
    üîê User queries their private chat with Trainly token
    Citations filtered for developer calls, full for direct user access.
    """

    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing Bearer token")

    token = authorization.replace("Bearer ", "")

    try:
        secret_key = os.getenv("TRAINLY_JWT_SECRET", "your-trainly-secret-key")
        claims = jwt.decode(token, secret_key, algorithms=["HS256"], audience="trainly-api")

        if claims["exp"] < time.time():
            raise HTTPException(status_code=401, detail="Token expired")

    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

    # Extract from token
    chat_id = claims["chat_id"]
    subchat_id = claims["subchat_id"]
    user_id = claims["user_id"]

    # Check scope
    if "chat.query" not in claims.get("scope", ""):
        raise HTTPException(status_code=403, detail="Insufficient scope")

    # Sanitize question
    sanitized_question = sanitize_with_xss_detection(
        request.question,
        allow_html=False,
        max_length=5000,
        context=f"user_query_{user_id[:8]}"
    )

    if not sanitized_question:
        raise HTTPException(status_code=400, detail="Invalid question")

    try:
        query_start_time = time.time()

        # For BYO OAuth endpoint, use simplified settings for now
        # The main privacy endpoint has full inheritance

        result = await answer_question_with_privacy_scope(
            question=sanitized_question,
            chat_id=subchat_id,
            end_user_id=user_id,
            app_id=chat_id,
            model="gpt-4o-mini",
            temperature=0.7,
            max_tokens=1000,
            custom_prompt=None
        )

        response_time = (time.time() - query_start_time) * 1000

        # Detect developer app call vs direct user access
        is_developer_call = True
        if req:
            origin = req.headers.get("origin", "")
            if "trainly.com" in origin or "localhost:3000" in origin:
                is_developer_call = False

        response = {
            "answer": result["answer"],
            "subchat_id": subchat_id,
            "access_type": "user_controlled",
            "privacy_note": "Generated from your private data only"
        }

        # Citation filtering for privacy protection
        if request.include_citations and result.get("context"):
            if is_developer_call:
                # Developer gets summary only (privacy protection)
                response["citations_summary"] = {
                    "sources_used": len(result["context"]),
                    "confidence": "high" if result["context"][0]["score"] > 0.8 else "medium",
                    "privacy_note": "Citations filtered for privacy - full access at trainly.com"
                }
            else:
                # Direct user gets full citations
                response["citations"] = [
                    {
                        "chunk_id": chunk["chunk_id"],
                        "snippet": chunk["chunk_text"][:500] + "...",
                        "score": chunk["score"]
                    }
                    for chunk in result["context"][:5]
                ]

        await track_api_query(chat_id, user_id, response_time, True)
        return response

    except Exception as e:
        logger.error(f"Query failed: {e}")
        raise HTTPException(status_code=500, detail="Query failed")

# ==============================================================================
# Scope Management API Endpoints
# ==============================================================================

@app.post("/v1/{chat_id}/scopes/configure")
async def configure_scopes(
    chat_id: str,
    scope_config: AppScopeConfig,
    credentials: HTTPAuthorizationCredentials = Depends(get_verified_chat_access)
):
    """
    Configure custom scopes for a chat.

    This allows developers to define custom scope fields like playlist_id, workspace_id, etc.
    that will be used to segment data within the chat.

    Example:
    {
        "scopes": [
            {
                "name": "playlist_id",
                "type": "string",
                "required": true,
                "description": "ID of the playlist this document belongs to"
            },
            {
                "name": "workspace_id",
                "type": "string",
                "required": false,
                "description": "Optional workspace identifier"
            }
        ]
    }
    """
    try:
        sanitized_chat_id = sanitize_chat_id(chat_id)
        if not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid chat_id format")

        # Validate scope names
        for scope_def in scope_config.scopes:
            if not validate_scope_name(scope_def.name):
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid scope name: {scope_def.name}. Must start with letter and contain only alphanumeric, underscores, hyphens."
                )

            if scope_def.type not in ["string", "number", "boolean"]:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid scope type: {scope_def.type}. Must be string, number, or boolean."
                )

        # Save the configuration
        success = await save_scope_config(sanitized_chat_id, scope_config)

        if success:
            logger.info(f"‚úÖ Scope configuration saved for chat {sanitized_chat_id}: {[s.name for s in scope_config.scopes]}")
            return {
                "success": True,
                "message": "Scope configuration saved successfully",
                "chat_id": sanitized_chat_id,
                "scopes": [s.dict() for s in scope_config.scopes]
            }
        else:
            raise HTTPException(status_code=500, detail="Failed to save scope configuration")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to configure scopes: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/v1/{chat_id}/scopes")
async def get_scopes(
    chat_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(get_verified_chat_access)
):
    """
    Get the current scope configuration for a chat.

    Returns the list of configured scopes and their definitions.
    """
    try:
        sanitized_chat_id = sanitize_chat_id(chat_id)
        if not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid chat_id format")

        scope_config = await get_scope_config(sanitized_chat_id)

        if scope_config:
            return {
                "chat_id": sanitized_chat_id,
                "scopes": [s.dict() for s in scope_config.scopes]
            }
        else:
            return {
                "chat_id": sanitized_chat_id,
                "scopes": []
            }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get scopes: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/v1/{chat_id}/scopes")
async def delete_scopes(
    chat_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(get_verified_chat_access)
):
    """
    Delete/clear all scope configurations for a chat.

    Note: This does not remove scope properties from existing nodes,
    it only clears the scope configuration.
    """
    try:
        sanitized_chat_id = sanitize_chat_id(chat_id)
        if not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid chat_id format")

        # Clear in-memory config
        if sanitized_chat_id in SCOPE_CONFIGS:
            del SCOPE_CONFIGS[sanitized_chat_id]

        # Also clear in Convex (you would implement this based on your Convex schema)
        logger.info(f"üóëÔ∏è Cleared scope configuration for chat {sanitized_chat_id}")

        return {
            "success": True,
            "message": "Scope configuration cleared",
            "chat_id": sanitized_chat_id
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete scopes: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# NOTE: FILE_PROCESSING_STATUS and _process_file_background have been removed.
# All file processing now goes through the Redis job queue (process_file_job).
# Status is tracked in Convex via the file_upload_queue table.

@app.post("/v1/{chat_id}/upload_with_scopes")
async def upload_file_with_scopes(
    chat_id: str,
    file: UploadFile = File(...),
    scope_values: str = Form("{}"),  # JSON string of scope values
    credentials: HTTPAuthorizationCredentials = Depends(get_verified_chat_access),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    """
    Upload a file with custom scope values.

    PERFORMANCE OPTIMIZED:
    - Text extraction happens immediately (fast)
    - Heavy processing (chunking, embedding, Neo4j) is queued for background workers
    - Returns immediately with processing status
    - Poll GET /v1/{chat_id}/files/{file_id}/status for completion

    Example scope_values:
    {
        "playlist_id": "playlist_123",
        "workspace_id": "workspace_456"
    }
    """
    try:
        import json

        sanitized_chat_id = sanitize_chat_id(chat_id)
        if not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid chat_id format")

        # Parse scope values
        try:
            parsed_scope_values = json.loads(scope_values)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="Invalid scope_values JSON")

        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file uploaded")

        sanitized_filename = sanitize_filename(file.filename)
        if not sanitized_filename:
            raise HTTPException(status_code=400, detail="Invalid filename")

        # Get actual file size
        file.file.seek(0, 2)  # Seek to end
        file_size = file.file.tell()
        file.file.seek(0)  # Reset to beginning

        if file_size > MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail=f"File too large (max {MAX_FILE_SIZE // (1024 * 1024)} MB)")

        # Extract text from file (this is fast, keep inline)
        file_content = await file.read()
        file_obj = io.BytesIO(file_content)

        # Determine file type and extract text
        if sanitized_filename.endswith('.pdf'):
            from pypdf import PdfReader
            reader = PdfReader(file_obj)
            text = "\n".join([page.extract_text() for page in reader.pages])
        elif sanitized_filename.endswith('.txt'):
            text = file_content.decode('utf-8')
        elif sanitized_filename.endswith('.docx'):
            import docx
            doc = docx.Document(file_obj)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        # Sanitize extracted text
        sanitized_text = sanitize_with_xss_detection(
            text,
            allow_html=False,
            max_length=1000000,
            context="file_upload_with_scopes"
        )

        if not sanitized_text and text:
            raise HTTPException(status_code=400, detail="File contains potentially malicious content")

        # Create file ID
        pdf_id = f"{sanitized_chat_id}_{sanitized_filename}_{int(time.time())}"
        file_queue_id = f"fq_{pdf_id}"

        # Try to enqueue for background processing (PERFORMANCE OPTIMIZATION)
        queue = get_file_queue()

        if queue is not None:
            # Queue is available - use async processing via Redis
            logger.info(f"üì§ Enqueueing file for background processing: {sanitized_filename}")

            try:
                job = queue.enqueue(
                    process_file_job,
                    file_queue_id,
                    sanitized_chat_id,
                    sanitized_filename,
                    sanitized_text,
                    file_size,
                    parsed_scope_values,
                    pdf_id,
                    job_timeout=1800,  # 30 minute timeout
                    result_ttl=86400,  # Keep result for 24 hours
                )

                logger.info(f"‚úÖ Job enqueued: {job.id} for {sanitized_filename}")

                return {
                    "success": True,
                    "file_id": pdf_id,
                    "file_queue_id": file_queue_id,
                    "job_id": job.id,
                    "status": "queued",
                    "filename": sanitized_filename,
                    "chat_id": sanitized_chat_id,
                    "scope_values": parsed_scope_values,
                    "size_bytes": file_size,
                    "status_endpoint": f"/v1/{sanitized_chat_id}/files/{pdf_id}/status",
                    "message": "File uploaded and queued for processing. Poll status endpoint for completion."
                }

            except Exception as e:
                logger.error(f"‚ùå Queue enqueue failed: {e}")
                raise HTTPException(
                    status_code=503,
                    detail="File processing queue unavailable. Please try again later."
                )

        # Redis queue is required for production - no fallback to blocking processing
        logger.error("‚ùå Redis queue not available - cannot process file")
        raise HTTPException(
            status_code=503,
            detail="File processing service is not available. Please contact support."
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to upload file with scopes: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/v1/{chat_id}/files/{file_id}/status")
async def get_file_status(
    chat_id: str,
    file_id: str,
    api_key: Optional[str] = Depends(get_optional_chat_access)
):
    """
    Get the processing status of a file.

    Returns:
        {
            "file_id": str,
            "status": "processing" | "ready" | "failed",
            "filename": str,
            "size_bytes": int,
            "error": str (if failed)
        }
    """
    try:
        # Sanitize chat_id and file_id first
        sanitized_chat_id = sanitize_chat_id(chat_id)
        if not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid chat_id format")

        sanitized_file_id = sanitize_text(file_id)
        if not sanitized_file_id:
            raise HTTPException(status_code=400, detail="Invalid file_id format")

        # Verify API key and chat access if credentials are provided
        if api_key:
            is_valid = await verify_api_key_and_chat(api_key, sanitized_chat_id)

            if not is_valid:
                raise HTTPException(
                    status_code=401,
                    detail="Invalid API key or chat not accessible."
                )

        # Check if file exists in Neo4j first (most reliable, works across workers)
        # This is the source of truth - if file is in Neo4j with chunks, it's ready
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                query = """
                MATCH (d:Document {id: $file_id})-[:HAS_CHUNK]->(c:Chunk)
                WHERE c.chatId = $chat_id
                WITH d, count(c) as chunk_count
                RETURN d.filename as filename, d.sizeBytes as size_bytes, chunk_count
                LIMIT 1
                """
                result = session.run(query, file_id=sanitized_file_id, chat_id=sanitized_chat_id)
                record = result.single()

                if record and record["chunk_count"] > 0:
                    # File exists in Neo4j with chunks - it's ready and queryable
                    return {
                        "file_id": sanitized_file_id,
                        "status": "ready",
                        "filename": record["filename"] or "Unknown",
                        "size_bytes": record["size_bytes"] or 0,
                        "chunk_count": record["chunk_count"]
                    }

        # File not in Neo4j yet - check if it matches expected pattern for processing
        # If file_id format matches expected pattern (chat_id_filename_timestamp), assume processing
        if sanitized_file_id.startswith(sanitized_chat_id + "_"):
            # Extract filename from file_id to provide better response
            parts = sanitized_file_id.split("_")
            if len(parts) >= 3:
                # Reconstruct filename (everything except last timestamp part)
                filename_part = "_".join(parts[1:-1])
                return {
                    "file_id": sanitized_file_id,
                    "status": "processing",
                    "filename": filename_part,
                    "message": "File is being processed"
                }

        # Check Redis queue for job status
        queue = get_file_queue()
        if queue is not None:
            try:
                from rq.job import Job
                # Try to find job by checking if file_queue_id exists
                file_queue_id = f"fq_{sanitized_file_id}"

                # Check registries for job status
                started_jobs = queue.started_job_registry.get_job_ids()
                for job_id in started_jobs:
                    try:
                        job = Job.fetch(job_id, connection=queue.connection)
                        if job.args and len(job.args) >= 1 and job.args[0] == file_queue_id:
                            return {
                                "file_id": sanitized_file_id,
                                "status": "processing",
                                "message": "File is being processed by worker"
                            }
                    except Exception:
                        pass

                # Check if job is queued
                queued_job_ids = queue.get_job_ids()
                for job_id in queued_job_ids:
                    try:
                        job = Job.fetch(job_id, connection=queue.connection)
                        if job.args and len(job.args) >= 1 and job.args[0] == file_queue_id:
                            return {
                                "file_id": sanitized_file_id,
                                "status": "queued",
                                "message": "File is queued for processing"
                            }
                    except Exception:
                        pass

                # Check failed jobs
                failed_jobs = queue.failed_job_registry.get_job_ids()
                for job_id in failed_jobs:
                    try:
                        job = Job.fetch(job_id, connection=queue.connection)
                        if job.args and len(job.args) >= 1 and job.args[0] == file_queue_id:
                            return {
                                "file_id": sanitized_file_id,
                                "status": "failed",
                                "error": str(job.exc_info) if job.exc_info else "Processing failed",
                                "message": "File processing failed"
                            }
                    except Exception:
                        pass

            except Exception as e:
                logger.warning(f"Could not check queue status: {e}")

        # File not found and doesn't match expected pattern
        raise HTTPException(status_code=404, detail=f"File {sanitized_file_id} not found")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get file status for {file_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get file status: {str(e)}")

@app.get("/v1/{chat_id}/files", response_model=FileListResponse)
async def api_list_files(
    chat_id: str,
    request: Request,
    credentials: HTTPAuthorizationCredentials = Depends(get_verified_chat_access)
):
    """
    External API endpoint for listing files in a chat's knowledge base.

    Requires a valid API key for the specific chat.

    Usage:
    GET https://api.trainlyai.com/v1/{chat_id}/files
    Authorization: Bearer tk_your_api_key
    """
    try:
        # Verify API key and chat access
        api_key = credentials.credentials

        # Sanitize chat_id
        sanitized_chat_id = sanitize_chat_id(chat_id)
        if not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid chat_id format")

        # Verify API key and chat access
        is_valid = await verify_api_key_and_chat(api_key, sanitized_chat_id)

        if not is_valid:
            raise HTTPException(
                status_code=401,
                detail="Invalid API key or chat not accessible. Please check: 1) Chat exists, 2) API access is enabled in chat settings, 3) API key is correct."
            )

        # Query Neo4j for all documents in the chat
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # Get all documents with their metadata
                query = """
                MATCH (d:Document)-[:HAS_CHUNK]->(c:Chunk)
                WHERE c.chatId = $chat_id
                WITH d, count(c) as chunk_count
                RETURN d.id as file_id, d.filename as filename,
                       d.uploadDate as upload_date, d.sizeBytes as size_bytes,
                       chunk_count
                ORDER BY d.uploadDate DESC
                """

                result = session.run(query, chat_id=sanitized_chat_id)

                files = []
                total_size = 0

                for record in result:
                    # Convert upload_date from timestamp to string if it's a number
                    upload_date_value = record["upload_date"]
                    if isinstance(upload_date_value, (int, float)):
                        # Convert timestamp to ISO format string
                        upload_date_str = datetime.fromtimestamp(upload_date_value / 1000).isoformat()
                    else:
                        upload_date_str = str(upload_date_value) if upload_date_value else "Unknown"

                    file_info = FileInfo(
                        file_id=record["file_id"],
                        filename=record["filename"] or "Unknown",
                        upload_date=upload_date_str,
                        size_bytes=record["size_bytes"] or 0,
                        chunk_count=record["chunk_count"]
                    )
                    files.append(file_info)
                    total_size += file_info.size_bytes

                logger.info(f"üìã Listed {len(files)} files for chat {sanitized_chat_id}")

                return FileListResponse(
                    success=True,
                    files=files,
                    total_files=len(files),
                    total_size_bytes=total_size
                )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to list files for chat {chat_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to list files: {str(e)}")

@app.delete("/v1/{chat_id}/files/{file_id}", response_model=FileDeleteResponse)
async def api_delete_file(
    chat_id: str,
    file_id: str,
    request: Request,
    credentials: HTTPAuthorizationCredentials = Depends(get_verified_chat_access)
):
    """
    External API endpoint for deleting a file from a chat's knowledge base.

    Requires a valid API key for the specific chat.

    Usage:
    DELETE https://api.trainlyai.com/v1/{chat_id}/files/{file_id}
    Authorization: Bearer tk_your_api_key
    """
    try:
        # Verify API key and chat access
        api_key = credentials.credentials

        # Sanitize chat_id and file_id
        sanitized_chat_id = sanitize_chat_id(chat_id)
        if not sanitized_chat_id:
            raise HTTPException(status_code=400, detail="Invalid chat_id format")

        sanitized_file_id = sanitize_text(file_id)
        if not sanitized_file_id:
            raise HTTPException(status_code=400, detail="Invalid file_id format")

        # Verify API key and chat access
        is_valid = await verify_api_key_and_chat(api_key, sanitized_chat_id)

        if not is_valid:
            raise HTTPException(
                status_code=401,
                detail="Invalid API key or chat not accessible. Please check: 1) Chat exists, 2) API access is enabled in chat settings, 3) API key is correct."
            )

        # Delete from Neo4j and get metadata
        with GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password)) as driver:
            with driver.session() as session:
                # First, get file info before deletion
                info_query = """
                MATCH (d:Document {id: $file_id})-[:HAS_CHUNK]->(c:Chunk)
                WHERE c.chatId = $chat_id
                WITH d, count(c) as chunk_count
                RETURN d.filename as filename, d.sizeBytes as size_bytes, chunk_count
                """

                info_result = session.run(info_query, file_id=sanitized_file_id, chat_id=sanitized_chat_id)
                file_info = info_result.single()

                if not file_info:
                    raise HTTPException(
                        status_code=404,
                        detail=f"File {sanitized_file_id} not found in chat"
                    )

                filename = file_info["filename"] or "Unknown"
                size_bytes = file_info["size_bytes"] or 0
                chunk_count = file_info["chunk_count"]

                # Now delete the document and all its chunks
                delete_query = """
                MATCH (d:Document {id: $file_id})-[r:HAS_CHUNK]->(c:Chunk)
                WHERE c.chatId = $chat_id
                DETACH DELETE d, c
                """

                delete_result = session.run(delete_query, file_id=sanitized_file_id, chat_id=sanitized_chat_id)
                summary = delete_result.consume()

                if summary.counters.nodes_deleted == 0:
                    raise HTTPException(
                        status_code=404,
                        detail=f"No document found with id: {sanitized_file_id}"
                    )

                logger.info(f"üóëÔ∏è Deleted file {filename} ({size_bytes} bytes, {chunk_count} chunks) from chat {sanitized_chat_id}")

                return FileDeleteResponse(
                    success=True,
                    message=f"File '{filename}' deleted successfully",
                    file_id=sanitized_file_id,
                    filename=filename,
                    chunks_deleted=chunk_count,
                    size_bytes_freed=size_bytes
                )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete file {file_id} from chat {chat_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete file: {str(e)}")

def run_worker():
    """
    Run the RQ worker for background file processing.

    Usage:
        python read_files.py worker

    This starts a worker that processes file ingestion jobs from the Redis queue.
    Scale by running multiple worker processes.
    """
    if not REDIS_AVAILABLE:
        print("‚ùå Redis/RQ not installed. Install with: pip install redis rq")
        sys.exit(1)

    conn = get_redis_connection()
    if conn is None:
        print("‚ùå Could not connect to Redis. Check REDIS_URL environment variable.")
        sys.exit(1)

    print("üöÄ Starting Trainly File Processing Worker...")
    print(f"üì° Connected to Redis: {REDIS_URL[:50]}...")
    print("üìã Listening on queue: file_ingestion")
    print("‚è±Ô∏è  Job timeout: 30 minutes")
    print("üí° Press Ctrl+C to stop")
    print("")

    try:
        import platform
        from rq import Worker, SimpleWorker

        # Use SimpleWorker on macOS to avoid fork() issues with Objective-C runtime
        if platform.system() == "Darwin":
            print("üçé macOS detected - using SimpleWorker (no forking)")
            worker = SimpleWorker(
                queues=['file_ingestion'],
                connection=conn,
                name=f"trainly_worker_{os.getpid()}"
            )
        else:
            worker = Worker(
                queues=['file_ingestion'],
                connection=conn,
                name=f"trainly_worker_{os.getpid()}"
            )

        worker.work(with_scheduler=False)
    except KeyboardInterrupt:
        print("\nüëã Worker stopped gracefully")
    except Exception as e:
        print(f"‚ùå Worker error: {e}")
        sys.exit(1)

def show_queue_status():
    """Show the current status of the file processing queue."""
    if not REDIS_AVAILABLE:
        print("‚ùå Redis/RQ not installed")
        return

    conn = get_redis_connection()
    if conn is None:
        print("‚ùå Could not connect to Redis")
        return

    from rq import Queue
    from rq.job import Job

    queue = Queue("file_ingestion", connection=conn)

    print("üìä File Ingestion Queue Status")
    print("=" * 40)
    print(f"Pending jobs:    {len(queue)}")
    print(f"Failed jobs:     {queue.failed_job_registry.count}")
    print(f"Finished jobs:   {queue.finished_job_registry.count}")
    print(f"Started jobs:    {queue.started_job_registry.count}")
    print(f"Deferred jobs:   {queue.deferred_job_registry.count}")
    print("")

    # Show recent jobs
    if len(queue) > 0:
        print("üìã Pending Jobs:")
        for job_id in queue.job_ids[:10]:  # Show first 10
            try:
                job = Job.fetch(job_id, connection=conn)
                print(f"  - {job_id[:20]}... | {job.func_name} | queued at {job.enqueued_at}")
            except Exception as e:
                print(f"  - {job_id[:20]}... | error: {e}")

if __name__ == "__main__":
    # Check for CLI commands
    if len(sys.argv) > 1:
        command = sys.argv[1].lower()

        if command == "worker":
            run_worker()
        elif command == "status":
            show_queue_status()
        elif command == "help":
            print("""
Trainly Backend CLI

Usage:
    python read_files.py              Start the FastAPI server
    python read_files.py worker       Start a background worker
    python read_files.py status       Show queue status
    python read_files.py help         Show this help message

Environment Variables:
    REDIS_URL                         Redis connection URL (required for workers)
    OPENAI_API_KEY                    OpenAI API key
    NEO4J_URI                         Neo4j database URI
    NEO4J_USER                        Neo4j username
    NEO4J_PASSWORD                    Neo4j password
    CONVEX_URL                        Convex deployment URL

Example:
    # Terminal 1: Start the API server
    python read_files.py

    # Terminal 2: Start a worker (can run multiple)
    python read_files.py worker
""")
        else:
            print(f"Unknown command: {command}")
            print("Run 'python read_files.py help' for usage")
            sys.exit(1)
    else:
        # Default: run the FastAPI server
        import uvicorn

        print("üöÄ Starting TeachAI GraphRAG API with Privacy-First Architecture...")
        print("üì° Internal endpoints: http://localhost:8000")
        print("üåê Legacy API endpoints:")
        print("   POST /v1/{chat_id}/answer_question")
        print("   POST /v1/{chat_id}/answer_question_stream")
        print("   GET  /v1/{chat_id}/info")
        print("   GET  /v1/health")
        print("üîí Trainly OAuth API endpoints:")
        print("   GET  /oauth/authorize                  (OAuth authorization)")
        print("   POST /oauth/token                      (Token exchange)")
        print("   POST /me/chats/query                   (User-controlled queries)")
        print("   POST /me/chats/files/presign           (Private file uploads)")
        print("   GET  /v1/privacy/health                (API health status)")
        print("üìö API Documentation: http://localhost:8000/docs")
        print("üîê Authentication: Trainly OAuth 2.0 (User-controlled tokens)")
        print("üõ°Ô∏è  Privacy Guarantee: Citations filtered, no raw file access for developers")
        print("üí° Simple OAuth: Lightweight implementation, complete privacy control")

        # Check Redis connection status
        if REDIS_AVAILABLE:
            conn = get_redis_connection()
            if conn:
                print("‚úÖ Redis connected - background file processing enabled")
                print("üí° Run 'python read_files.py worker' in another terminal to start processing")
            else:
                print("‚ö†Ô∏è Redis not connected - files will process synchronously")
        else:
            print("‚ö†Ô∏è Redis not installed - files will process synchronously")

        uvicorn.run(app, host="0.0.0.0", port=8000)